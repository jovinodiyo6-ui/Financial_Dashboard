from pathlib import Path
from flask import Flask, request, jsonify
from flask_sqlalchemy import SQLAlchemy
from flask_jwt_extended import JWTManager, create_access_token, jwt_required, get_jwt_identity
from flask_bcrypt import Bcrypt
from flask_cors import CORS
from sqlalchemy import text
from sqlalchemy.exc import IntegrityError, OperationalError, SQLAlchemyError
from functools import wraps
import hashlib
import datetime
import json
import os
import secrets
import base64
import smtplib
import pandas as pd
import re
import urllib.request
import urllib.error
from dotenv import load_dotenv
from email.message import EmailMessage
from io import StringIO

try:
    import pdfplumber
except Exception:  # pragma: no cover - optional dependency
    pdfplumber = None

try:
    import docx
except Exception:  # pragma: no cover - optional dependency
    docx = None

try:
    from flask_limiter import Limiter
    from flask_limiter.util import get_remote_address
except Exception:  # pragma: no cover - runtime fallback for missing optional dependency
    Limiter = None

    def get_remote_address():
        return "0.0.0.0"

try:
    import stripe
except Exception:  # pragma: no cover - optional dependency
    stripe = None

try:
    from celery import Celery, Task
except Exception:  # pragma: no cover - optional dependency
    Celery = None
    Task = object

load_dotenv()

app = Flask(__name__)

allowed_origins = os.getenv("CORS_ORIGINS", "*")
CORS(app, origins=[o.strip() for o in allowed_origins.split(",") if o.strip()] if allowed_origins != "*" else "*")

# ---------------- CONFIG ----------------

# Support both SQLite (dev) and PostgreSQL (production)
DATABASE_URL = os.getenv("DATABASE_URL", "sqlite:///saas.db")
if DATABASE_URL.startswith("postgres://"):
    DATABASE_URL = DATABASE_URL.replace("postgres://", "postgresql://", 1)
if DATABASE_URL.startswith("postgresql://"):
    DATABASE_URL = DATABASE_URL.replace("postgresql://", "postgresql+psycopg://", 1)
if DATABASE_URL.startswith("postgresql://") and "sslmode=" not in DATABASE_URL:
    separator = "&" if "?" in DATABASE_URL else "?"
    DATABASE_URL = f"{DATABASE_URL}{separator}sslmode=require"
if DATABASE_URL.startswith("postgresql+psycopg://") and "sslmode=" not in DATABASE_URL:
    separator = "&" if "?" in DATABASE_URL else "?"
    DATABASE_URL = f"{DATABASE_URL}{separator}sslmode=require"

JWT_SECRET = os.getenv("JWT_SECRET_KEY", "")
ENV = os.getenv("FLASK_ENV", "development")

if ENV == "production" and not JWT_SECRET:
    print("WARNING: JWT_SECRET_KEY is missing in production; using development fallback secret.")
if ENV == "production" and DATABASE_URL.startswith("sqlite"):
    print("WARNING: production DATABASE_URL is SQLite. Configure PostgreSQL for reliability.")

app.config["SQLALCHEMY_DATABASE_URI"] = DATABASE_URL
app.config["JWT_SECRET_KEY"] = JWT_SECRET or "dev-only-secret-change-me"
app.config["SECRET_KEY"] = os.getenv("FLASK_SECRET_KEY", JWT_SECRET or "dev-flask-secret-change-me")
app.config["JWT_ACCESS_TOKEN_EXPIRES"] = datetime.timedelta(days=7)
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["MAX_CONTENT_LENGTH"] = 5 * 1024 * 1024  # 5 MB upload cap
app.config["SESSION_COOKIE_SECURE"] = ENV == "production"
app.config["SESSION_COOKIE_HTTPONLY"] = True
app.config["SESSION_COOKIE_SAMESITE"] = "Lax"
app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {
    "pool_pre_ping": True,
    "pool_recycle": 300,
    "pool_timeout": 20,
    "pool_size": 5,
    "max_overflow": 5,
}

db = SQLAlchemy(app)
jwt = JWTManager(app)
bcrypt = Bcrypt(app)
if Limiter:
    limiter = Limiter(
        key_func=get_remote_address,
        default_limits=[],
        storage_uri=os.getenv("RATE_LIMIT_STORAGE_URI", "memory://"),
    )
    limiter.init_app(app)
else:
    class _NoopLimiter:
        def limit(self, _rule):
            def _decorator(fn):
                return fn

            return _decorator

    limiter = _NoopLimiter()

STRIPE_SECRET_KEY = os.getenv("STRIPE_SECRET_KEY", "").strip()
STRIPE_WEBHOOK_SECRET = os.getenv("STRIPE_WEBHOOK_SECRET", "").strip()
STRIPE_PRICE_IDS = {
    "pro": os.getenv("STRIPE_PRICE_PRO", "").strip(),
    "ai": os.getenv("STRIPE_PRICE_AI", "").strip(),
}
REDIS_URL = os.getenv("REDIS_URL", "redis://localhost:6379/0")
ASYNC_JOB_MODE = os.getenv("ASYNC_JOB_MODE", "inline" if ENV != "production" else "celery").strip().lower()
MPESA_ENV = os.getenv("MPESA_ENV", "sandbox").strip().lower()
MPESA_CONSUMER_KEY = os.getenv("MPESA_CONSUMER_KEY", "").strip()
MPESA_CONSUMER_SECRET = os.getenv("MPESA_CONSUMER_SECRET", "").strip()
MPESA_SHORTCODE = os.getenv("MPESA_SHORTCODE", "").strip()
MPESA_PASSKEY = os.getenv("MPESA_PASSKEY", "").strip()
MPESA_CALLBACK_URL = os.getenv("MPESA_CALLBACK_URL", "").strip()
MPESA_INITIATOR_NAME = os.getenv("MPESA_INITIATOR_NAME", "").strip()
MPESA_SECURITY_CREDENTIAL = os.getenv("MPESA_SECURITY_CREDENTIAL", "").strip()
MPESA_BASE_URL = "https://api.safaricom.co.ke" if MPESA_ENV == "production" else "https://sandbox.safaricom.co.ke"

if stripe and STRIPE_SECRET_KEY:
    stripe.api_key = STRIPE_SECRET_KEY

celery_app = None
if Celery:
    class FlaskTask(Task):
        def __call__(self, *args, **kwargs):
            with app.app_context():
                return self.run(*args, **kwargs)

    celery_app = Celery(
        app.import_name,
        broker=os.getenv("CELERY_BROKER_URL", REDIS_URL),
        backend=os.getenv("CELERY_RESULT_BACKEND", REDIS_URL),
        task_cls=FlaskTask,
    )
    celery_app.conf.update(
        task_serializer="json",
        accept_content=["json"],
        result_serializer="json",
        timezone="UTC",
        enable_utc=True,
    )

# ---------------- DATABASE ----------------

class Organization(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    usage = db.Column(db.Integer, default=0)
    billing_email = db.Column(db.String(120), nullable=True)
    plan_code = db.Column(db.String(20), nullable=False, default="free")
    subscription_status = db.Column(db.String(20), nullable=False, default="free")
    stripe_customer_id = db.Column(db.String(120), nullable=True)
    stripe_subscription_id = db.Column(db.String(120), nullable=True)
    max_companies = db.Column(db.Integer, nullable=False, default=1)
    ai_assistant_enabled = db.Column(db.Boolean, nullable=False, default=False)
    subscription_updated_at = db.Column(db.DateTime(timezone=True), nullable=True)


class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password = db.Column(db.String(200), nullable=False)
    role = db.Column(db.String(20), nullable=False)
    org_id = db.Column(db.Integer, nullable=False)
    default_company_id = db.Column(db.Integer, nullable=True)


class Report(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    org_id = db.Column(db.Integer, nullable=False)
    company_id = db.Column(db.Integer, nullable=True)
    data = db.Column(db.Text, nullable=False)


class Company(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    org_id = db.Column(db.Integer, nullable=False)
    name = db.Column(db.String(120), nullable=False)
    business_type = db.Column(db.String(50), nullable=False, default="sole_proprietor")


class UserCompanyMembership(db.Model):
    __table_args__ = (db.UniqueConstraint("user_id", "company_id", name="uq_user_company_membership"),)

    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, nullable=False)
    company_id = db.Column(db.Integer, nullable=False)
    role = db.Column(db.String(20), nullable=False, default="member")
    is_default = db.Column(db.Boolean, nullable=False, default=False)
    created_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )


class CompanyPartner(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    company_id = db.Column(db.Integer, nullable=False)
    name = db.Column(db.String(120), nullable=False)
    display_order = db.Column(db.Integer, nullable=False, default=1)


class CompanyOnboardingState(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    company_id = db.Column(db.Integer, unique=True, nullable=False)
    is_configured = db.Column(db.Boolean, nullable=False, default=False)
    configured_at = db.Column(db.DateTime(timezone=True), nullable=True)


class APIKey(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    org_id = db.Column(db.Integer, nullable=False)
    key_hash = db.Column(db.String(200), nullable=False)


class AuditLog(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, nullable=False)
    company_id = db.Column(db.Integer, nullable=True)
    action = db.Column(db.String(200), nullable=False)
    time = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
    )


class ActiveSession(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, unique=True, nullable=False)
    last_seen = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )


class PasswordResetToken(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, nullable=False)
    token_hash = db.Column(db.String(64), unique=True, nullable=False)
    expires_at = db.Column(db.DateTime(timezone=True), nullable=False)
    used_at = db.Column(db.DateTime(timezone=True), nullable=True)
    created_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )


class BackgroundJob(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    org_id = db.Column(db.Integer, nullable=False)
    company_id = db.Column(db.Integer, nullable=False)
    requested_by = db.Column(db.Integer, nullable=False)
    job_type = db.Column(db.String(40), nullable=False)
    status = db.Column(db.String(20), nullable=False, default="queued")
    provider = db.Column(db.String(20), nullable=False, default="inline")
    payload_json = db.Column(db.Text, nullable=True)
    result_json = db.Column(db.Text, nullable=True)
    error_message = db.Column(db.Text, nullable=True)
    started_at = db.Column(db.DateTime(timezone=True), nullable=True)
    completed_at = db.Column(db.DateTime(timezone=True), nullable=True)
    created_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )


class BillingPaymentRequest(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    org_id = db.Column(db.Integer, nullable=False)
    company_id = db.Column(db.Integer, nullable=False)
    requested_by = db.Column(db.Integer, nullable=False)
    provider = db.Column(db.String(20), nullable=False, default="mpesa")
    plan_code = db.Column(db.String(20), nullable=False)
    currency_code = db.Column(db.String(8), nullable=False, default="KES")
    amount = db.Column(db.Float, nullable=False, default=0.0)
    phone_number = db.Column(db.String(30), nullable=True)
    external_reference = db.Column(db.String(120), nullable=True)
    merchant_request_id = db.Column(db.String(120), nullable=True)
    checkout_request_id = db.Column(db.String(120), nullable=True)
    status = db.Column(db.String(20), nullable=False, default="pending")
    provider_response_json = db.Column(db.Text, nullable=True)
    callback_payload_json = db.Column(db.Text, nullable=True)
    created_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )
    updated_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        onupdate=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )


class Invoice(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    org_id = db.Column(db.Integer, nullable=False)
    company_id = db.Column(db.Integer, nullable=False)
    invoice_number = db.Column(db.String(40), nullable=False, unique=True)
    customer_name = db.Column(db.String(120), nullable=False)
    customer_email = db.Column(db.String(120), nullable=True)
    status = db.Column(db.String(20), nullable=False, default="draft")
    issue_date = db.Column(db.Date, nullable=False)
    due_date = db.Column(db.Date, nullable=False)
    subtotal = db.Column(db.Float, nullable=False, default=0.0)
    tax_rate = db.Column(db.Float, nullable=False, default=0.0)
    tax_amount = db.Column(db.Float, nullable=False, default=0.0)
    total_amount = db.Column(db.Float, nullable=False, default=0.0)
    balance_due = db.Column(db.Float, nullable=False, default=0.0)
    notes = db.Column(db.Text, nullable=True)
    created_by = db.Column(db.Integer, nullable=False)
    last_sent_at = db.Column(db.DateTime(timezone=True), nullable=True)
    paid_at = db.Column(db.DateTime(timezone=True), nullable=True)
    created_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )


class InvoiceItem(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    invoice_id = db.Column(db.Integer, nullable=False)
    description = db.Column(db.String(200), nullable=False)
    quantity = db.Column(db.Float, nullable=False, default=1.0)
    unit_price = db.Column(db.Float, nullable=False, default=0.0)
    amount = db.Column(db.Float, nullable=False, default=0.0)


class CustomerPayment(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    org_id = db.Column(db.Integer, nullable=False)
    company_id = db.Column(db.Integer, nullable=False)
    invoice_id = db.Column(db.Integer, nullable=False)
    amount = db.Column(db.Float, nullable=False, default=0.0)
    reference = db.Column(db.String(120), nullable=True)
    source = db.Column(db.String(30), nullable=False, default="manual")
    notes = db.Column(db.Text, nullable=True)
    payment_date = db.Column(db.Date, nullable=False)
    bank_transaction_id = db.Column(db.Integer, nullable=True)
    created_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )


class VendorBill(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    org_id = db.Column(db.Integer, nullable=False)
    company_id = db.Column(db.Integer, nullable=False)
    bill_number = db.Column(db.String(40), nullable=False, unique=True)
    vendor_name = db.Column(db.String(120), nullable=False)
    status = db.Column(db.String(20), nullable=False, default="draft")
    issue_date = db.Column(db.Date, nullable=False)
    due_date = db.Column(db.Date, nullable=False)
    subtotal = db.Column(db.Float, nullable=False, default=0.0)
    tax_rate = db.Column(db.Float, nullable=False, default=0.0)
    tax_amount = db.Column(db.Float, nullable=False, default=0.0)
    total_amount = db.Column(db.Float, nullable=False, default=0.0)
    balance_due = db.Column(db.Float, nullable=False, default=0.0)
    notes = db.Column(db.Text, nullable=True)
    created_by = db.Column(db.Integer, nullable=False)
    approved_at = db.Column(db.DateTime(timezone=True), nullable=True)
    paid_at = db.Column(db.DateTime(timezone=True), nullable=True)
    created_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )


class VendorBillItem(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    bill_id = db.Column(db.Integer, nullable=False)
    description = db.Column(db.String(200), nullable=False)
    quantity = db.Column(db.Float, nullable=False, default=1.0)
    unit_price = db.Column(db.Float, nullable=False, default=0.0)
    amount = db.Column(db.Float, nullable=False, default=0.0)


class VendorPayment(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    org_id = db.Column(db.Integer, nullable=False)
    company_id = db.Column(db.Integer, nullable=False)
    bill_id = db.Column(db.Integer, nullable=False)
    amount = db.Column(db.Float, nullable=False, default=0.0)
    reference = db.Column(db.String(120), nullable=True)
    source = db.Column(db.String(30), nullable=False, default="manual")
    notes = db.Column(db.Text, nullable=True)
    payment_date = db.Column(db.Date, nullable=False)
    bank_transaction_id = db.Column(db.Integer, nullable=True)
    created_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )


class BankFeedTransaction(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    org_id = db.Column(db.Integer, nullable=False)
    company_id = db.Column(db.Integer, nullable=False)
    posted_at = db.Column(db.Date, nullable=False)
    description = db.Column(db.String(255), nullable=False)
    amount = db.Column(db.Float, nullable=False, default=0.0)
    reference = db.Column(db.String(120), nullable=True)
    status = db.Column(db.String(20), nullable=False, default="unmatched")
    matched_invoice_id = db.Column(db.Integer, nullable=True)
    matched_bill_id = db.Column(db.Integer, nullable=True)
    raw_payload = db.Column(db.Text, nullable=True)
    created_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )


class BankConnection(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    org_id = db.Column(db.Integer, nullable=False)
    company_id = db.Column(db.Integer, nullable=False)
    provider = db.Column(db.String(30), nullable=False, default="plaid")
    item_id = db.Column(db.String(120), nullable=False, unique=True)
    access_token = db.Column(db.String(255), nullable=False)
    institution_name = db.Column(db.String(120), nullable=True)
    status = db.Column(db.String(20), nullable=False, default="connected")
    sync_cursor = db.Column(db.String(255), nullable=True)
    created_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )
    updated_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        onupdate=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )


class TaxProfile(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    org_id = db.Column(db.Integer, nullable=False)
    company_id = db.Column(db.Integer, nullable=False, unique=True)
    jurisdiction_code = db.Column(db.String(40), nullable=False, default="generic")
    filing_frequency = db.Column(db.String(20), nullable=False, default="monthly")
    registration_number = db.Column(db.String(80), nullable=True)
    currency_code = db.Column(db.String(8), nullable=False, default="USD")
    sales_tax_name = db.Column(db.String(40), nullable=False, default="Sales Tax")
    purchase_tax_name = db.Column(db.String(40), nullable=False, default="Purchase Tax Credit")
    indirect_tax_rate = db.Column(db.Float, nullable=False, default=16.0)
    income_tax_rate = db.Column(db.Float, nullable=False, default=30.0)
    period_start_month = db.Column(db.Integer, nullable=False, default=1)
    created_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )
    updated_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        onupdate=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )


class LedgerAccount(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    org_id = db.Column(db.Integer, nullable=False)
    company_id = db.Column(db.Integer, nullable=False)
    code = db.Column(db.String(20), nullable=False)
    name = db.Column(db.String(120), nullable=False)
    category = db.Column(db.String(20), nullable=False)
    subtype = db.Column(db.String(40), nullable=True)
    normal_balance = db.Column(db.String(10), nullable=False, default="debit")
    description = db.Column(db.Text, nullable=True)
    is_system = db.Column(db.Boolean, nullable=False, default=False)
    is_active = db.Column(db.Boolean, nullable=False, default=True)
    created_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )
    updated_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        onupdate=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )


class JournalEntry(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    org_id = db.Column(db.Integer, nullable=False)
    company_id = db.Column(db.Integer, nullable=False)
    entry_number = db.Column(db.String(40), nullable=False, unique=True)
    entry_date = db.Column(db.Date, nullable=False)
    memo = db.Column(db.String(255), nullable=False)
    reference = db.Column(db.String(120), nullable=True)
    source_type = db.Column(db.String(40), nullable=False, default="manual")
    source_id = db.Column(db.Integer, nullable=True)
    status = db.Column(db.String(20), nullable=False, default="posted")
    reverses_entry_id = db.Column(db.Integer, nullable=True)
    created_by = db.Column(db.Integer, nullable=False)
    created_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )


class JournalLine(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    journal_entry_id = db.Column(db.Integer, nullable=False)
    account_id = db.Column(db.Integer, nullable=False)
    project_id = db.Column(db.Integer, nullable=True)
    line_number = db.Column(db.Integer, nullable=False, default=1)
    description = db.Column(db.String(255), nullable=True)
    debit = db.Column(db.Float, nullable=False, default=0.0)
    credit = db.Column(db.Float, nullable=False, default=0.0)


class VendorProfile(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    org_id = db.Column(db.Integer, nullable=False)
    company_id = db.Column(db.Integer, nullable=False)
    vendor_name = db.Column(db.String(120), nullable=False)
    email = db.Column(db.String(120), nullable=True)
    tax_id = db.Column(db.String(80), nullable=True)
    default_payment_rail = db.Column(db.String(30), nullable=False, default="ach")
    remittance_reference = db.Column(db.String(120), nullable=True)
    bank_last4 = db.Column(db.String(4), nullable=True)
    is_1099_eligible = db.Column(db.Boolean, nullable=False, default=False)
    tax_form_type = db.Column(db.String(20), nullable=False, default="1099-NEC")
    tin_status = db.Column(db.String(20), nullable=False, default="pending")
    created_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )
    updated_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        onupdate=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )


class BillDisbursement(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    org_id = db.Column(db.Integer, nullable=False)
    company_id = db.Column(db.Integer, nullable=False)
    bill_id = db.Column(db.Integer, nullable=False)
    vendor_profile_id = db.Column(db.Integer, nullable=True)
    payment_rail = db.Column(db.String(30), nullable=False, default="ach")
    status = db.Column(db.String(20), nullable=False, default="scheduled")
    scheduled_date = db.Column(db.Date, nullable=False)
    processed_at = db.Column(db.DateTime(timezone=True), nullable=True)
    amount = db.Column(db.Float, nullable=False, default=0.0)
    reference = db.Column(db.String(120), nullable=True)
    confirmation_code = db.Column(db.String(40), nullable=True)
    compliance_status = db.Column(db.String(20), nullable=False, default="ready")
    created_by = db.Column(db.Integer, nullable=False)
    created_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )


class ReconciliationRule(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    org_id = db.Column(db.Integer, nullable=False)
    company_id = db.Column(db.Integer, nullable=False)
    name = db.Column(db.String(120), nullable=False)
    keyword = db.Column(db.String(120), nullable=True)
    direction = db.Column(db.String(20), nullable=False, default="any")
    min_amount = db.Column(db.Float, nullable=True)
    max_amount = db.Column(db.Float, nullable=True)
    auto_action = db.Column(db.String(30), nullable=False, default="suggest_account")
    target_reference = db.Column(db.String(80), nullable=True)
    exception_type = db.Column(db.String(40), nullable=True)
    priority = db.Column(db.Integer, nullable=False, default=100)
    is_active = db.Column(db.Boolean, nullable=False, default=True)
    created_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )


class ReconciliationException(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    org_id = db.Column(db.Integer, nullable=False)
    company_id = db.Column(db.Integer, nullable=False)
    bank_transaction_id = db.Column(db.Integer, nullable=False)
    exception_type = db.Column(db.String(40), nullable=False)
    notes = db.Column(db.Text, nullable=True)
    status = db.Column(db.String(20), nullable=False, default="open")
    created_by = db.Column(db.Integer, nullable=False)
    resolved_at = db.Column(db.DateTime(timezone=True), nullable=True)
    created_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )


class TaxFiling(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    org_id = db.Column(db.Integer, nullable=False)
    company_id = db.Column(db.Integer, nullable=False)
    jurisdiction_code = db.Column(db.String(40), nullable=False)
    filing_frequency = db.Column(db.String(20), nullable=False)
    filing_type = db.Column(db.String(30), nullable=False, default="indirect_tax")
    period_start = db.Column(db.Date, nullable=False)
    period_end = db.Column(db.Date, nullable=False)
    status = db.Column(db.String(20), nullable=False, default="prepared")
    reference = db.Column(db.String(80), nullable=True)
    payload_json = db.Column(db.Text, nullable=True)
    prepared_by = db.Column(db.Integer, nullable=False)
    prepared_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )
    submitted_at = db.Column(db.DateTime(timezone=True), nullable=True)
    created_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )


class EmployeeProfile(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    org_id = db.Column(db.Integer, nullable=False)
    company_id = db.Column(db.Integer, nullable=False)
    full_name = db.Column(db.String(120), nullable=False)
    email = db.Column(db.String(120), nullable=True)
    pay_type = db.Column(db.String(20), nullable=False, default="hourly")
    hourly_rate = db.Column(db.Float, nullable=False, default=0.0)
    salary_amount = db.Column(db.Float, nullable=False, default=0.0)
    withholding_rate = db.Column(db.Float, nullable=False, default=0.0)
    benefit_rate = db.Column(db.Float, nullable=False, default=0.0)
    is_active = db.Column(db.Boolean, nullable=False, default=True)
    created_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )
    updated_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        onupdate=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )


class ContractorProfile(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    org_id = db.Column(db.Integer, nullable=False)
    company_id = db.Column(db.Integer, nullable=False)
    full_name = db.Column(db.String(120), nullable=False)
    email = db.Column(db.String(120), nullable=True)
    tax_id = db.Column(db.String(80), nullable=True)
    default_rate = db.Column(db.Float, nullable=False, default=0.0)
    is_1099_eligible = db.Column(db.Boolean, nullable=False, default=True)
    tax_form_type = db.Column(db.String(20), nullable=False, default="1099-NEC")
    is_active = db.Column(db.Boolean, nullable=False, default=True)
    created_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )
    updated_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        onupdate=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )


class TimeEntry(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    org_id = db.Column(db.Integer, nullable=False)
    company_id = db.Column(db.Integer, nullable=False)
    employee_id = db.Column(db.Integer, nullable=True)
    contractor_id = db.Column(db.Integer, nullable=True)
    project_id = db.Column(db.Integer, nullable=True)
    work_date = db.Column(db.Date, nullable=False)
    hours = db.Column(db.Float, nullable=False, default=0.0)
    hourly_cost = db.Column(db.Float, nullable=False, default=0.0)
    billable_rate = db.Column(db.Float, nullable=False, default=0.0)
    description = db.Column(db.String(255), nullable=True)
    status = db.Column(db.String(20), nullable=False, default="submitted")
    created_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )


class MileageEntry(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    org_id = db.Column(db.Integer, nullable=False)
    company_id = db.Column(db.Integer, nullable=False)
    employee_id = db.Column(db.Integer, nullable=True)
    contractor_id = db.Column(db.Integer, nullable=True)
    project_id = db.Column(db.Integer, nullable=True)
    trip_date = db.Column(db.Date, nullable=False)
    miles = db.Column(db.Float, nullable=False, default=0.0)
    rate_per_mile = db.Column(db.Float, nullable=False, default=0.0)
    purpose = db.Column(db.String(255), nullable=True)
    status = db.Column(db.String(20), nullable=False, default="submitted")
    created_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )


class PayrollRun(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    org_id = db.Column(db.Integer, nullable=False)
    company_id = db.Column(db.Integer, nullable=False)
    payroll_number = db.Column(db.String(40), nullable=False, unique=True)
    period_start = db.Column(db.Date, nullable=False)
    period_end = db.Column(db.Date, nullable=False)
    pay_date = db.Column(db.Date, nullable=False)
    status = db.Column(db.String(20), nullable=False, default="processed")
    gross_pay = db.Column(db.Float, nullable=False, default=0.0)
    withholding_total = db.Column(db.Float, nullable=False, default=0.0)
    benefit_total = db.Column(db.Float, nullable=False, default=0.0)
    mileage_reimbursement_total = db.Column(db.Float, nullable=False, default=0.0)
    net_cash = db.Column(db.Float, nullable=False, default=0.0)
    created_by = db.Column(db.Integer, nullable=False)
    created_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )


class PayrollLine(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    payroll_run_id = db.Column(db.Integer, nullable=False)
    employee_id = db.Column(db.Integer, nullable=False)
    regular_hours = db.Column(db.Float, nullable=False, default=0.0)
    gross_pay = db.Column(db.Float, nullable=False, default=0.0)
    withholding_amount = db.Column(db.Float, nullable=False, default=0.0)
    benefit_amount = db.Column(db.Float, nullable=False, default=0.0)
    mileage_reimbursement = db.Column(db.Float, nullable=False, default=0.0)
    net_pay = db.Column(db.Float, nullable=False, default=0.0)


class InventoryItem(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    org_id = db.Column(db.Integer, nullable=False)
    company_id = db.Column(db.Integer, nullable=False)
    sku = db.Column(db.String(60), nullable=False)
    name = db.Column(db.String(120), nullable=False)
    category = db.Column(db.String(80), nullable=True)
    quantity_on_hand = db.Column(db.Float, nullable=False, default=0.0)
    reorder_point = db.Column(db.Float, nullable=False, default=0.0)
    reorder_quantity = db.Column(db.Float, nullable=False, default=0.0)
    unit_cost = db.Column(db.Float, nullable=False, default=0.0)
    unit_price = db.Column(db.Float, nullable=False, default=0.0)
    preferred_vendor_name = db.Column(db.String(120), nullable=True)
    created_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )
    updated_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        onupdate=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )


class PurchaseOrder(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    org_id = db.Column(db.Integer, nullable=False)
    company_id = db.Column(db.Integer, nullable=False)
    po_number = db.Column(db.String(40), nullable=False, unique=True)
    vendor_name = db.Column(db.String(120), nullable=False)
    status = db.Column(db.String(20), nullable=False, default="draft")
    issue_date = db.Column(db.Date, nullable=False)
    expected_date = db.Column(db.Date, nullable=True)
    notes = db.Column(db.Text, nullable=True)
    created_by = db.Column(db.Integer, nullable=False)
    created_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )


class PurchaseOrderLine(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    purchase_order_id = db.Column(db.Integer, nullable=False)
    inventory_item_id = db.Column(db.Integer, nullable=True)
    sku = db.Column(db.String(60), nullable=True)
    description = db.Column(db.String(200), nullable=False)
    quantity = db.Column(db.Float, nullable=False, default=0.0)
    unit_cost = db.Column(db.Float, nullable=False, default=0.0)
    received_quantity = db.Column(db.Float, nullable=False, default=0.0)


class InventoryMovement(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    org_id = db.Column(db.Integer, nullable=False)
    company_id = db.Column(db.Integer, nullable=False)
    inventory_item_id = db.Column(db.Integer, nullable=False)
    project_id = db.Column(db.Integer, nullable=True)
    movement_type = db.Column(db.String(30), nullable=False)
    quantity_delta = db.Column(db.Float, nullable=False, default=0.0)
    unit_cost = db.Column(db.Float, nullable=False, default=0.0)
    reference = db.Column(db.String(120), nullable=True)
    occurred_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )


class Project(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    org_id = db.Column(db.Integer, nullable=False)
    company_id = db.Column(db.Integer, nullable=False)
    project_code = db.Column(db.String(40), nullable=False, unique=True)
    name = db.Column(db.String(120), nullable=False)
    customer_name = db.Column(db.String(120), nullable=True)
    status = db.Column(db.String(20), nullable=False, default="active")
    budget_revenue = db.Column(db.Float, nullable=False, default=0.0)
    budget_cost = db.Column(db.Float, nullable=False, default=0.0)
    notes = db.Column(db.Text, nullable=True)
    created_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )
    updated_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        onupdate=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )


class ProjectCostEntry(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    org_id = db.Column(db.Integer, nullable=False)
    company_id = db.Column(db.Integer, nullable=False)
    project_id = db.Column(db.Integer, nullable=False)
    entry_type = db.Column(db.String(20), nullable=False, default="cost")
    description = db.Column(db.String(255), nullable=False)
    amount = db.Column(db.Float, nullable=False, default=0.0)
    reference = db.Column(db.String(120), nullable=True)
    work_date = db.Column(db.Date, nullable=False)
    created_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )


class IntegrationConnection(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    org_id = db.Column(db.Integer, nullable=False)
    company_id = db.Column(db.Integer, nullable=False)
    provider = db.Column(db.String(40), nullable=False)
    category = db.Column(db.String(40), nullable=False)
    status = db.Column(db.String(20), nullable=False, default="available")
    config_json = db.Column(db.Text, nullable=True)
    last_synced_at = db.Column(db.DateTime(timezone=True), nullable=True)
    created_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )
    updated_at = db.Column(
        db.DateTime(timezone=True),
        default=lambda: datetime.datetime.now(datetime.UTC),
        onupdate=lambda: datetime.datetime.now(datetime.UTC),
        nullable=False,
    )


with app.app_context():
    # Avoid multi-worker startup races in production. In production, schema
    # should be managed explicitly (migrations/init job), not at app import time.
    if ENV != "production":
        try:
            db.create_all()
            inspector = db.inspect(db.engine)

            def ensure_column(table_name, column_name, column_sql):
                existing_columns = {column["name"] for column in inspector.get_columns(table_name)}
                if column_name not in existing_columns:
                    db.session.execute(text(f"ALTER TABLE {table_name} ADD COLUMN {column_sql}"))
                    db.session.commit()

            ensure_column("report", "company_id", "company_id INTEGER")
            ensure_column("organization", "billing_email", "billing_email VARCHAR(120)")
            ensure_column("organization", "plan_code", "plan_code VARCHAR(20) DEFAULT 'free'")
            ensure_column("organization", "subscription_status", "subscription_status VARCHAR(20) DEFAULT 'free'")
            ensure_column("organization", "stripe_customer_id", "stripe_customer_id VARCHAR(120)")
            ensure_column("organization", "stripe_subscription_id", "stripe_subscription_id VARCHAR(120)")
            ensure_column("organization", "max_companies", "max_companies INTEGER DEFAULT 1")
            ensure_column("organization", "ai_assistant_enabled", "ai_assistant_enabled BOOLEAN DEFAULT 0")
            ensure_column("organization", "subscription_updated_at", "subscription_updated_at DATETIME")
            ensure_column("user", "default_company_id", "default_company_id INTEGER")
            ensure_column("audit_log", "company_id", "company_id INTEGER")
        except OperationalError as exc:
            if "already exists" not in str(exc).lower():
                raise


def run_schema_compatibility_sync():
    try:
        db.create_all()
        inspector = db.inspect(db.engine)

        def ensure_column(table_name, column_name, column_sql):
            existing_columns = {column["name"] for column in inspector.get_columns(table_name)}
            if column_name not in existing_columns:
                db.session.execute(text(f"ALTER TABLE {table_name} ADD COLUMN {column_sql}"))
                db.session.commit()

        ensure_column("report", "company_id", "company_id INTEGER")
        ensure_column("organization", "billing_email", "billing_email VARCHAR(120)")
        ensure_column("organization", "plan_code", "plan_code VARCHAR(20) DEFAULT 'free'")
        ensure_column("organization", "subscription_status", "subscription_status VARCHAR(20) DEFAULT 'free'")
        ensure_column("organization", "stripe_customer_id", "stripe_customer_id VARCHAR(120)")
        ensure_column("organization", "stripe_subscription_id", "stripe_subscription_id VARCHAR(120)")
        ensure_column("organization", "max_companies", "max_companies INTEGER DEFAULT 1")
        ensure_column("organization", "ai_assistant_enabled", "ai_assistant_enabled BOOLEAN DEFAULT 0")
        ensure_column("organization", "subscription_updated_at", "subscription_updated_at DATETIME")
        ensure_column("user", "default_company_id", "default_company_id INTEGER")
        ensure_column("audit_log", "company_id", "company_id INTEGER")
    except OperationalError as exc:
        if "already exists" not in str(exc).lower():
            raise

# ---------------- LIMITS ----------------

FREE_USAGE_LIMIT = 5
MAINTENANCE_DEFAULT_MESSAGE = "[System Under Maintainance]"
VALID_ROLES = {"owner", "admin", "accountant", "manager", "cashier", "member"}
VALID_MEMBERSHIP_ROLES = set(VALID_ROLES)
INVOICE_OPEN_STATUSES = {"draft", "sent", "partial", "overdue"}
BILL_OPEN_STATUSES = {"draft", "approved", "partial", "overdue"}
INVOICE_SETTLED_STATUSES = {"paid", "cancelled"}
BILL_SETTLED_STATUSES = {"paid", "cancelled"}
VALID_TAX_FILING_FREQUENCIES = {"monthly", "quarterly", "annual"}
VALID_ACCOUNT_CATEGORIES = {"asset", "liability", "equity", "revenue", "expense"}
VALID_PAYMENT_RAILS = {"ach", "wire", "card", "check", "mobile_money"}
VALID_RECONCILIATION_DIRECTIONS = {"any", "inflow", "outflow"}
VALID_RECONCILIATION_ACTIONS = {"suggest_account", "flag_exception"}
VALID_TAX_FILING_TYPES = {"indirect_tax", "income_tax", "payroll_tax"}
VALID_PAY_TYPES = {"hourly", "salary"}
VALID_INTEGRATION_STATUSES = {"available", "connected", "attention"}
VALID_BUSINESS_TYPES = {"sole_proprietor", "partnership", "manufacturing"}
PASSWORD_RESET_EXPIRY_MINUTES = int(os.getenv("PASSWORD_RESET_EXPIRY_MINUTES", "30"))
JOB_TYPES = {"finance_digest", "tax_filing_package", "accountant_brief"}
JOB_TERMINAL_STATUSES = {"completed", "failed"}
ROLE_PERMISSION_MAP = {
    "owner": {"company:view", "company:manage", "user:manage", "billing:manage", "finance:manage", "jobs:run", "ai:ask"},
    "admin": {"company:view", "company:manage", "user:manage", "billing:manage", "finance:manage", "jobs:run", "ai:ask"},
    "accountant": {"company:view", "finance:manage", "jobs:run", "ai:ask"},
    "manager": {"company:view", "finance:manage", "jobs:run", "ai:ask"},
    "cashier": {"company:view", "finance:manage"},
    "member": {"company:view"},
}
PLAN_DEFINITIONS = {
    "free": {
        "code": "free",
        "label": "Starter",
        "price_monthly": 0,
        "local_price_kes": 0,
        "summary": "For students, freelancers, and first-time founders building habits before they need advanced finance operations.",
        "max_companies": 1,
        "ai_enabled": False,
        "features": ["1 company", "core accounting", "manual imports"],
    },
    "pro": {
        "code": "pro",
        "label": "Pro",
        "price_monthly": 20,
        "local_price_kes": 900,
        "summary": "For serious small businesses that need workflow automation, exports, bank feeds, and deeper operating control.",
        "max_companies": 5,
        "ai_enabled": False,
        "features": ["5 companies", "team workflows", "automations", "async jobs"],
    },
    "ai": {
        "code": "ai",
        "label": "AI CFO",
        "price_monthly": 50,
        "local_price_kes": 1500,
        "summary": "For owners who want proactive alerts, cash forecasting, and an always-available finance copilot.",
        "max_companies": 25,
        "ai_enabled": True,
        "features": ["25 companies", "AI CFO chat", "proactive alerts", "scenario forecasting"],
    },
}
INTEGRATION_CATALOG = [
    {"provider": "plaid", "category": "banking", "description": "Direct bank feeds and transaction sync."},
    {"provider": "stripe", "category": "payments", "description": "Card payments and customer collections."},
    {"provider": "google_drive", "category": "documents", "description": "Shared workpapers, invoices, and filing packs."},
    {"provider": "slack", "category": "collaboration", "description": "Approval alerts, daily cash, and close notifications."},
    {"provider": "power_bi", "category": "analytics", "description": "Publish accountant and board reporting datasets."},
]
DEFAULT_CHART_OF_ACCOUNTS = [
    {"code": "1000", "name": "Cash", "category": "asset", "subtype": "current", "normal_balance": "debit"},
    {"code": "1100", "name": "Accounts Receivable", "category": "asset", "subtype": "current", "normal_balance": "debit"},
    {"code": "1200", "name": "Inventory Asset", "category": "asset", "subtype": "current", "normal_balance": "debit"},
    {"code": "1250", "name": "Input Tax Receivable", "category": "asset", "subtype": "current", "normal_balance": "debit"},
    {"code": "1300", "name": "Prepaid Expenses", "category": "asset", "subtype": "current", "normal_balance": "debit"},
    {"code": "1500", "name": "Property and Equipment", "category": "asset", "subtype": "non-current", "normal_balance": "debit"},
    {"code": "2000", "name": "Accounts Payable", "category": "liability", "subtype": "current", "normal_balance": "credit"},
    {"code": "2100", "name": "Sales Tax Payable", "category": "liability", "subtype": "current", "normal_balance": "credit"},
    {"code": "2200", "name": "Payroll Withholding Payable", "category": "liability", "subtype": "current", "normal_balance": "credit"},
    {"code": "2300", "name": "Accrued Expenses", "category": "liability", "subtype": "current", "normal_balance": "credit"},
    {"code": "2400", "name": "Contractor Payable", "category": "liability", "subtype": "current", "normal_balance": "credit"},
    {"code": "3000", "name": "Owner Equity", "category": "equity", "subtype": "equity", "normal_balance": "credit"},
    {"code": "3100", "name": "Retained Earnings", "category": "equity", "subtype": "equity", "normal_balance": "credit"},
    {"code": "4000", "name": "Sales Revenue", "category": "revenue", "subtype": "operating", "normal_balance": "credit"},
    {"code": "4100", "name": "Service Revenue", "category": "revenue", "subtype": "operating", "normal_balance": "credit"},
    {"code": "5000", "name": "Cost of Goods Sold", "category": "expense", "subtype": "operating", "normal_balance": "debit"},
    {"code": "5100", "name": "Payroll Expense", "category": "expense", "subtype": "operating", "normal_balance": "debit"},
    {"code": "5200", "name": "Operating Expense", "category": "expense", "subtype": "operating", "normal_balance": "debit"},
    {"code": "5300", "name": "Tax Expense", "category": "expense", "subtype": "other", "normal_balance": "debit"},
    {"code": "5400", "name": "Mileage Reimbursement Expense", "category": "expense", "subtype": "operating", "normal_balance": "debit"},
    {"code": "5500", "name": "Software and SaaS Expense", "category": "expense", "subtype": "operating", "normal_balance": "debit"},
    {"code": "5600", "name": "Contractor Expense", "category": "expense", "subtype": "operating", "normal_balance": "debit"},
]
TAX_JURISDICTION_LIBRARY = {
    "generic": {
        "name": "Generic Indirect Tax",
        "filing_type": "indirect_tax",
        "return_labels": ["taxable_sales", "tax_collected", "tax_credit", "net_tax_due"],
    },
    "ke-vat": {
        "name": "Kenya VAT",
        "filing_type": "indirect_tax",
        "return_labels": ["vatable_sales", "output_vat", "input_vat", "vat_payable"],
    },
    "us-sales-tax": {
        "name": "United States Sales Tax",
        "filing_type": "indirect_tax",
        "return_labels": ["taxable_sales", "sales_tax_collected", "exempt_sales", "net_sales_tax_due"],
    },
    "uk-vat": {
        "name": "United Kingdom VAT",
        "filing_type": "indirect_tax",
        "return_labels": ["vat_due_sales", "vat_reclaimed", "net_vat_due", "total_sales_ex_vat"],
    },
    "ca-gst": {
        "name": "Canada GST/HST",
        "filing_type": "indirect_tax",
        "return_labels": ["gst_hst_collected", "input_tax_credits", "net_gst_hst_due", "taxable_supplies"],
    },
    "au-gst": {
        "name": "Australia GST",
        "filing_type": "indirect_tax",
        "return_labels": ["gst_collected", "gst_credits", "net_gst_due", "taxable_sales"],
    },
}
PLAID_ENVIRONMENTS = {
    "sandbox": "https://sandbox.plaid.com",
    "development": "https://development.plaid.com",
    "production": "https://production.plaid.com",
}

# ---------------- HELPERS ----------------

def error_response(message, status=400):
    return {"error": message}, status


@app.after_request
def add_security_headers(response):
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["X-Frame-Options"] = "DENY"
    response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
    response.headers["Permissions-Policy"] = "camera=(), microphone=(), geolocation=()"
    response.headers["Cache-Control"] = "no-store"
    if request.is_secure or request.headers.get("X-Forwarded-Proto", "").lower() == "https":
        response.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"
    return response


@app.errorhandler(429)
def ratelimit_handler(_):
    return error_response("too many requests, try again later", 429)


def safe_commit():
    try:
        db.session.commit()
        return True
    except SQLAlchemyError:
        db.session.rollback()
        return False


def get_plan_definition(plan_code=None):
    normalized_code = str(plan_code or "free").strip().lower()
    return PLAN_DEFINITIONS.get(normalized_code, PLAN_DEFINITIONS["free"])


def apply_subscription_plan(org, plan_code, status=None, stripe_customer_id=None, stripe_subscription_id=None):
    plan = get_plan_definition(plan_code)
    org.plan_code = plan["code"]
    org.max_companies = int(plan["max_companies"])
    org.ai_assistant_enabled = bool(plan["ai_enabled"])
    org.subscription_status = (status or ("free" if plan["code"] == "free" else "active")).strip().lower()
    if stripe_customer_id:
        org.stripe_customer_id = stripe_customer_id
    if stripe_subscription_id is not None:
        org.stripe_subscription_id = stripe_subscription_id or None
    org.subscription_updated_at = datetime.datetime.now(datetime.UTC)
    return plan


def serialize_subscription(org):
    plan = get_plan_definition(org.plan_code)
    return {
        "plan_code": plan["code"],
        "plan_label": plan["label"],
        "price_monthly": plan["price_monthly"],
        "max_companies": int(org.max_companies or plan["max_companies"]),
        "subscription_status": org.subscription_status or ("free" if plan["code"] == "free" else "active"),
        "ai_enabled": bool(org.ai_assistant_enabled),
        "billing_email": org.billing_email or "",
        "stripe_customer_id": org.stripe_customer_id or "",
        "stripe_configured": bool(stripe and STRIPE_SECRET_KEY),
        "features": plan["features"],
    }


def get_user_from_token():
    try:
        user_id = int(get_jwt_identity())
    except (TypeError, ValueError):
        return None
    return db.session.get(User, user_id)


def build_access_token(user):
    org = db.session.get(Organization, user.org_id) if user.org_id else None
    subscription = serialize_subscription(org) if org else serialize_subscription(Organization(name="", usage=0))
    return create_access_token(
        identity=str(user.id),
        additional_claims={
            "email": user.email,
            "org_id": user.org_id,
            "role": user.role,
            "default_company_id": user.default_company_id,
            "plan_code": subscription["plan_code"],
        },
    )


def roles_required(*allowed_roles):
    allowed = set(allowed_roles)

    def decorator(fn):
        @wraps(fn)
        def wrapper(*args, **kwargs):
            user = get_user_from_token()
            if not user:
                return error_response("invalid token", 401)
            if user.role not in allowed:
                return error_response("not allowed", 403)
            return fn(*args, **kwargs)

        return wrapper

    return decorator


def plan_required(minimum_plan):
    def decorator(fn):
        @wraps(fn)
        def wrapper(*args, **kwargs):
            user = get_user_from_token()
            if not user:
                return error_response("invalid token", 401)
            denial = ensure_plan_access(user, minimum_plan)
            if denial:
                return denial
            return fn(*args, **kwargs)

        return wrapper

    return decorator


def log(user_id, action, company_id=None):
    db.session.add(AuditLog(user_id=user_id, company_id=company_id, action=action))
    # Avoid crashing request flow because of non-critical audit logging issues.
    safe_commit()


def normalize_company_ids(raw_value):
    if raw_value is None:
        return []
    if isinstance(raw_value, str) and raw_value.strip() == "":
        return []

    raw_items = raw_value if isinstance(raw_value, list) else [raw_value]
    company_ids = []
    seen = set()
    for entry in raw_items:
        company_id = parse_company_id(entry)
        if company_id is None or company_id in seen:
            continue
        seen.add(company_id)
        company_ids.append(company_id)
    return company_ids


def membership_company_map(company_ids):
    if not company_ids:
        return {}
    companies = Company.query.filter(Company.id.in_(company_ids)).all()
    return {company.id: company for company in companies}


def membership_sort_key(membership):
    return (0 if membership.is_default else 1, membership.company_id, membership.id)


def assign_company_membership(user, company, role=None, is_default=False):
    membership_role = (role or user.role or "member").strip().lower()
    if membership_role not in VALID_MEMBERSHIP_ROLES:
        membership_role = "member"

    membership = UserCompanyMembership.query.filter_by(user_id=user.id, company_id=company.id).first()
    if not membership:
        membership = UserCompanyMembership(user_id=user.id, company_id=company.id, role=membership_role)
        db.session.add(membership)
        db.session.flush()
    else:
        membership.role = membership_role

    if is_default:
        existing_defaults = UserCompanyMembership.query.filter_by(user_id=user.id, is_default=True).all()
        for existing in existing_defaults:
            if existing.id != membership.id:
                existing.is_default = False
        membership.is_default = True
        user.default_company_id = company.id
    elif user.default_company_id is None:
        membership.is_default = True
        user.default_company_id = company.id

    return membership


def ensure_user_company_memberships(user):
    memberships = UserCompanyMembership.query.filter_by(user_id=user.id).all()
    companies = Company.query.filter_by(org_id=user.org_id).order_by(Company.id.asc()).all()
    if not companies:
        companies = [get_or_create_default_company(user.org_id)]

    changed = False
    if not memberships:
        if user.role in {"owner", "admin"}:
            default_company_id = user.default_company_id or companies[0].id
            for company in companies:
                assign_company_membership(user, company, role=user.role, is_default=company.id == default_company_id)
            changed = True
        else:
            default_company = next((company for company in companies if company.id == user.default_company_id), companies[0])
            assign_company_membership(user, default_company, role=user.role, is_default=True)
            changed = True
        memberships = UserCompanyMembership.query.filter_by(user_id=user.id).all()

    membership_map = {membership.company_id: membership for membership in memberships}
    if user.role in {"owner", "admin"}:
        for company in companies:
            if company.id not in membership_map:
                assign_company_membership(user, company, role=user.role, is_default=False)
                changed = True
        memberships = UserCompanyMembership.query.filter_by(user_id=user.id).all()

    memberships = sorted(memberships, key=membership_sort_key)
    if memberships and not any(membership.is_default for membership in memberships):
        memberships[0].is_default = True
        changed = True
    if memberships and user.default_company_id not in {membership.company_id for membership in memberships}:
        default_membership = next((membership for membership in memberships if membership.is_default), memberships[0])
        user.default_company_id = default_membership.company_id
        changed = True

    if changed:
        safe_commit()
        memberships = sorted(UserCompanyMembership.query.filter_by(user_id=user.id).all(), key=membership_sort_key)
    return memberships


def company_membership_for_user(user, company_id):
    memberships = ensure_user_company_memberships(user)
    for membership in memberships:
        if membership.company_id == company_id:
            return membership
    return None


def accessible_companies_for_user(user):
    memberships = ensure_user_company_memberships(user)
    company_ids = [membership.company_id for membership in memberships]
    companies = membership_company_map(company_ids)
    items = []
    for membership in sorted(memberships, key=membership_sort_key):
        company = companies.get(membership.company_id)
        if company:
            items.append((company, membership))
    return items


def has_company_permission(user, company, action):
    membership = company_membership_for_user(user, company.id)
    if not membership:
        return False
    role = membership.role or user.role or "member"
    return action in ROLE_PERMISSION_MAP.get(role, set())


def normalize_membership_specs(raw_memberships, raw_company_ids, fallback_company_id, fallback_role):
    specs = []
    if isinstance(raw_memberships, list) and raw_memberships:
        for raw_item in raw_memberships:
            item = raw_item or {}
            company_id = parse_company_id(item.get("company_id"))
            if company_id is None:
                continue
            role = (item.get("role") or fallback_role or "member").strip().lower()
            if role not in VALID_MEMBERSHIP_ROLES:
                raise ValueError("invalid membership role")
            specs.append(
                {
                    "company_id": company_id,
                    "role": role,
                    "is_default": parse_bool(item.get("is_default"), False),
                }
            )
    else:
        company_ids = normalize_company_ids(raw_company_ids)
        if not company_ids and fallback_company_id is not None:
            company_ids = [fallback_company_id]
        for index, company_id in enumerate(company_ids):
            specs.append(
                {
                    "company_id": company_id,
                    "role": fallback_role,
                    "is_default": index == 0,
                }
            )

    unique_specs = []
    seen = set()
    for spec in specs:
        if spec["company_id"] in seen:
            continue
        seen.add(spec["company_id"])
        unique_specs.append(spec)

    if not unique_specs:
        raise ValueError("at least one company membership is required")
    if not any(spec["is_default"] for spec in unique_specs):
        unique_specs[0]["is_default"] = True
    return unique_specs


def apply_user_company_memberships(user, membership_specs):
    specs = membership_specs or []
    if not specs:
        raise ValueError("at least one company membership is required")

    company_ids = [spec["company_id"] for spec in specs]
    company_map = membership_company_map(company_ids)
    if len(company_map) != len(company_ids):
        raise ValueError("one or more companies were not found")
    invalid_company = next((company for company in company_map.values() if company.org_id != user.org_id), None)
    if invalid_company:
        raise ValueError("company does not belong to this organization")

    existing = {membership.company_id: membership for membership in UserCompanyMembership.query.filter_by(user_id=user.id).all()}
    desired_ids = set(company_ids)
    for company_id, membership in existing.items():
        if company_id not in desired_ids:
            db.session.delete(membership)

    default_company_id = next((spec["company_id"] for spec in specs if spec["is_default"]), specs[0]["company_id"])
    for spec in specs:
        membership = existing.get(spec["company_id"])
        if not membership:
            membership = UserCompanyMembership(user_id=user.id, company_id=spec["company_id"])
            db.session.add(membership)
        membership.role = spec["role"]
        membership.is_default = spec["company_id"] == default_company_id

    user.default_company_id = default_company_id
    return company_map.get(default_company_id)


def serialize_membership(membership, company=None):
    company = company or db.session.get(Company, membership.company_id)
    return {
        "company_id": membership.company_id,
        "company_name": company.name if company else "",
        "role": membership.role,
        "is_default": bool(membership.is_default),
    }


def serialize_user(user):
    memberships = ensure_user_company_memberships(user)
    company_map = membership_company_map([membership.company_id for membership in memberships])
    return {
        "id": user.id,
        "email": user.email,
        "role": user.role,
        "org_id": user.org_id,
        "default_company_id": user.default_company_id,
        "memberships": [serialize_membership(membership, company_map.get(membership.company_id)) for membership in memberships],
    }


def hash_key(raw_key):
    return hashlib.sha256(raw_key.encode()).hexdigest()


def frontend_base_url():
    explicit = os.getenv("FRONTEND_URL", "").strip()
    if explicit:
        return explicit.rstrip("/")
    vercel_url = os.getenv("VERCEL_URL", "").strip()
    if vercel_url:
        return f"https://{vercel_url}".rstrip("/")
    return "http://localhost:5173"


def backend_base_url():
    explicit = os.getenv("PUBLIC_API_URL", "").strip()
    if explicit:
        return explicit.rstrip("/")
    return "http://localhost:5000"


def should_return_password_reset_preview():
    return ENV != "production"


def plan_rank(plan_code):
    ordering = {"free": 0, "pro": 1, "ai": 2}
    return ordering.get(str(plan_code or "free").strip().lower(), 0)


def org_has_plan(org, minimum_plan):
    return plan_rank(org.plan_code if org else "free") >= plan_rank(minimum_plan)


def ensure_plan_access(user, minimum_plan):
    org = db.session.get(Organization, user.org_id)
    if not org_has_plan(org, minimum_plan):
        label = get_plan_definition(minimum_plan)["label"]
        return error_response(f"{label} plan required", 403)
    return None


def normalize_phone_number(raw_value):
    digits = re.sub(r"\D", "", str(raw_value or ""))
    if digits.startswith("0") and len(digits) == 10:
        digits = f"254{digits[1:]}"
    if digits.startswith("254") and len(digits) == 12:
        return digits
    if digits.startswith("7") and len(digits) == 9:
        return f"254{digits}"
    raise ValueError("phone_number must be a valid Kenya mobile number")


def mpesa_is_configured():
    return all([MPESA_CONSUMER_KEY, MPESA_CONSUMER_SECRET, MPESA_SHORTCODE, MPESA_PASSKEY])


def call_json_api(url, payload=None, headers=None, method=None):
    request_headers = {"Content-Type": "application/json", **(headers or {})}
    body = None if payload is None else json.dumps(payload).encode("utf-8")
    request_method = method or ("POST" if body is not None else "GET")
    req = urllib.request.Request(url, data=body, headers=request_headers, method=request_method)
    with urllib.request.urlopen(req, timeout=25) as response:
        charset = response.headers.get_content_charset() or "utf-8"
        raw = response.read().decode(charset)
    return json.loads(raw or "{}")


def fetch_mpesa_access_token():
    credentials = f"{MPESA_CONSUMER_KEY}:{MPESA_CONSUMER_SECRET}".encode("utf-8")
    encoded_credentials = base64.b64encode(credentials).decode("utf-8")
    req = urllib.request.Request(
        f"{MPESA_BASE_URL}/oauth/v1/generate?grant_type=client_credentials",
        headers={"Authorization": f"Basic {encoded_credentials}"},
        method="GET",
    )
    with urllib.request.urlopen(req, timeout=25) as response:
        charset = response.headers.get_content_charset() or "utf-8"
        raw = response.read().decode(charset)
    payload = json.loads(raw or "{}")
    return payload.get("access_token")


def current_mpesa_timestamp():
    return datetime.datetime.now(datetime.UTC).strftime("%Y%m%d%H%M%S")


def build_mpesa_password(timestamp):
    raw_password = f"{MPESA_SHORTCODE}{MPESA_PASSKEY}{timestamp}".encode("utf-8")
    return base64.b64encode(raw_password).decode("utf-8")


def initiate_mpesa_stk_push(plan_code, phone_number):
    plan = get_plan_definition(plan_code)
    amount = int(plan["local_price_kes"])
    reference = f"FIN-{plan_code.upper()}-{secrets.token_hex(4).upper()}"

    if not mpesa_is_configured():
        return {
            "provider": "mpesa",
            "status": "preview",
            "amount": amount,
            "currency_code": "KES",
            "plan_code": plan_code,
            "reference": reference,
            "merchant_request_id": f"preview-merchant-{secrets.token_hex(3)}",
            "checkout_request_id": f"preview-checkout-{secrets.token_hex(3)}",
            "customer_message": f"Preview mode: ask {phone_number} to approve a KES {amount} M-Pesa subscription payment.",
            "raw": {"preview": True},
        }

    timestamp = current_mpesa_timestamp()
    callback_url = MPESA_CALLBACK_URL or f"{backend_base_url()}/billing/mpesa/callback"
    access_token = fetch_mpesa_access_token()
    if not access_token:
        raise ValueError("unable to obtain M-Pesa access token")

    payload = {
        "BusinessShortCode": MPESA_SHORTCODE,
        "Password": build_mpesa_password(timestamp),
        "Timestamp": timestamp,
        "TransactionType": "CustomerPayBillOnline",
        "Amount": amount,
        "PartyA": phone_number,
        "PartyB": MPESA_SHORTCODE,
        "PhoneNumber": phone_number,
        "CallBackURL": callback_url,
        "AccountReference": reference,
        "TransactionDesc": f"{plan['label']} monthly subscription",
    }
    response = call_json_api(
        f"{MPESA_BASE_URL}/mpesa/stkpush/v1/processrequest",
        payload=payload,
        headers={"Authorization": f"Bearer {access_token}"},
    )
    return {
        "provider": "mpesa",
        "status": "pending",
        "amount": amount,
        "currency_code": "KES",
        "plan_code": plan_code,
        "reference": reference,
        "merchant_request_id": response.get("MerchantRequestID"),
        "checkout_request_id": response.get("CheckoutRequestID"),
        "customer_message": response.get("CustomerMessage") or response.get("ResponseDescription") or "STK push initiated.",
        "raw": response,
    }


def serialize_billing_payment_request(payment_request):
    provider_response = {}
    callback_payload = {}
    try:
        provider_response = json.loads(payment_request.provider_response_json or "{}")
    except (TypeError, ValueError):
        provider_response = {}
    try:
        callback_payload = json.loads(payment_request.callback_payload_json or "{}")
    except (TypeError, ValueError):
        callback_payload = {}
    return {
        "id": payment_request.id,
        "provider": payment_request.provider,
        "plan_code": payment_request.plan_code,
        "currency_code": payment_request.currency_code,
        "amount": round(float(payment_request.amount or 0), 2),
        "phone_number": payment_request.phone_number or "",
        "external_reference": payment_request.external_reference or "",
        "merchant_request_id": payment_request.merchant_request_id or "",
        "checkout_request_id": payment_request.checkout_request_id or "",
        "status": payment_request.status,
        "provider_response": provider_response,
        "callback_payload": callback_payload,
        "created_at": payment_request.created_at.isoformat() if payment_request.created_at else None,
        "updated_at": payment_request.updated_at.isoformat() if payment_request.updated_at else None,
    }


def build_password_reset_link(raw_token):
    return f"{frontend_base_url()}/?auth=reset&token={raw_token}"


def expire_existing_password_reset_tokens(user_id):
    now = datetime.datetime.now(datetime.UTC)
    tokens = PasswordResetToken.query.filter_by(user_id=user_id, used_at=None).all()
    for token in tokens:
        token.used_at = now


def issue_password_reset_token(user):
    raw_token = secrets.token_urlsafe(32)
    expire_existing_password_reset_tokens(user.id)
    record = PasswordResetToken(
        user_id=user.id,
        token_hash=hash_key(raw_token),
        expires_at=datetime.datetime.now(datetime.UTC) + datetime.timedelta(minutes=PASSWORD_RESET_EXPIRY_MINUTES),
    )
    db.session.add(record)
    return raw_token


def get_valid_password_reset_record(raw_token):
    token_hash = hash_key(raw_token)
    record = PasswordResetToken.query.filter_by(token_hash=token_hash).first()
    if not record:
        return None
    if record.used_at is not None:
        return None
    expires_at = record.expires_at
    if expires_at.tzinfo is None:
        expires_at = expires_at.replace(tzinfo=datetime.UTC)
    if expires_at <= datetime.datetime.now(datetime.UTC):
        record.used_at = datetime.datetime.now(datetime.UTC)
        safe_commit()
        return None
    return record


def send_password_reset_email(user, reset_link):
    smtp_host = os.getenv("SMTP_HOST", "").strip()
    smtp_port = int(os.getenv("SMTP_PORT", "587"))
    smtp_username = os.getenv("SMTP_USERNAME", "").strip()
    smtp_password = os.getenv("SMTP_PASSWORD", "")
    smtp_from = (os.getenv("SMTP_FROM_EMAIL", "").strip() or smtp_username).strip()
    if not smtp_host or not smtp_from:
        return False

    message = EmailMessage()
    message["Subject"] = "Reset your FinancePro password"
    message["From"] = smtp_from
    message["To"] = user.email
    message.set_content(
        "\n".join(
            [
                "A password reset was requested for your FinancePro account.",
                "",
                f"Reset link: {reset_link}",
                "",
                f"This link expires in {PASSWORD_RESET_EXPIRY_MINUTES} minutes.",
                "If you did not request this change, you can ignore this email.",
            ]
        )
    )

    try:
        with smtplib.SMTP(smtp_host, smtp_port, timeout=15) as smtp:
            smtp.ehlo()
            if os.getenv("SMTP_USE_TLS", "true").strip().lower() not in {"false", "0", "no"}:
                smtp.starttls()
                smtp.ehlo()
            if smtp_username:
                smtp.login(smtp_username, smtp_password)
            smtp.send_message(message)
        return True
    except Exception:
        return False


def touch_session(user_id):
    now = datetime.datetime.now(datetime.UTC)
    session = ActiveSession.query.filter_by(user_id=user_id).first()
    if session:
        session.last_seen = now
    else:
        db.session.add(ActiveSession(user_id=user_id, last_seen=now))
    safe_commit()


def clear_session(user_id):
    session = ActiveSession.query.filter_by(user_id=user_id).first()
    if session:
        db.session.delete(session)
    safe_commit()


def clear_all_sessions(user_id):
    ActiveSession.query.filter_by(user_id=user_id).delete()
    safe_commit()


def active_user_count_for_org(org_id, online_window_minutes=5):
    cutoff = datetime.datetime.now(datetime.UTC) - datetime.timedelta(minutes=online_window_minutes)
    return (
        db.session.query(ActiveSession.id)
        .join(User, ActiveSession.user_id == User.id)
        .filter(User.org_id == org_id, ActiveSession.last_seen >= cutoff)
        .count()
    )


def aggregate_org_reports(org_id, company_id=None):
    query = Report.query.filter_by(org_id=org_id)
    if company_id is not None:
        query = query.filter_by(company_id=company_id)
    reports = query.all()
    revenue_total = 0.0
    expense_total = 0.0
    assets_total = 0.0

    for report in reports:
        try:
            payload = json.loads(report.data or "{}")
        except (TypeError, ValueError):
            payload = {}

        revenue_total += float(payload.get("revenue", 0) or 0)
        expense_total += float(payload.get("expenses", payload.get("expense", 0)) or 0)
        assets_total += float(payload.get("total_assets", payload.get("totalAssets", 0)) or 0)

    return {
        "revenue": round(revenue_total, 2),
        "expenses": round(expense_total, 2),
        "profit": round(revenue_total - expense_total, 2),
        "total_assets": round(assets_total, 2),
    }


def get_or_create_default_company(org_id, name="Main Company"):
    company = Company.query.filter_by(org_id=org_id).order_by(Company.id.asc()).first()
    if company:
        seed_chart_of_accounts(company)
        seed_integration_connections(company)
        return company

    company = Company(org_id=org_id, name=name, business_type="sole_proprietor")
    db.session.add(company)
    db.session.flush()
    seed_chart_of_accounts(company)
    seed_integration_connections(company)
    safe_commit()
    return company


def normalize_business_type(raw_value):
    business_type = str(raw_value or "sole_proprietor").strip().lower()
    if business_type not in VALID_BUSINESS_TYPES:
        return None
    return business_type


def normalize_partner_names(raw_value):
    if not isinstance(raw_value, list):
        return []

    names = []
    seen = set()
    for entry in raw_value:
        name = " ".join(str(entry or "").split())
        if not name:
            continue
        lower_name = name.lower()
        if lower_name in seen:
            raise ValueError("partner names must be unique")
        seen.add(lower_name)
        names.append(name)
    return names


def validate_company_setup(business_type, partner_names):
    if business_type not in VALID_BUSINESS_TYPES:
        return "invalid business type"
    if business_type == "partnership" and len(partner_names) < 2:
        return "partnerships require at least two partner names"
    return None


def replace_company_partners(company, partner_names):
    existing = CompanyPartner.query.filter_by(company_id=company.id).all()
    for partner in existing:
        db.session.delete(partner)

    for index, partner_name in enumerate(partner_names, start=1):
        db.session.add(
            CompanyPartner(
                company_id=company.id,
                name=partner_name,
                display_order=index,
            )
        )


def set_company_onboarding_state(company_id, is_configured):
    state = CompanyOnboardingState.query.filter_by(company_id=company_id).first()
    if not state:
        state = CompanyOnboardingState(company_id=company_id)
        db.session.add(state)
    state.is_configured = bool(is_configured)
    state.configured_at = datetime.datetime.now(datetime.UTC) if is_configured else None
    return state


def is_company_onboarding_complete(company_id):
    state = CompanyOnboardingState.query.filter_by(company_id=company_id).first()
    return bool(state and state.is_configured)


def get_company_partner_names(company_id):
    partners = (
        CompanyPartner.query.filter_by(company_id=company_id)
        .order_by(CompanyPartner.display_order.asc(), CompanyPartner.id.asc())
        .all()
    )
    return [partner.name for partner in partners]


def serialize_company(company, membership=None):
    partner_names = get_company_partner_names(company.id)
    return {
        "id": company.id,
        "name": company.name,
        "business_type": company.business_type,
        "partner_names": partner_names,
        "partner_count": len(partner_names),
        "onboarding_complete": is_company_onboarding_complete(company.id),
        "membership_role": membership.role if membership else None,
        "is_default": bool(membership.is_default) if membership else False,
    }


def parse_company_id(raw_value):
    if raw_value in {None, ""}:
        return None
    try:
        return int(raw_value)
    except (TypeError, ValueError):
        return None


def resolve_company_for_user(user, raw_company_id=None):
    memberships = ensure_user_company_memberships(user)
    if not memberships:
        return None

    company_id = parse_company_id(raw_company_id)
    if company_id is None:
        default_membership = next((membership for membership in memberships if membership.is_default), memberships[0])
        company_id = default_membership.company_id

    membership = next((row for row in memberships if row.company_id == company_id), None)
    if not membership:
        return None
    return Company.query.filter_by(id=company_id, org_id=user.org_id).first()


def today_utc_date():
    return datetime.datetime.now(datetime.UTC).date()


def iso_date(value):
    return value.isoformat() if value else None


def parse_iso_date(raw_value, field_name, default=None):
    if raw_value in {None, ""}:
        return default
    if isinstance(raw_value, datetime.datetime):
        return raw_value.date()
    if isinstance(raw_value, datetime.date):
        return raw_value
    try:
        return datetime.date.fromisoformat(str(raw_value))
    except ValueError as exc:
        raise ValueError(f"{field_name} must be in YYYY-MM-DD format") from exc


def parse_money(value, field_name):
    try:
        amount = float(value)
    except (TypeError, ValueError) as exc:
        raise ValueError(f"{field_name} must be numeric") from exc
    return round(amount, 2)


def parse_bool(value, default=False):
    if value in {None, ""}:
        return default
    if isinstance(value, bool):
        return value
    return str(value).strip().lower() in {"1", "true", "yes", "on"}


def normalize_document_items(items, document_name):
    if not isinstance(items, list) or not items:
        raise ValueError(f"{document_name} requires at least one line item")

    normalized_items = []
    subtotal = 0.0
    for index, item in enumerate(items, start=1):
        payload = item or {}
        description = (payload.get("description") or f"Line {index}").strip()
        quantity = parse_money(payload.get("quantity", 0), f"line {index} quantity")
        unit_price = parse_money(payload.get("unit_price", 0), f"line {index} unit_price")
        if quantity <= 0:
            raise ValueError(f"line {index} quantity must be greater than 0")
        if unit_price < 0:
            raise ValueError(f"line {index} unit_price cannot be negative")

        amount = round(quantity * unit_price, 2)
        subtotal += amount
        normalized_items.append(
            {
                "description": description,
                "quantity": quantity,
                "unit_price": unit_price,
                "amount": amount,
            }
        )

    return normalized_items, round(subtotal, 2)


def generate_document_number(model_class, company_id, prefix):
    next_number = model_class.query.filter_by(company_id=company_id).count() + 1
    return f"{prefix}-{int(company_id):03d}-{next_number:05d}"


def invoice_items_for(invoice_id):
    return InvoiceItem.query.filter_by(invoice_id=invoice_id).order_by(InvoiceItem.id.asc()).all()


def bill_items_for(bill_id):
    return VendorBillItem.query.filter_by(bill_id=bill_id).order_by(VendorBillItem.id.asc()).all()


def serialize_line_items(items):
    return [
        {
            "id": item.id,
            "description": item.description,
            "quantity": round(float(item.quantity or 0), 2),
            "unit_price": round(float(item.unit_price or 0), 2),
            "amount": round(float(item.amount or 0), 2),
        }
        for item in items
    ]


def total_customer_payments(invoice_id):
    rows = CustomerPayment.query.filter_by(invoice_id=invoice_id).all()
    return round(sum(float(row.amount or 0) for row in rows), 2)


def total_vendor_payments(bill_id):
    rows = VendorPayment.query.filter_by(bill_id=bill_id).all()
    return round(sum(float(row.amount or 0) for row in rows), 2)


def refresh_invoice_status(invoice):
    previous_status = invoice.status
    previous_balance = round(float(invoice.balance_due or 0), 2)
    payments_total = total_customer_payments(invoice.id)
    invoice.balance_due = round(max(0.0, float(invoice.total_amount or 0) - payments_total), 2)
    today = today_utc_date()

    if previous_status == "cancelled":
        return previous_status != invoice.status or previous_balance != invoice.balance_due

    if invoice.balance_due <= 0.009:
        invoice.balance_due = 0.0
        invoice.status = "paid"
        invoice.paid_at = invoice.paid_at or datetime.datetime.now(datetime.UTC)
    elif payments_total > 0:
        invoice.status = "partial"
    elif invoice.status == "draft":
        invoice.status = "draft"
    elif invoice.due_date and invoice.due_date < today:
        invoice.status = "overdue"
    else:
        invoice.status = "sent"

    return previous_status != invoice.status or previous_balance != invoice.balance_due


def refresh_bill_status(bill):
    previous_status = bill.status
    previous_balance = round(float(bill.balance_due or 0), 2)
    payments_total = total_vendor_payments(bill.id)
    bill.balance_due = round(max(0.0, float(bill.total_amount or 0) - payments_total), 2)
    today = today_utc_date()

    if previous_status == "cancelled":
        return previous_status != bill.status or previous_balance != bill.balance_due

    if bill.balance_due <= 0.009:
        bill.balance_due = 0.0
        bill.status = "paid"
        bill.paid_at = bill.paid_at or datetime.datetime.now(datetime.UTC)
    elif payments_total > 0:
        bill.status = "partial"
    elif bill.status == "draft":
        bill.status = "draft"
    elif bill.due_date and bill.due_date < today:
        bill.status = "overdue"
    else:
        bill.status = "approved"

    return previous_status != bill.status or previous_balance != bill.balance_due


def refresh_finance_documents(company_id):
    dirty = False
    invoices = Invoice.query.filter_by(company_id=company_id).all()
    bills = VendorBill.query.filter_by(company_id=company_id).all()

    for invoice in invoices:
        dirty = refresh_invoice_status(invoice) or dirty
    for bill in bills:
        dirty = refresh_bill_status(bill) or dirty

    if dirty:
        safe_commit()

    return invoices, bills


def serialize_invoice(invoice):
    paid_amount = round(float(invoice.total_amount or 0) - float(invoice.balance_due or 0), 2)
    return {
        "id": invoice.id,
        "invoice_number": invoice.invoice_number,
        "customer_name": invoice.customer_name,
        "customer_email": invoice.customer_email,
        "status": invoice.status,
        "issue_date": iso_date(invoice.issue_date),
        "due_date": iso_date(invoice.due_date),
        "subtotal": round(float(invoice.subtotal or 0), 2),
        "tax_rate": round(float(invoice.tax_rate or 0), 2),
        "tax_amount": round(float(invoice.tax_amount or 0), 2),
        "total_amount": round(float(invoice.total_amount or 0), 2),
        "balance_due": round(float(invoice.balance_due or 0), 2),
        "paid_amount": paid_amount,
        "notes": invoice.notes or "",
        "overdue": invoice.status == "overdue",
        "items": serialize_line_items(invoice_items_for(invoice.id)),
    }


def serialize_bill(bill):
    paid_amount = round(float(bill.total_amount or 0) - float(bill.balance_due or 0), 2)
    return {
        "id": bill.id,
        "bill_number": bill.bill_number,
        "vendor_name": bill.vendor_name,
        "status": bill.status,
        "issue_date": iso_date(bill.issue_date),
        "due_date": iso_date(bill.due_date),
        "subtotal": round(float(bill.subtotal or 0), 2),
        "tax_rate": round(float(bill.tax_rate or 0), 2),
        "tax_amount": round(float(bill.tax_amount or 0), 2),
        "total_amount": round(float(bill.total_amount or 0), 2),
        "balance_due": round(float(bill.balance_due or 0), 2),
        "paid_amount": paid_amount,
        "notes": bill.notes or "",
        "overdue": bill.status == "overdue",
        "items": serialize_line_items(bill_items_for(bill.id)),
    }


def serialize_bank_transaction(transaction):
    amount = round(float(transaction.amount or 0), 2)
    return {
        "id": transaction.id,
        "posted_at": iso_date(transaction.posted_at),
        "description": transaction.description,
        "amount": amount,
        "absolute_amount": round(abs(amount), 2),
        "direction": "inflow" if amount >= 0 else "outflow",
        "reference": transaction.reference or "",
        "status": transaction.status,
        "matched_invoice_id": transaction.matched_invoice_id,
        "matched_bill_id": transaction.matched_bill_id,
    }


def normalize_bank_feed_dataframe(df):
    rename_map = {}
    for column in df.columns:
        key = str(column).strip().lower()
        if key in {"date", "posted_at", "posted date", "transaction_date", "value date"}:
            rename_map[column] = "posted_at"
        elif key in {"description", "details", "memo", "narration"}:
            rename_map[column] = "description"
        elif key in {"amount", "value"}:
            rename_map[column] = "amount"
        elif key in {"debit", "withdrawal", "money_out"}:
            rename_map[column] = "debit"
        elif key in {"credit", "deposit", "money_in"}:
            rename_map[column] = "credit"
        elif key in {"reference", "ref", "transaction_id"}:
            rename_map[column] = "reference"

    normalized = df.rename(columns=rename_map).copy()
    missing = {"posted_at", "description"}.difference(normalized.columns)
    if missing:
        raise ValueError(f"bank feed missing required columns: {', '.join(sorted(missing))}")
    if "amount" not in normalized.columns and not {"debit", "credit"}.intersection(normalized.columns):
        raise ValueError("bank feed must include amount or debit/credit columns")

    normalized["posted_at"] = pd.to_datetime(normalized["posted_at"], errors="coerce").dt.date
    if normalized["posted_at"].isna().any():
        raise ValueError("bank feed date column must contain valid dates")

    def to_numeric_series(series):
        return pd.to_numeric(series.astype(str).str.replace(",", "", regex=False).str.strip(), errors="coerce")

    if "amount" in normalized.columns:
        normalized["amount"] = to_numeric_series(normalized["amount"])
    else:
        credit = to_numeric_series(normalized.get("credit", pd.Series([0] * len(normalized))))
        debit = to_numeric_series(normalized.get("debit", pd.Series([0] * len(normalized))))
        normalized["amount"] = credit.fillna(0) - debit.fillna(0)

    if normalized["amount"].isna().any():
        raise ValueError("bank feed amount column must contain numeric values")

    if "reference" not in normalized.columns:
        normalized["reference"] = ""

    normalized["description"] = normalized["description"].astype(str).str.strip()
    return normalized[["posted_at", "description", "amount", "reference"]]


def calculate_aging(documents, key_name):
    buckets = {
        "current": 0.0,
        "days_1_30": 0.0,
        "days_31_60": 0.0,
        "days_61_90": 0.0,
        "days_90_plus": 0.0,
    }
    items = []
    today = today_utc_date()

    for document in documents:
        balance_due = round(float(document.balance_due or 0), 2)
        if balance_due <= 0 or document.status in {"draft", "cancelled", "paid"}:
            continue

        days_past_due = max(0, (today - document.due_date).days) if document.due_date else 0
        if days_past_due == 0:
            buckets["current"] += balance_due
        elif days_past_due <= 30:
            buckets["days_1_30"] += balance_due
        elif days_past_due <= 60:
            buckets["days_31_60"] += balance_due
        elif days_past_due <= 90:
            buckets["days_61_90"] += balance_due
        else:
            buckets["days_90_plus"] += balance_due

        items.append(
            {
                "id": document.id,
                key_name: getattr(document, key_name),
                "counterparty": getattr(document, "customer_name", None) or getattr(document, "vendor_name", ""),
                "due_date": iso_date(document.due_date),
                "status": document.status,
                "balance_due": balance_due,
                "days_past_due": days_past_due,
            }
        )

    return {
        "buckets": {key: round(value, 2) for key, value in buckets.items()},
        "total_open": round(sum(buckets.values()), 2),
        "overdue_count": sum(1 for item in items if item["days_past_due"] > 0),
        "items": items,
    }


def calculate_finance_summary(company):
    invoices, bills = refresh_finance_documents(company.id)
    today = today_utc_date()

    invoice_payments = CustomerPayment.query.filter_by(company_id=company.id).all()
    bill_payments = VendorPayment.query.filter_by(company_id=company.id).all()
    bank_transactions = BankFeedTransaction.query.filter_by(company_id=company.id).all()

    open_receivables = sum(invoice.balance_due for invoice in invoices if invoice.status in {"sent", "partial", "overdue"})
    overdue_receivables = sum(invoice.balance_due for invoice in invoices if invoice.status == "overdue")
    open_payables = sum(bill.balance_due for bill in bills if bill.status in {"approved", "partial", "overdue"})
    overdue_payables = sum(bill.balance_due for bill in bills if bill.status == "overdue")

    collected_this_month = sum(
        payment.amount
        for payment in invoice_payments
        if payment.payment_date and payment.payment_date.year == today.year and payment.payment_date.month == today.month
    )
    paid_this_month = sum(
        payment.amount
        for payment in bill_payments
        if payment.payment_date and payment.payment_date.year == today.year and payment.payment_date.month == today.month
    )

    sales_tax_due = sum(invoice.tax_amount for invoice in invoices if invoice.status not in {"draft", "cancelled"})
    purchase_tax_credit = sum(bill.tax_amount for bill in bills if bill.status not in {"draft", "cancelled"})

    return {
        "open_receivables": round(open_receivables, 2),
        "overdue_receivables": round(overdue_receivables, 2),
        "open_payables": round(open_payables, 2),
        "overdue_payables": round(overdue_payables, 2),
        "collected_this_month": round(collected_this_month, 2),
        "paid_this_month": round(paid_this_month, 2),
        "invoice_count": len(invoices),
        "bill_count": len(bills),
        "overdue_invoice_count": sum(1 for invoice in invoices if invoice.status == "overdue"),
        "overdue_bill_count": sum(1 for bill in bills if bill.status == "overdue"),
        "bank_unmatched_count": sum(1 for transaction in bank_transactions if transaction.status == "unmatched"),
        "bank_net_flow": round(sum(float(transaction.amount or 0) for transaction in bank_transactions), 2),
        "sales_tax_due": round(sales_tax_due, 2),
        "purchase_tax_credit": round(purchase_tax_credit, 2),
        "net_tax_due": round(sales_tax_due - purchase_tax_credit, 2),
    }


def calculate_tax_summary(company, profile=None):
    profile = profile or get_or_create_tax_profile(company)
    invoices, bills = refresh_finance_documents(company.id)
    sales_tax_collected = sum(invoice.tax_amount for invoice in invoices if invoice.status not in {"draft", "cancelled"})
    purchase_tax_credit = sum(bill.tax_amount for bill in bills if bill.status not in {"draft", "cancelled"})
    taxable_profit = sum(invoice.subtotal for invoice in invoices if invoice.status not in {"draft", "cancelled"}) - sum(
        bill.subtotal for bill in bills if bill.status not in {"draft", "cancelled"}
    )
    estimated_income_tax = max(0.0, taxable_profit) * (float(profile.income_tax_rate or 0) / 100)
    return {
        "jurisdiction_code": profile.jurisdiction_code,
        "filing_frequency": profile.filing_frequency,
        "currency_code": profile.currency_code,
        "sales_tax_name": profile.sales_tax_name,
        "purchase_tax_name": profile.purchase_tax_name,
        "income_tax_rate": round(float(profile.income_tax_rate or 0), 2),
        "indirect_tax_rate": round(float(profile.indirect_tax_rate or 0), 2),
        "sales_tax_collected": round(sales_tax_collected, 2),
        "purchase_tax_credit": round(purchase_tax_credit, 2),
        "net_tax_due": round(sales_tax_collected - purchase_tax_credit, 2),
        "taxable_profit": round(taxable_profit, 2),
        "estimated_income_tax": round(estimated_income_tax, 2),
        "effective_tax_rate": (float(profile.income_tax_rate or 0) / 100) if taxable_profit > 0 else 0.0,
    }


def build_reconciliation_suggestions(company):
    invoices, bills = refresh_finance_documents(company.id)
    invoice_candidates = [invoice for invoice in invoices if invoice.status in {"sent", "partial", "overdue"} and invoice.balance_due > 0]
    bill_candidates = [bill for bill in bills if bill.status in {"approved", "partial", "overdue"} and bill.balance_due > 0]
    transactions = (
        BankFeedTransaction.query.filter(
            BankFeedTransaction.company_id == company.id,
            BankFeedTransaction.status.in_(["unmatched", "rule_matched"]),
        )
        .order_by(BankFeedTransaction.posted_at.desc(), BankFeedTransaction.id.desc())
        .all()
    )

    suggestions = []
    for transaction in transactions:
        direction_candidates = invoice_candidates if float(transaction.amount or 0) >= 0 else bill_candidates
        best_match = None
        best_confidence = 0.0
        for document in direction_candidates:
            open_balance = round(float(document.balance_due or 0), 2)
            transaction_amount = round(abs(float(transaction.amount or 0)), 2)
            if transaction_amount <= 0 or open_balance <= 0 or transaction_amount > open_balance + 0.01:
                continue

            counterparty = getattr(document, "customer_name", None) or getattr(document, "vendor_name", "")
            description = (transaction.description or "").lower()
            counterparty_tokens = [token for token in counterparty.lower().split() if len(token) > 2]
            text_match = any(token in description for token in counterparty_tokens)
            amount_gap = abs(open_balance - transaction_amount)
            exact_amount = amount_gap <= 0.01
            date_proximity = abs((transaction.posted_at - document.due_date).days) if document.due_date else 30
            confidence = 0.35
            confidence += 0.35 if exact_amount else 0.15
            confidence += 0.2 if text_match else 0.0
            confidence += 0.1 if date_proximity <= 14 else 0.0

            if confidence > best_confidence:
                best_confidence = confidence
                best_match = document

        if best_match:
            suggestions.append(
                {
                    "transaction": serialize_bank_transaction(transaction),
                    "entity_type": "invoice" if float(transaction.amount or 0) >= 0 else "bill",
                    "entity_id": best_match.id,
                    "document_number": getattr(best_match, "invoice_number", None) or getattr(best_match, "bill_number", ""),
                    "counterparty": getattr(best_match, "customer_name", None) or getattr(best_match, "vendor_name", ""),
                    "open_balance": round(float(best_match.balance_due or 0), 2),
                    "confidence": round(best_confidence, 2),
                }
            )

    return sorted(suggestions, key=lambda item: item["confidence"], reverse=True)[:20]


def get_invoice_for_user(user, invoice_id):
    return Invoice.query.filter_by(id=invoice_id, org_id=user.org_id).first()


def get_bill_for_user(user, bill_id):
    return VendorBill.query.filter_by(id=bill_id, org_id=user.org_id).first()


def get_bank_transaction_for_user(user, transaction_id):
    return BankFeedTransaction.query.filter_by(id=transaction_id, org_id=user.org_id).first()


def apply_customer_payment(invoice, amount, payment_date, reference="", source="manual", notes="", bank_transaction_id=None):
    if amount <= 0:
        raise ValueError("payment amount must be greater than 0")
    refresh_invoice_status(invoice)
    if amount - float(invoice.balance_due or 0) > 0.01:
        raise ValueError("payment exceeds invoice balance")

    db.session.add(
        CustomerPayment(
            org_id=invoice.org_id,
            company_id=invoice.company_id,
            invoice_id=invoice.id,
            amount=round(amount, 2),
            reference=(reference or "").strip() or None,
            source=source,
            notes=notes or None,
            payment_date=payment_date,
            bank_transaction_id=bank_transaction_id,
        )
    )
    db.session.flush()
    refresh_invoice_status(invoice)


def apply_vendor_payment(bill, amount, payment_date, reference="", source="manual", notes="", bank_transaction_id=None):
    if amount <= 0:
        raise ValueError("payment amount must be greater than 0")
    refresh_bill_status(bill)
    if amount - float(bill.balance_due or 0) > 0.01:
        raise ValueError("payment exceeds bill balance")

    db.session.add(
        VendorPayment(
            org_id=bill.org_id,
            company_id=bill.company_id,
            bill_id=bill.id,
            amount=round(amount, 2),
            reference=(reference or "").strip() or None,
            source=source,
            notes=notes or None,
            payment_date=payment_date,
            bank_transaction_id=bank_transaction_id,
        )
    )
    db.session.flush()
    refresh_bill_status(bill)


def plaid_enabled():
    return bool(os.getenv("PLAID_CLIENT_ID") and os.getenv("PLAID_SECRET"))


def plaid_base_url():
    env = (os.getenv("PLAID_ENV", "sandbox") or "sandbox").strip().lower()
    return PLAID_ENVIRONMENTS.get(env, PLAID_ENVIRONMENTS["sandbox"])


def plaid_products():
    raw = (os.getenv("PLAID_PRODUCTS", "transactions") or "transactions").strip()
    return [value.strip() for value in raw.split(",") if value.strip()]


def plaid_country_codes():
    raw = (os.getenv("PLAID_COUNTRY_CODES", "US") or "US").strip()
    return [value.strip().upper() for value in raw.split(",") if value.strip()]


def call_plaid(endpoint, payload):
    if not plaid_enabled():
        raise ValueError("Plaid is not configured. Set PLAID_CLIENT_ID and PLAID_SECRET to enable direct bank connections.")

    body = json.dumps(
        {
            "client_id": os.getenv("PLAID_CLIENT_ID"),
            "secret": os.getenv("PLAID_SECRET"),
            **payload,
        }
    ).encode("utf-8")
    req = urllib.request.Request(
        url=f"{plaid_base_url()}{endpoint}",
        data=body,
        headers={"Content-Type": "application/json"},
        method="POST",
    )

    try:
        with urllib.request.urlopen(req, timeout=30) as response:
            return json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        raw = exc.read().decode("utf-8", errors="ignore")
        try:
            payload = json.loads(raw)
            message = payload.get("error_message") or payload.get("display_message") or raw or "Plaid request failed"
        except (TypeError, ValueError):
            message = raw or "Plaid request failed"
        raise ValueError(message) from exc
    except urllib.error.URLError as exc:
        raise ValueError("Unable to reach Plaid. Check network access and PLAID_ENV.") from exc


def serialize_bank_connection(connection):
    return {
        "id": connection.id,
        "provider": connection.provider,
        "institution_name": connection.institution_name or "Connected bank",
        "status": connection.status,
        "item_id": connection.item_id,
        "company_id": connection.company_id,
        "created_at": connection.created_at.isoformat() if connection.created_at else None,
        "updated_at": connection.updated_at.isoformat() if connection.updated_at else None,
    }


def get_or_create_tax_profile(company):
    profile = TaxProfile.query.filter_by(company_id=company.id).first()
    if profile:
        return profile

    profile = TaxProfile(
        org_id=company.org_id,
        company_id=company.id,
        jurisdiction_code=(os.getenv("DEFAULT_TAX_JURISDICTION", "generic") or "generic").strip().lower(),
        filing_frequency=(os.getenv("DEFAULT_TAX_FILING_FREQUENCY", "monthly") or "monthly").strip().lower(),
        currency_code=(os.getenv("DEFAULT_TAX_CURRENCY", "USD") or "USD").strip().upper(),
        sales_tax_name=(os.getenv("DEFAULT_SALES_TAX_NAME", "Sales Tax") or "Sales Tax").strip(),
        purchase_tax_name=(os.getenv("DEFAULT_PURCHASE_TAX_NAME", "Purchase Tax Credit") or "Purchase Tax Credit").strip(),
        indirect_tax_rate=float(os.getenv("DEFAULT_INDIRECT_TAX_RATE", "16") or 16),
        income_tax_rate=float(os.getenv("DEFAULT_INCOME_TAX_RATE", "30") or 30),
        period_start_month=int(os.getenv("DEFAULT_TAX_PERIOD_START_MONTH", "1") or 1),
    )
    db.session.add(profile)
    safe_commit()
    return profile


def serialize_tax_profile(profile):
    return {
        "jurisdiction_code": profile.jurisdiction_code,
        "filing_frequency": profile.filing_frequency,
        "registration_number": profile.registration_number or "",
        "currency_code": profile.currency_code,
        "sales_tax_name": profile.sales_tax_name,
        "purchase_tax_name": profile.purchase_tax_name,
        "indirect_tax_rate": round(float(profile.indirect_tax_rate or 0), 2),
        "income_tax_rate": round(float(profile.income_tax_rate or 0), 2),
        "period_start_month": int(profile.period_start_month or 1),
    }


def current_tax_period(profile):
    today = today_utc_date()
    if profile.filing_frequency == "annual":
        start = datetime.date(today.year, max(1, min(12, profile.period_start_month or 1)), 1)
        if today < start:
            start = datetime.date(today.year - 1, start.month, 1)
        end = datetime.date(start.year + 1, start.month, 1) - datetime.timedelta(days=1)
        return start, end

    if profile.filing_frequency == "quarterly":
        quarter_index = (today.month - 1) // 3
        start_month = quarter_index * 3 + 1
        start = datetime.date(today.year, start_month, 1)
        if start_month == 10:
            end = datetime.date(today.year + 1, 1, 1) - datetime.timedelta(days=1)
        else:
            end = datetime.date(today.year, start_month + 3, 1) - datetime.timedelta(days=1)
        return start, end

    start = datetime.date(today.year, today.month, 1)
    if today.month == 12:
        end = datetime.date(today.year + 1, 1, 1) - datetime.timedelta(days=1)
    else:
        end = datetime.date(today.year, today.month + 1, 1) - datetime.timedelta(days=1)
    return start, end


def build_tax_filing_preview(company, profile, period_start=None, period_end=None):
    start = parse_iso_date(period_start, "period_start") if period_start else None
    end = parse_iso_date(period_end, "period_end") if period_end else None
    if not start or not end:
        start, end = current_tax_period(profile)

    invoices, bills = refresh_finance_documents(company.id)
    scoped_invoices = [
        invoice for invoice in invoices
        if invoice.status not in {"draft", "cancelled"} and invoice.issue_date and start <= invoice.issue_date <= end
    ]
    scoped_bills = [
        bill for bill in bills
        if bill.status not in {"draft", "cancelled"} and bill.issue_date and start <= bill.issue_date <= end
    ]

    sales_tax_collected = sum(float(invoice.tax_amount or 0) for invoice in scoped_invoices)
    purchase_tax_credit = sum(float(bill.tax_amount or 0) for bill in scoped_bills)
    taxable_profit = sum(float(invoice.subtotal or 0) for invoice in scoped_invoices) - sum(float(bill.subtotal or 0) for bill in scoped_bills)
    estimated_income_tax = max(0.0, taxable_profit) * (float(profile.income_tax_rate or 0) / 100)

    return {
        "period_start": iso_date(start),
        "period_end": iso_date(end),
        "jurisdiction_code": profile.jurisdiction_code,
        "filing_frequency": profile.filing_frequency,
        "sales_tax_name": profile.sales_tax_name,
        "purchase_tax_name": profile.purchase_tax_name,
        "currency_code": profile.currency_code,
        "sales_tax_collected": round(sales_tax_collected, 2),
        "purchase_tax_credit": round(purchase_tax_credit, 2),
        "net_tax_due": round(sales_tax_collected - purchase_tax_credit, 2),
        "taxable_profit": round(taxable_profit, 2),
        "estimated_income_tax": round(estimated_income_tax, 2),
        "documents": {
            "sales_documents": len(scoped_invoices),
            "purchase_documents": len(scoped_bills),
        },
    }


def seed_chart_of_accounts(company):
    created = 0
    for template in DEFAULT_CHART_OF_ACCOUNTS:
        existing = LedgerAccount.query.filter_by(company_id=company.id, code=template["code"]).first()
        if existing:
            continue
        db.session.add(
            LedgerAccount(
                org_id=company.org_id,
                company_id=company.id,
                code=template["code"],
                name=template["name"],
                category=template["category"],
                subtype=template.get("subtype"),
                normal_balance=template.get("normal_balance", "debit"),
                description=template.get("description"),
                is_system=True,
                is_active=True,
            )
        )
        created += 1
    if created:
        db.session.flush()
    return created


def serialize_ledger_account(account):
    return {
        "id": account.id,
        "code": account.code,
        "name": account.name,
        "category": account.category,
        "subtype": account.subtype or "",
        "normal_balance": account.normal_balance,
        "is_system": bool(account.is_system),
        "is_active": bool(account.is_active),
        "description": account.description or "",
    }


def journal_lines_for(entry_id):
    return JournalLine.query.filter_by(journal_entry_id=entry_id).order_by(JournalLine.line_number.asc(), JournalLine.id.asc()).all()


def get_company_account(company_id, account_id=None, account_code=None):
    if account_id is not None:
        return LedgerAccount.query.filter_by(company_id=company_id, id=account_id).first()
    if account_code:
        return LedgerAccount.query.filter_by(company_id=company_id, code=str(account_code).strip()).first()
    return None


def generate_journal_number(company_id):
    next_number = JournalEntry.query.filter_by(company_id=company_id).count() + 1
    return f"JE-{int(company_id):03d}-{next_number:05d}"


def serialize_journal_entry(entry):
    lines = []
    for line in journal_lines_for(entry.id):
        account = db.session.get(LedgerAccount, line.account_id)
        lines.append(
            {
                "id": line.id,
                "account_id": line.account_id,
                "account_code": account.code if account else "",
                "account_name": account.name if account else "",
                "project_id": line.project_id,
                "description": line.description or "",
                "debit": round(float(line.debit or 0), 2),
                "credit": round(float(line.credit or 0), 2),
            }
        )

    return {
        "id": entry.id,
        "entry_number": entry.entry_number,
        "entry_date": iso_date(entry.entry_date),
        "memo": entry.memo,
        "reference": entry.reference or "",
        "source_type": entry.source_type,
        "source_id": entry.source_id,
        "status": entry.status,
        "reverses_entry_id": entry.reverses_entry_id,
        "lines": lines,
    }


def normalize_journal_lines(company, lines):
    if not isinstance(lines, list) or len(lines) < 2:
        raise ValueError("journal entry requires at least two lines")

    normalized = []
    debit_total = 0.0
    credit_total = 0.0

    for index, raw_line in enumerate(lines, start=1):
        payload = raw_line or {}
        account = get_company_account(company.id, payload.get("account_id"), payload.get("account_code"))
        if not account:
            raise ValueError(f"journal line {index} references an unknown account")

        debit = parse_money(payload.get("debit", 0), f"journal line {index} debit")
        credit = parse_money(payload.get("credit", 0), f"journal line {index} credit")
        if debit > 0 and credit > 0:
            raise ValueError(f"journal line {index} cannot contain both debit and credit")
        if debit <= 0 and credit <= 0:
            raise ValueError(f"journal line {index} must contain a debit or credit amount")

        normalized.append(
            {
                "account_id": account.id,
                "project_id": payload.get("project_id"),
                "description": (payload.get("description") or "").strip() or None,
                "debit": debit,
                "credit": credit,
            }
        )
        debit_total += debit
        credit_total += credit

    if round(debit_total, 2) != round(credit_total, 2):
        raise ValueError("journal entry debits and credits must balance")

    return normalized, round(debit_total, 2), round(credit_total, 2)


def post_journal_entry(company, user, entry_date, memo, lines, source_type="manual", source_id=None, reference=None):
    seed_chart_of_accounts(company)
    normalized_lines, _, _ = normalize_journal_lines(company, lines)
    entry = JournalEntry(
        org_id=company.org_id,
        company_id=company.id,
        entry_number=generate_journal_number(company.id),
        entry_date=entry_date,
        memo=memo,
        reference=(reference or "").strip() or None,
        source_type=(source_type or "manual").strip().lower(),
        source_id=source_id,
        status="posted",
        created_by=user.id,
    )
    db.session.add(entry)
    db.session.flush()

    for index, line in enumerate(normalized_lines, start=1):
        db.session.add(
            JournalLine(
                journal_entry_id=entry.id,
                account_id=line["account_id"],
                project_id=line["project_id"],
                line_number=index,
                description=line["description"],
                debit=line["debit"],
                credit=line["credit"],
            )
        )

    db.session.flush()
    return entry


def source_entry_exists(company_id, source_type, source_id):
    if source_id is None:
        return False
    return (
        JournalEntry.query.filter_by(
            company_id=company_id,
            source_type=(source_type or "").strip().lower(),
            source_id=source_id,
        )
        .filter(JournalEntry.status.in_(["posted", "reversed"]))
        .first()
        is not None
    )


def post_operational_entry(company, user, source_type, source_id, memo, lines, entry_date=None, reference=None):
    if source_entry_exists(company.id, source_type, source_id):
        return None
    return post_journal_entry(
        company,
        user,
        entry_date or today_utc_date(),
        memo=memo,
        lines=lines,
        source_type=source_type,
        source_id=source_id,
        reference=reference,
    )


def reverse_source_entries(company, user, source_type, source_id, memo):
    entries = (
        JournalEntry.query.filter_by(
            company_id=company.id,
            source_type=(source_type or "").strip().lower(),
            source_id=source_id,
            status="posted",
        )
        .order_by(JournalEntry.id.asc())
        .all()
    )
    reversed_entries = []
    for entry in entries:
        reversal_lines = []
        for line in journal_lines_for(entry.id):
            reversal_lines.append(
                {
                    "account_id": line.account_id,
                    "project_id": line.project_id,
                    "description": f"Reversal of {entry.entry_number}",
                    "debit": round(float(line.credit or 0), 2),
                    "credit": round(float(line.debit or 0), 2),
                }
            )
        reversal = post_journal_entry(
            company,
            user,
            today_utc_date(),
            memo=memo,
            lines=reversal_lines,
            source_type=f"{source_type}_reversal",
            source_id=source_id,
            reference=entry.entry_number,
        )
        reversal.reverses_entry_id = entry.id
        entry.status = "reversed"
        reversed_entries.append(reversal)
    return reversed_entries


def build_trial_balance(company):
    seed_chart_of_accounts(company)
    accounts = LedgerAccount.query.filter_by(company_id=company.id).order_by(LedgerAccount.code.asc()).all()
    lines = (
        db.session.query(JournalLine, JournalEntry)
        .join(JournalEntry, JournalLine.journal_entry_id == JournalEntry.id)
        .filter(JournalEntry.company_id == company.id, JournalEntry.status.in_(["posted", "reversed"]))
        .all()
    )
    balance_map = {account.id: {"debit": 0.0, "credit": 0.0} for account in accounts}
    for line, _entry in lines:
        bucket = balance_map.setdefault(line.account_id, {"debit": 0.0, "credit": 0.0})
        bucket["debit"] += float(line.debit or 0)
        bucket["credit"] += float(line.credit or 0)

    items = []
    debit_total = 0.0
    credit_total = 0.0
    for account in accounts:
        totals = balance_map.get(account.id, {"debit": 0.0, "credit": 0.0})
        debit_total += totals["debit"]
        credit_total += totals["credit"]
        items.append(
            {
                **serialize_ledger_account(account),
                "debit_total": round(totals["debit"], 2),
                "credit_total": round(totals["credit"], 2),
                "net_balance": round(totals["debit"] - totals["credit"], 2),
            }
        )

    return {
        "items": items,
        "debit_total": round(debit_total, 2),
        "credit_total": round(credit_total, 2),
        "balanced": round(debit_total, 2) == round(credit_total, 2),
    }


def build_account_register(company, account):
    rows = (
        db.session.query(JournalLine, JournalEntry)
        .join(JournalEntry, JournalLine.journal_entry_id == JournalEntry.id)
        .filter(JournalEntry.company_id == company.id, JournalLine.account_id == account.id)
        .order_by(JournalEntry.entry_date.asc(), JournalEntry.id.asc(), JournalLine.id.asc())
        .all()
    )

    running_balance = 0.0
    items = []
    for line, entry in rows:
        debit = round(float(line.debit or 0), 2)
        credit = round(float(line.credit or 0), 2)
        movement = debit - credit if account.normal_balance == "debit" else credit - debit
        running_balance = round(running_balance + movement, 2)
        items.append(
            {
                "entry_id": entry.id,
                "entry_number": entry.entry_number,
                "entry_date": iso_date(entry.entry_date),
                "memo": entry.memo,
                "reference": entry.reference or "",
                "debit": debit,
                "credit": credit,
                "running_balance": running_balance,
                "status": entry.status,
            }
        )

    return {
        "account": serialize_ledger_account(account),
        "items": items,
        "ending_balance": running_balance,
    }


def build_accounting_overview(company):
    trial_balance = build_trial_balance(company)
    recent_entries = (
        JournalEntry.query.filter_by(company_id=company.id)
        .order_by(JournalEntry.entry_date.desc(), JournalEntry.id.desc())
        .limit(8)
        .all()
    )
    return {
        "account_count": len(trial_balance["items"]),
        "journal_count": JournalEntry.query.filter_by(company_id=company.id).count(),
        "trial_balance": trial_balance,
        "recent_entries": [serialize_journal_entry(entry) for entry in recent_entries],
    }


def serialize_vendor_profile(vendor):
    return {
        "id": vendor.id,
        "vendor_name": vendor.vendor_name,
        "email": vendor.email or "",
        "tax_id": vendor.tax_id or "",
        "default_payment_rail": vendor.default_payment_rail,
        "remittance_reference": vendor.remittance_reference or "",
        "bank_last4": vendor.bank_last4 or "",
        "is_1099_eligible": bool(vendor.is_1099_eligible),
        "tax_form_type": vendor.tax_form_type,
        "tin_status": vendor.tin_status,
    }


def get_or_create_vendor_profile(company, vendor_name, defaults=None):
    normalized_name = (vendor_name or "").strip()
    if not normalized_name:
        raise ValueError("vendor_name is required")
    vendor = VendorProfile.query.filter_by(company_id=company.id, vendor_name=normalized_name).first()
    if vendor:
        return vendor

    defaults = defaults or {}
    vendor = VendorProfile(
        org_id=company.org_id,
        company_id=company.id,
        vendor_name=normalized_name,
        email=(defaults.get("email") or "").strip().lower() or None,
        tax_id=(defaults.get("tax_id") or "").strip() or None,
        default_payment_rail=(defaults.get("default_payment_rail") or "ach").strip().lower(),
        remittance_reference=(defaults.get("remittance_reference") or "").strip() or None,
        bank_last4=(defaults.get("bank_last4") or "").strip()[-4:] or None,
        is_1099_eligible=parse_bool(defaults.get("is_1099_eligible"), False),
        tax_form_type=(defaults.get("tax_form_type") or "1099-NEC").strip().upper(),
        tin_status=(defaults.get("tin_status") or "pending").strip().lower(),
    )
    db.session.add(vendor)
    db.session.flush()
    return vendor


def serialize_disbursement(disbursement):
    bill = db.session.get(VendorBill, disbursement.bill_id)
    vendor = db.session.get(VendorProfile, disbursement.vendor_profile_id) if disbursement.vendor_profile_id else None
    return {
        "id": disbursement.id,
        "bill_id": disbursement.bill_id,
        "bill_number": bill.bill_number if bill else "",
        "vendor_name": vendor.vendor_name if vendor else (bill.vendor_name if bill else ""),
        "payment_rail": disbursement.payment_rail,
        "status": disbursement.status,
        "scheduled_date": iso_date(disbursement.scheduled_date),
        "processed_at": disbursement.processed_at.isoformat() if disbursement.processed_at else None,
        "amount": round(float(disbursement.amount or 0), 2),
        "reference": disbursement.reference or "",
        "confirmation_code": disbursement.confirmation_code or "",
        "compliance_status": disbursement.compliance_status,
    }


def build_1099_summary(company, year=None):
    scoped_year = year or today_utc_date().year
    vendors = VendorProfile.query.filter_by(company_id=company.id).order_by(VendorProfile.vendor_name.asc()).all()
    items = []
    total_reportable = 0.0
    for vendor in vendors:
        payments_total = 0.0
        vendor_bills = VendorBill.query.filter_by(company_id=company.id, vendor_name=vendor.vendor_name).all()
        for bill in vendor_bills:
            payments = VendorPayment.query.filter_by(bill_id=bill.id).all()
            for payment in payments:
                if payment.payment_date and payment.payment_date.year == scoped_year:
                    payments_total += float(payment.amount or 0)

        if payments_total <= 0 and not vendor.is_1099_eligible:
            continue

        total_reportable += payments_total
        items.append(
            {
                **serialize_vendor_profile(vendor),
                "year": scoped_year,
                "reportable_payments": round(payments_total, 2),
                "threshold_status": "ready" if payments_total >= 600 else "below_threshold",
            }
        )

    return {
        "year": scoped_year,
        "items": items,
        "reportable_total": round(total_reportable, 2),
        "ready_count": sum(1 for item in items if item["threshold_status"] == "ready"),
    }


def build_bill_pay_summary(company):
    disbursements = BillDisbursement.query.filter_by(company_id=company.id).order_by(BillDisbursement.id.desc()).all()
    return {
        "scheduled_count": sum(1 for item in disbursements if item.status == "scheduled"),
        "processing_count": sum(1 for item in disbursements if item.status == "processing"),
        "completed_count": sum(1 for item in disbursements if item.status == "completed"),
        "scheduled_amount": round(sum(float(item.amount or 0) for item in disbursements if item.status == "scheduled"), 2),
        "completed_amount": round(sum(float(item.amount or 0) for item in disbursements if item.status == "completed"), 2),
        "items": [serialize_disbursement(item) for item in disbursements[:20]],
    }


def serialize_reconciliation_rule(rule):
    return {
        "id": rule.id,
        "name": rule.name,
        "keyword": rule.keyword or "",
        "direction": rule.direction,
        "min_amount": round(float(rule.min_amount or 0), 2) if rule.min_amount is not None else None,
        "max_amount": round(float(rule.max_amount or 0), 2) if rule.max_amount is not None else None,
        "auto_action": rule.auto_action,
        "target_reference": rule.target_reference or "",
        "exception_type": rule.exception_type or "",
        "priority": rule.priority,
        "is_active": bool(rule.is_active),
    }


def serialize_reconciliation_exception(exception):
    transaction = db.session.get(BankFeedTransaction, exception.bank_transaction_id)
    return {
        "id": exception.id,
        "bank_transaction_id": exception.bank_transaction_id,
        "transaction": serialize_bank_transaction(transaction) if transaction else None,
        "exception_type": exception.exception_type,
        "notes": exception.notes or "",
        "status": exception.status,
        "resolved_at": exception.resolved_at.isoformat() if exception.resolved_at else None,
    }


def rule_matches_transaction(rule, transaction):
    description = (transaction.description or "").lower()
    absolute_amount = round(abs(float(transaction.amount or 0)), 2)
    direction = "inflow" if float(transaction.amount or 0) >= 0 else "outflow"

    if rule.direction != "any" and direction != rule.direction:
        return False
    if rule.keyword and rule.keyword.strip().lower() not in description:
        return False
    if rule.min_amount is not None and absolute_amount < float(rule.min_amount):
        return False
    if rule.max_amount is not None and absolute_amount > float(rule.max_amount):
        return False
    return True


def auto_apply_reconciliation_rules(company, user=None):
    rules = (
        ReconciliationRule.query.filter_by(company_id=company.id, is_active=True)
        .order_by(ReconciliationRule.priority.asc(), ReconciliationRule.id.asc())
        .all()
    )
    if not rules:
        return {"matched": 0, "exceptions": 0, "items": []}

    transactions = (
        BankFeedTransaction.query.filter(
            BankFeedTransaction.company_id == company.id,
            BankFeedTransaction.status.in_(["unmatched", "rule_matched"]),
        )
        .order_by(BankFeedTransaction.posted_at.desc(), BankFeedTransaction.id.desc())
        .all()
    )

    matched = 0
    exceptions = 0
    items = []
    actor_id = user.id if user else 0
    for transaction in transactions:
        selected_rule = None
        for rule in rules:
            if rule_matches_transaction(rule, transaction):
                selected_rule = rule
                break
        if not selected_rule:
            continue

        payload = {
            "transaction": serialize_bank_transaction(transaction),
            "rule": serialize_reconciliation_rule(selected_rule),
        }
        if selected_rule.auto_action == "flag_exception":
            existing_exception = ReconciliationException.query.filter_by(
                company_id=company.id,
                bank_transaction_id=transaction.id,
                status="open",
            ).first()
            if not existing_exception:
                db.session.add(
                    ReconciliationException(
                        org_id=company.org_id,
                        company_id=company.id,
                        bank_transaction_id=transaction.id,
                        exception_type=(selected_rule.exception_type or "review_required").strip() or "review_required",
                        notes=f"Raised by rule {selected_rule.name}",
                        created_by=actor_id,
                    )
                )
                exceptions += 1
            transaction.status = "exception"
            payload["result"] = "exception"
        else:
            transaction.status = "rule_matched"
            matched += 1
            payload["result"] = "rule_matched"
        items.append(payload)

    if items:
        safe_commit()
    return {"matched": matched, "exceptions": exceptions, "items": items}


def build_reconciliation_workspace(company):
    transactions = (
        BankFeedTransaction.query.filter_by(company_id=company.id)
        .order_by(BankFeedTransaction.posted_at.desc(), BankFeedTransaction.id.desc())
        .limit(80)
        .all()
    )
    rules = (
        ReconciliationRule.query.filter_by(company_id=company.id)
        .order_by(ReconciliationRule.priority.asc(), ReconciliationRule.id.asc())
        .all()
    )
    exceptions = (
        ReconciliationException.query.filter_by(company_id=company.id)
        .order_by(ReconciliationException.created_at.desc(), ReconciliationException.id.desc())
        .limit(20)
        .all()
    )
    return {
        "summary": {
            "unmatched": sum(1 for row in transactions if row.status == "unmatched"),
            "rule_matched": sum(1 for row in transactions if row.status == "rule_matched"),
            "matched": sum(1 for row in transactions if row.status == "matched"),
            "exceptions": sum(1 for row in transactions if row.status == "exception"),
        },
        "transactions": [serialize_bank_transaction(row) for row in transactions],
        "rules": [serialize_reconciliation_rule(rule) for rule in rules],
        "exceptions": [serialize_reconciliation_exception(item) for item in exceptions],
    }


def serialize_tax_filing(filing):
    payload = {}
    try:
        payload = json.loads(filing.payload_json or "{}")
    except (TypeError, ValueError):
        payload = {}

    return {
        "id": filing.id,
        "jurisdiction_code": filing.jurisdiction_code,
        "filing_frequency": filing.filing_frequency,
        "filing_type": filing.filing_type,
        "period_start": iso_date(filing.period_start),
        "period_end": iso_date(filing.period_end),
        "status": filing.status,
        "reference": filing.reference or "",
        "prepared_at": filing.prepared_at.isoformat() if filing.prepared_at else None,
        "submitted_at": filing.submitted_at.isoformat() if filing.submitted_at else None,
        "payload": payload,
    }


def build_tax_filing_package(company, profile, period_start=None, period_end=None):
    preview = build_tax_filing_preview(company, profile, period_start=period_start, period_end=period_end)
    jurisdiction = TAX_JURISDICTION_LIBRARY.get(profile.jurisdiction_code, TAX_JURISDICTION_LIBRARY["generic"])
    labels = jurisdiction["return_labels"]
    taxable_sales = round(sum(float(invoice.subtotal or 0) for invoice in refresh_finance_documents(company.id)[0]), 2)
    line_values = {
        "taxable_sales": taxable_sales,
        "tax_collected": preview["sales_tax_collected"],
        "tax_credit": preview["purchase_tax_credit"],
        "net_tax_due": preview["net_tax_due"],
        "vatable_sales": taxable_sales,
        "output_vat": preview["sales_tax_collected"],
        "input_vat": preview["purchase_tax_credit"],
        "vat_payable": preview["net_tax_due"],
        "sales_tax_collected": preview["sales_tax_collected"],
        "exempt_sales": 0.0,
        "net_sales_tax_due": preview["net_tax_due"],
        "vat_due_sales": preview["sales_tax_collected"],
        "vat_reclaimed": preview["purchase_tax_credit"],
        "net_vat_due": preview["net_tax_due"],
        "total_sales_ex_vat": taxable_sales,
        "gst_hst_collected": preview["sales_tax_collected"],
        "input_tax_credits": preview["purchase_tax_credit"],
        "net_gst_hst_due": preview["net_tax_due"],
        "taxable_supplies": taxable_sales,
        "gst_collected": preview["sales_tax_collected"],
        "gst_credits": preview["purchase_tax_credit"],
        "net_gst_due": preview["net_tax_due"],
    }
    return {
        "jurisdiction": {
            "code": profile.jurisdiction_code,
            "name": jurisdiction["name"],
            "filing_type": jurisdiction["filing_type"],
        },
        "profile": serialize_tax_profile(profile),
        "preview": preview,
        "return_lines": [{"label": label, "value": round(float(line_values.get(label, 0) or 0), 2)} for label in labels],
    }


def worker_display_name(employee_id=None, contractor_id=None):
    if employee_id:
        employee = db.session.get(EmployeeProfile, employee_id)
        return employee.full_name if employee else "Employee"
    if contractor_id:
        contractor = db.session.get(ContractorProfile, contractor_id)
        return contractor.full_name if contractor else "Contractor"
    return "Unknown Worker"


def serialize_employee(employee):
    return {
        "id": employee.id,
        "full_name": employee.full_name,
        "email": employee.email or "",
        "pay_type": employee.pay_type,
        "hourly_rate": round(float(employee.hourly_rate or 0), 2),
        "salary_amount": round(float(employee.salary_amount or 0), 2),
        "withholding_rate": round(float(employee.withholding_rate or 0), 2),
        "benefit_rate": round(float(employee.benefit_rate or 0), 2),
        "is_active": bool(employee.is_active),
    }


def serialize_contractor(contractor):
    return {
        "id": contractor.id,
        "full_name": contractor.full_name,
        "email": contractor.email or "",
        "tax_id": contractor.tax_id or "",
        "default_rate": round(float(contractor.default_rate or 0), 2),
        "is_1099_eligible": bool(contractor.is_1099_eligible),
        "tax_form_type": contractor.tax_form_type,
        "is_active": bool(contractor.is_active),
    }


def serialize_time_entry(entry):
    return {
        "id": entry.id,
        "employee_id": entry.employee_id,
        "contractor_id": entry.contractor_id,
        "project_id": entry.project_id,
        "worker_name": worker_display_name(entry.employee_id, entry.contractor_id),
        "work_date": iso_date(entry.work_date),
        "hours": round(float(entry.hours or 0), 2),
        "hourly_cost": round(float(entry.hourly_cost or 0), 2),
        "billable_rate": round(float(entry.billable_rate or 0), 2),
        "status": entry.status,
        "description": entry.description or "",
    }


def serialize_mileage_entry(entry):
    reimbursement = round(float(entry.miles or 0) * float(entry.rate_per_mile or 0), 2)
    return {
        "id": entry.id,
        "employee_id": entry.employee_id,
        "contractor_id": entry.contractor_id,
        "project_id": entry.project_id,
        "worker_name": worker_display_name(entry.employee_id, entry.contractor_id),
        "trip_date": iso_date(entry.trip_date),
        "miles": round(float(entry.miles or 0), 2),
        "rate_per_mile": round(float(entry.rate_per_mile or 0), 4),
        "reimbursement": reimbursement,
        "purpose": entry.purpose or "",
        "status": entry.status,
    }


def serialize_payroll_run(payroll):
    lines = PayrollLine.query.filter_by(payroll_run_id=payroll.id).order_by(PayrollLine.id.asc()).all()
    return {
        "id": payroll.id,
        "payroll_number": payroll.payroll_number,
        "period_start": iso_date(payroll.period_start),
        "period_end": iso_date(payroll.period_end),
        "pay_date": iso_date(payroll.pay_date),
        "status": payroll.status,
        "gross_pay": round(float(payroll.gross_pay or 0), 2),
        "withholding_total": round(float(payroll.withholding_total or 0), 2),
        "benefit_total": round(float(payroll.benefit_total or 0), 2),
        "mileage_reimbursement_total": round(float(payroll.mileage_reimbursement_total or 0), 2),
        "net_cash": round(float(payroll.net_cash or 0), 2),
        "lines": [
            {
                "id": line.id,
                "employee_id": line.employee_id,
                "employee_name": worker_display_name(employee_id=line.employee_id),
                "regular_hours": round(float(line.regular_hours or 0), 2),
                "gross_pay": round(float(line.gross_pay or 0), 2),
                "withholding_amount": round(float(line.withholding_amount or 0), 2),
                "benefit_amount": round(float(line.benefit_amount or 0), 2),
                "mileage_reimbursement": round(float(line.mileage_reimbursement or 0), 2),
                "net_pay": round(float(line.net_pay or 0), 2),
            }
            for line in lines
        ],
    }


def generate_payroll_number(company_id):
    next_number = PayrollRun.query.filter_by(company_id=company_id).count() + 1
    return f"PAY-{int(company_id):03d}-{next_number:05d}"


def build_workforce_overview(company):
    today = today_utc_date()
    month_start = datetime.date(today.year, today.month, 1)
    employees = EmployeeProfile.query.filter_by(company_id=company.id).all()
    contractors = ContractorProfile.query.filter_by(company_id=company.id).all()
    time_entries = TimeEntry.query.filter_by(company_id=company.id).all()
    mileage_entries = MileageEntry.query.filter_by(company_id=company.id).all()
    payroll_runs = PayrollRun.query.filter_by(company_id=company.id).all()

    return {
        "employee_count": sum(1 for row in employees if row.is_active),
        "contractor_count": sum(1 for row in contractors if row.is_active),
        "hours_this_month": round(
            sum(float(row.hours or 0) for row in time_entries if row.work_date and row.work_date >= month_start),
            2,
        ),
        "mileage_this_month": round(
            sum(float(row.miles or 0) for row in mileage_entries if row.trip_date and row.trip_date >= month_start),
            2,
        ),
        "payroll_this_month": round(
            sum(float(run.net_cash or 0) for run in payroll_runs if run.pay_date and run.pay_date >= month_start),
            2,
        ),
        "contractor_1099_exposure": round(
            sum(float(row.hours or 0) * float(row.hourly_cost or 0) for row in time_entries if row.contractor_id),
            2,
        ),
    }


def serialize_inventory_item(item):
    return {
        "id": item.id,
        "sku": item.sku,
        "name": item.name,
        "category": item.category or "",
        "quantity_on_hand": round(float(item.quantity_on_hand or 0), 2),
        "reorder_point": round(float(item.reorder_point or 0), 2),
        "reorder_quantity": round(float(item.reorder_quantity or 0), 2),
        "unit_cost": round(float(item.unit_cost or 0), 2),
        "unit_price": round(float(item.unit_price or 0), 2),
        "inventory_value": round(float(item.quantity_on_hand or 0) * float(item.unit_cost or 0), 2),
        "preferred_vendor_name": item.preferred_vendor_name or "",
        "needs_reorder": float(item.quantity_on_hand or 0) <= float(item.reorder_point or 0),
    }


def purchase_order_lines_for(po_id):
    return PurchaseOrderLine.query.filter_by(purchase_order_id=po_id).order_by(PurchaseOrderLine.id.asc()).all()


def serialize_purchase_order(po):
    lines = purchase_order_lines_for(po.id)
    return {
        "id": po.id,
        "po_number": po.po_number,
        "vendor_name": po.vendor_name,
        "status": po.status,
        "issue_date": iso_date(po.issue_date),
        "expected_date": iso_date(po.expected_date),
        "notes": po.notes or "",
        "items": [
            {
                "id": line.id,
                "inventory_item_id": line.inventory_item_id,
                "sku": line.sku or "",
                "description": line.description,
                "quantity": round(float(line.quantity or 0), 2),
                "unit_cost": round(float(line.unit_cost or 0), 2),
                "received_quantity": round(float(line.received_quantity or 0), 2),
            }
            for line in lines
        ],
        "ordered_total": round(sum(float(line.quantity or 0) * float(line.unit_cost or 0) for line in lines), 2),
        "received_total": round(sum(float(line.received_quantity or 0) * float(line.unit_cost or 0) for line in lines), 2),
    }


def serialize_inventory_movement(movement):
    item = db.session.get(InventoryItem, movement.inventory_item_id)
    return {
        "id": movement.id,
        "inventory_item_id": movement.inventory_item_id,
        "sku": item.sku if item else "",
        "item_name": item.name if item else "",
        "project_id": movement.project_id,
        "movement_type": movement.movement_type,
        "quantity_delta": round(float(movement.quantity_delta or 0), 2),
        "unit_cost": round(float(movement.unit_cost or 0), 2),
        "reference": movement.reference or "",
        "occurred_at": movement.occurred_at.isoformat() if movement.occurred_at else None,
    }


def build_inventory_summary(company):
    items = InventoryItem.query.filter_by(company_id=company.id).order_by(InventoryItem.name.asc()).all()
    purchase_orders = PurchaseOrder.query.filter_by(company_id=company.id).order_by(PurchaseOrder.id.desc()).limit(12).all()
    movements = InventoryMovement.query.filter_by(company_id=company.id).order_by(InventoryMovement.id.desc()).limit(20).all()
    return {
        "item_count": len(items),
        "inventory_value": round(sum(float(item.quantity_on_hand or 0) * float(item.unit_cost or 0) for item in items), 2),
        "low_stock_count": sum(1 for item in items if float(item.quantity_on_hand or 0) <= float(item.reorder_point or 0)),
        "open_purchase_orders": sum(1 for po in purchase_orders if po.status in {"draft", "ordered", "partial"}),
        "reorder_items": [serialize_inventory_item(item) for item in items if float(item.quantity_on_hand or 0) <= float(item.reorder_point or 0)],
        "items": [serialize_inventory_item(item) for item in items[:20]],
        "purchase_orders": [serialize_purchase_order(po) for po in purchase_orders],
        "movements": [serialize_inventory_movement(movement) for movement in movements],
    }


def receive_purchase_order_items(company, user, purchase_order, receipt_lines):
    if purchase_order.status not in {"ordered", "partial"}:
        raise ValueError("purchase order must be ordered before receiving")

    normalized_receipts = {}
    for raw_line in receipt_lines or []:
        line_id = raw_line.get("line_id")
        if line_id is None:
            continue
        normalized_receipts[int(line_id)] = parse_money(raw_line.get("quantity", 0), f"receipt quantity for line {line_id}")

    if not normalized_receipts:
        raise ValueError("receipt requires at least one line quantity")

    lines = purchase_order_lines_for(purchase_order.id)
    received_any = False
    total_receipt_value = 0.0
    for line in lines:
        receipt_quantity = normalized_receipts.get(line.id, 0.0)
        if receipt_quantity <= 0:
            continue
        outstanding = round(float(line.quantity or 0) - float(line.received_quantity or 0), 2)
        if receipt_quantity - outstanding > 0.01:
            raise ValueError("receipt quantity exceeds outstanding quantity")
        line.received_quantity = round(float(line.received_quantity or 0) + receipt_quantity, 2)
        received_any = True
        total_receipt_value += round(receipt_quantity * float(line.unit_cost or 0), 2)

        item = db.session.get(InventoryItem, line.inventory_item_id) if line.inventory_item_id else InventoryItem.query.filter_by(company_id=company.id, sku=line.sku).first()
        if item:
            item.quantity_on_hand = round(float(item.quantity_on_hand or 0) + receipt_quantity, 2)
            db.session.add(
                InventoryMovement(
                    org_id=company.org_id,
                    company_id=company.id,
                    inventory_item_id=item.id,
                    movement_type="purchase_receipt",
                    quantity_delta=receipt_quantity,
                    unit_cost=float(line.unit_cost or 0),
                    reference=purchase_order.po_number,
                )
            )

    if not received_any:
        raise ValueError("no receipt quantity greater than zero was provided")

    all_received = all(round(float(line.received_quantity or 0), 2) >= round(float(line.quantity or 0), 2) for line in lines)
    purchase_order.status = "received" if all_received else "partial"

    post_operational_entry(
        company,
        user,
        source_type="purchase_order_receipt",
        source_id=purchase_order.id,
        memo=f"Inventory receipt for {purchase_order.po_number}",
        lines=[
            {"account_code": "1200", "debit": total_receipt_value, "credit": 0, "description": purchase_order.vendor_name},
            {"account_code": "2000", "debit": 0, "credit": total_receipt_value, "description": purchase_order.vendor_name},
        ],
        entry_date=today_utc_date(),
        reference=purchase_order.po_number,
    )


def serialize_project(project):
    return {
        "id": project.id,
        "project_code": project.project_code,
        "name": project.name,
        "customer_name": project.customer_name or "",
        "status": project.status,
        "budget_revenue": round(float(project.budget_revenue or 0), 2),
        "budget_cost": round(float(project.budget_cost or 0), 2),
        "notes": project.notes or "",
    }


def serialize_project_cost_entry(entry):
    return {
        "id": entry.id,
        "project_id": entry.project_id,
        "entry_type": entry.entry_type,
        "description": entry.description,
        "amount": round(float(entry.amount or 0), 2),
        "reference": entry.reference or "",
        "work_date": iso_date(entry.work_date),
    }


def build_project_summary(company):
    projects = Project.query.filter_by(company_id=company.id).order_by(Project.updated_at.desc(), Project.id.desc()).all()
    items = []
    total_revenue = 0.0
    total_cost = 0.0
    for project in projects:
        manual_entries = ProjectCostEntry.query.filter_by(project_id=project.id).all()
        time_entries = TimeEntry.query.filter_by(company_id=company.id, project_id=project.id).all()
        mileage_entries = MileageEntry.query.filter_by(company_id=company.id, project_id=project.id).all()
        inventory_movements = InventoryMovement.query.filter_by(company_id=company.id, project_id=project.id).all()

        manual_revenue = sum(float(entry.amount or 0) for entry in manual_entries if entry.entry_type == "revenue")
        manual_cost = sum(float(entry.amount or 0) for entry in manual_entries if entry.entry_type != "revenue")
        time_revenue = sum(float(entry.hours or 0) * float(entry.billable_rate or 0) for entry in time_entries)
        time_cost = sum(float(entry.hours or 0) * float(entry.hourly_cost or 0) for entry in time_entries)
        mileage_cost = sum(float(entry.miles or 0) * float(entry.rate_per_mile or 0) for entry in mileage_entries)
        inventory_cost = sum(abs(float(entry.quantity_delta or 0)) * float(entry.unit_cost or 0) for entry in inventory_movements if float(entry.quantity_delta or 0) < 0)
        actual_revenue = round(manual_revenue + time_revenue, 2)
        actual_cost = round(manual_cost + time_cost + mileage_cost + inventory_cost, 2)
        margin = round(actual_revenue - actual_cost, 2)
        total_revenue += actual_revenue
        total_cost += actual_cost
        items.append(
            {
                **serialize_project(project),
                "actual_revenue": actual_revenue,
                "actual_cost": actual_cost,
                "margin": margin,
                "budget_delta_revenue": round(actual_revenue - float(project.budget_revenue or 0), 2),
                "budget_delta_cost": round(actual_cost - float(project.budget_cost or 0), 2),
            }
        )

    return {
        "items": items,
        "total_revenue": round(total_revenue, 2),
        "total_cost": round(total_cost, 2),
        "total_margin": round(total_revenue - total_cost, 2),
    }


def serialize_integration_connection(connection):
    payload = {}
    try:
        payload = json.loads(connection.config_json or "{}")
    except (TypeError, ValueError):
        payload = {}
    return {
        "id": connection.id,
        "provider": connection.provider,
        "category": connection.category,
        "status": connection.status,
        "config": payload,
        "last_synced_at": connection.last_synced_at.isoformat() if connection.last_synced_at else None,
    }


def seed_integration_connections(company):
    created = 0
    for item in INTEGRATION_CATALOG:
        existing = IntegrationConnection.query.filter_by(company_id=company.id, provider=item["provider"]).first()
        if existing:
            continue
        db.session.add(
            IntegrationConnection(
                org_id=company.org_id,
                company_id=company.id,
                provider=item["provider"],
                category=item["category"],
                status="available",
                config_json=json.dumps({"description": item["description"]}),
            )
        )
        created += 1
    if created:
        db.session.flush()
    return created


def build_accountant_toolkit(company):
    trial_balance = build_trial_balance(company)
    projects = build_project_summary(company)
    inventory = build_inventory_summary(company)
    workforce = build_workforce_overview(company)
    finance = calculate_finance_summary(company)
    tax = calculate_tax_summary(company)
    return {
        "trial_balance": trial_balance,
        "projects": projects,
        "inventory": {
            "inventory_value": inventory["inventory_value"],
            "low_stock_count": inventory["low_stock_count"],
            "open_purchase_orders": inventory["open_purchase_orders"],
        },
        "workforce": workforce,
        "receivables": finance["open_receivables"],
        "payables": finance["open_payables"],
        "tax_due": tax["net_tax_due"],
    }


def sync_plaid_connection(connection):
    cursor = connection.sync_cursor
    has_more = True
    total_added = 0
    total_modified = 0
    total_removed = 0

    while has_more:
        response = call_plaid(
            "/transactions/sync",
            {
                "access_token": connection.access_token,
                "cursor": cursor,
            },
        )

        for transaction in response.get("added", []):
            plaid_transaction_id = transaction.get("transaction_id")
            if not plaid_transaction_id:
                continue

            posted_at = parse_iso_date(transaction.get("date"), "date", today_utc_date())
            raw_amount = round(float(transaction.get("amount") or 0), 2)
            amount = round(-raw_amount, 2)
            existing = BankFeedTransaction.query.filter_by(
                org_id=connection.org_id,
                company_id=connection.company_id,
                reference=f"plaid:{plaid_transaction_id}",
            ).first()
            description = (
                transaction.get("merchant_name")
                or transaction.get("name")
                or transaction.get("authorized_description")
                or "Plaid transaction"
            )

            if existing:
                existing.posted_at = posted_at
                existing.description = description
                existing.amount = amount
                existing.raw_payload = json.dumps(transaction)
            else:
                db.session.add(
                    BankFeedTransaction(
                        org_id=connection.org_id,
                        company_id=connection.company_id,
                        posted_at=posted_at,
                        description=description,
                        amount=amount,
                        reference=f"plaid:{plaid_transaction_id}",
                        raw_payload=json.dumps(transaction),
                    )
                )
                total_added += 1

        for transaction in response.get("modified", []):
            plaid_transaction_id = transaction.get("transaction_id")
            if not plaid_transaction_id:
                continue
            existing = BankFeedTransaction.query.filter_by(
                org_id=connection.org_id,
                company_id=connection.company_id,
                reference=f"plaid:{plaid_transaction_id}",
            ).first()
            if not existing:
                continue
            existing.posted_at = parse_iso_date(transaction.get("date"), "date", existing.posted_at)
            existing.description = (
                transaction.get("merchant_name")
                or transaction.get("name")
                or transaction.get("authorized_description")
                or existing.description
            )
            existing.amount = round(-float(transaction.get("amount") or 0), 2)
            existing.raw_payload = json.dumps(transaction)
            total_modified += 1

        for removed in response.get("removed", []):
            plaid_transaction_id = removed.get("transaction_id")
            if not plaid_transaction_id:
                continue
            existing = BankFeedTransaction.query.filter_by(
                org_id=connection.org_id,
                company_id=connection.company_id,
                reference=f"plaid:{plaid_transaction_id}",
            ).first()
            if existing:
                db.session.delete(existing)
                total_removed += 1

        cursor = response.get("next_cursor")
        has_more = bool(response.get("has_more"))

    connection.sync_cursor = cursor or connection.sync_cursor
    connection.status = "connected"
    connection.updated_at = datetime.datetime.now(datetime.UTC)

    if not safe_commit():
        raise ValueError("database error while syncing Plaid transactions")

    return {
        "added": total_added,
        "modified": total_modified,
        "removed": total_removed,
        "cursor": connection.sync_cursor,
    }


def maintenance_state():
    enabled = (os.getenv("MAINTENANCE_MODE", "0").strip().lower() in {"1", "true", "yes", "on"})
    message = (os.getenv("MAINTENANCE_MESSAGE", MAINTENANCE_DEFAULT_MESSAGE) or MAINTENANCE_DEFAULT_MESSAGE).strip()
    return {"maintenance": enabled, "message": message}


def default_subtype_for(ledger_type):
    if ledger_type in {"asset", "liability"}:
        return "current"
    if ledger_type in {"revenue", "expense"}:
        return "operating"
    return "equity"


LEDGER_COLUMN_ALIASES = {
    "account": {"account", "name", "account name", "description"},
    "type": {"type"},
    "subtype": {"subtype", "class", "category"},
    "amount": {"amount", "value"},
    "debit": {"debit", "dr"},
    "credit": {"credit", "cr"},
    "depreciation": {"depreciation", "depreciation_amount"},
}


def normalize_column_label(value):
    return re.sub(r"\s+", " ", str(value or "").strip().lower())


def detect_column_role(column):
    key = normalize_column_label(column)
    for role, aliases in LEDGER_COLUMN_ALIASES.items():
        if key in aliases:
            return role
    return None


def uploaded_file_seek(uploaded_file):
    seek = getattr(uploaded_file, "seek", None)
    if callable(seek):
        seek(0)
        return

    stream = getattr(uploaded_file, "stream", None)
    if stream is not None and hasattr(stream, "seek"):
        stream.seek(0)


def read_tabular_dataframe(uploaded_file, reader, **kwargs):
    frame = reader(uploaded_file, **kwargs)
    roles = {detect_column_role(column) for column in frame.columns}
    roles.discard(None)

    # Trial balances are often uploaded as raw rows with no header row at all.
    if not roles and frame.shape[1] <= 2:
        uploaded_file_seek(uploaded_file)
        return reader(uploaded_file, header=None, **kwargs)

    return frame


def is_blank_cell(value):
    if value is None:
        return True
    if isinstance(value, float) and pd.isna(value):
        return True
    return str(value).strip() == ""


def parse_numeric_cell(value):
    if value is None:
        return None
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        if pd.isna(value):
            return None
        return float(value)

    text = str(value).strip()
    if not text:
        return None

    negative = False
    if text.startswith("(") and text.endswith(")"):
        negative = True
        text = text[1:-1].strip()

    if text.endswith("-"):
        negative = True
        text = text[:-1].strip()

    text = re.sub(r"(?i)\b(dr|cr)\b", "", text)
    text = text.replace(",", "").replace("$", "").replace("€", "").replace("£", "").strip()
    if not text:
        return None

    try:
        number = float(text)
    except ValueError:
        return None
    return -number if negative else number


def normalize_account_key(account_name):
    return re.sub(r"[^a-z0-9]+", " ", str(account_name or "").strip().lower()).strip()


def should_skip_derived_label(account_name):
    key = normalize_account_key(account_name)
    if not key:
        return True

    derived_labels = {
        "cost of raw materials consumed",
        "prime cost",
        "total factory overheads",
        "cost of goods manufactured",
        "cost of goods manufactured 445",
    }
    return key in derived_labels


def infer_trial_balance_account(account_name):
    raw_name = str(account_name or "").strip()
    key = normalize_account_key(raw_name)
    title = raw_name.title() or "Unclassified Entry"

    if not key:
        return {"account": "Unclassified Entry", "type": "expense", "subtype": "operating"}

    if "closing raw material" in key:
        return {"account": "Closing Raw Materials", "type": "asset", "subtype": "current"}
    if "opening raw material" in key or "raw materials opening" in key:
        return {"account": "Raw Materials Opening Stock", "type": "asset", "subtype": "current"}
    if "opening work in progress" in key or "opening wip" in key:
        return {"account": "Opening Work in Progress", "type": "asset", "subtype": "current"}
    if "closing work in progress" in key or "closing wip" in key:
        return {"account": "Closing Work in Progress", "type": "asset", "subtype": "current"}
    if "closing stock" in key:
        return {"account": "Closing Stock", "type": "asset", "subtype": "current"}
    if key.startswith("stock ") or key == "stock" or "opening stock" in key:
        return {"account": "Opening Stock", "type": "asset", "subtype": "current"}

    if "returns inward" in key or "return inward" in key:
        return {"account": "Sales Returns", "type": "expense", "subtype": "operating"}
    if "returns outward" in key or "return outward" in key:
        return {"account": "Returns Outwards", "type": "expense", "subtype": "operating"}
    if "carriage inward" in key:
        return {"account": "Carriage Inwards", "type": "expense", "subtype": "operating"}
    if "carriage outward" in key:
        return {"account": "Carriage Outwards", "type": "expense", "subtype": "operating"}
    if "direct manufacturing labor" in key or "direct manufacturing labour" in key:
        return {"account": "Direct Manufacturing Labor", "type": "expense", "subtype": "operating"}
    if "factory indirect labor" in key or "factory indirect labour" in key:
        return {"account": "Factory Indirect Labor", "type": "expense", "subtype": "operating"}
    if "factory utilit" in key:
        return {"account": "Factory Utilities", "type": "expense", "subtype": "operating"}
    if "depreciation of factory equipment" in key:
        return {"account": "Depreciation of Factory Equipment", "type": "expense", "subtype": "operating"}
    if "purchase" in key:
        return {"account": "Purchases", "type": "expense", "subtype": "operating"}
    if "sales" in key or "turnover" in key:
        return {"account": "Sales Revenue", "type": "revenue", "subtype": "operating"}
    if "salary" in key or "wages" in key or "payroll" in key:
        return {"account": "Payroll Expenses", "type": "expense", "subtype": "operating"}
    if key == "rent" or "rent expense" in key:
        return {"account": "Rent Expense", "type": "expense", "subtype": "operating"}
    if "insurance" in key:
        return {"account": "Insurance Premiums", "type": "expense", "subtype": "operating"}
    if "motor expense" in key:
        return {"account": "Motor Expenses", "type": "expense", "subtype": "operating"}
    if "lighting" in key or "heating" in key or "utilities" in key:
        return {"account": "Utilities Expense", "type": "expense", "subtype": "operating"}
    if "general expense" in key:
        return {"account": "General Expenses", "type": "expense", "subtype": "operating"}
    if "office expense" in key:
        return {"account": "Office Expenses", "type": "expense", "subtype": "operating"}
    if "bad debt" in key:
        return {"account": "Bad Debts", "type": "expense", "subtype": "other"}
    if "discount" in key:
        return {"account": "Discounts", "type": "expense", "subtype": "operating"}
    if "depreciation" in key:
        return {"account": "Depreciation Expense", "type": "expense", "subtype": "other"}
    if "interest received" in key:
        return {"account": "Interest Received", "type": "revenue", "subtype": "other"}
    if "interest paid" in key or "interest on loan" in key or "interest on borrowing" in key:
        return {"account": "Interest Paid on Loans", "type": "expense", "subtype": "other"}

    if "premises" in key or "land" in key or "building" in key:
        return {"account": "Land and Buildings", "type": "asset", "subtype": "non-current"}
    if "motor vehicle" in key or "vehicle" in key:
        return {"account": "Vehicles", "type": "asset", "subtype": "non-current"}
    if "fixture" in key or "fitting" in key or "equipment" in key:
        return {"account": "Equipment", "type": "asset", "subtype": "non-current"}
    if "inventory" in key:
        return {"account": "Inventory", "type": "asset", "subtype": "current"}
    if "debtor" in key or "receivable" in key:
        return {"account": "Accounts Receivable", "type": "asset", "subtype": "current"}
    if "cash" in key or "bank" in key:
        return {"account": "Cash and Cash Equivalents", "type": "asset", "subtype": "current"}

    if "creditor" in key or "payable" in key:
        return {"account": "Accounts Payable", "type": "liability", "subtype": "current"}
    if "loan" in key or "borrowing" in key:
        return {"account": "Bank Loan", "type": "liability", "subtype": "non-current"}
    if "capital" in key or "equity" in key:
        return {"account": "Owner Capital", "type": "capital", "subtype": "equity"}
    if "drawing" in key:
        return {"account": "Drawings", "type": "drawings", "subtype": "equity"}

    if any(keyword in key for keyword in {"income", "revenue"}):
        return {"account": title, "type": "revenue", "subtype": "other"}
    if any(keyword in key for keyword in {"expense", "cost"}):
        return {"account": title, "type": "expense", "subtype": "operating"}

    return {"account": title, "type": "expense", "subtype": "operating"}


def aggregate_ledger_dataframe(df):
    return (
        df.groupby(["account", "type", "subtype"], as_index=False, dropna=False)[["amount", "depreciation"]]
        .sum()
        .sort_values(["type", "account"], kind="stable")
        .reset_index(drop=True)
    )


def normalize_structured_ledger_dataframe(df):
    rename_map = {}
    debit_column = None
    credit_column = None

    for column in df.columns:
        role = detect_column_role(column)
        if role in {"account", "type", "subtype", "amount", "depreciation"} and role not in rename_map.values():
            rename_map[column] = role
        elif role == "debit" and debit_column is None:
            debit_column = column
        elif role == "credit" and credit_column is None:
            credit_column = column

    normalized = df.rename(columns=rename_map).copy()
    normalized = normalized.dropna(how="all")
    if normalized.empty:
        raise ValueError("file has no readable ledger rows")

    if "amount" not in normalized.columns:
        if debit_column is None and credit_column is None:
            raise ValueError("missing required columns: amount")

        debit_values = (
            pd.to_numeric(normalized[debit_column], errors="coerce").fillna(0).abs()
            if debit_column is not None
            else 0
        )
        credit_values = (
            pd.to_numeric(normalized[credit_column], errors="coerce").fillna(0).abs()
            if credit_column is not None
            else 0
        )
        normalized["amount"] = debit_values + credit_values

    normalized["amount"] = normalized["amount"].apply(parse_numeric_cell)
    if normalized["amount"].isna().any():
        raise ValueError("amount column must contain numeric values")

    if "account" not in normalized.columns:
        if "type" not in normalized.columns:
            raise ValueError("missing required columns: account or type")
        normalized["account"] = normalized["type"].astype(str).str.title()
    else:
        normalized["account"] = normalized["account"].astype(str).str.strip().replace("", pd.NA)
        if normalized["account"].isna().all() and "type" not in normalized.columns:
            raise ValueError("account column must contain ledger names")

    inferred = pd.DataFrame(normalized["account"].fillna("").map(infer_trial_balance_account).tolist())

    if "type" not in normalized.columns:
        normalized["type"] = inferred["type"]
    else:
        normalized["type"] = normalized["type"].astype(str).str.lower().str.strip()
        blank_type = normalized["type"].eq("")
        normalized.loc[blank_type, "type"] = inferred.loc[blank_type, "type"]

    if "subtype" not in normalized.columns:
        normalized["subtype"] = inferred["subtype"]
    else:
        normalized["subtype"] = normalized["subtype"].astype(str).str.lower().str.strip()
        blank_subtype = normalized["subtype"].eq("")
        normalized.loc[blank_subtype, "subtype"] = inferred.loc[blank_subtype, "subtype"]
        still_blank = normalized["subtype"].eq("")
        normalized.loc[still_blank, "subtype"] = normalized.loc[still_blank, "type"].map(default_subtype_for)

    blank_account = normalized["account"].isna() | normalized["account"].astype(str).str.strip().eq("")
    normalized.loc[blank_account, "account"] = inferred.loc[blank_account, "account"]

    if "depreciation" not in normalized.columns:
        normalized["depreciation"] = 0.0
    else:
        normalized["depreciation"] = normalized["depreciation"].apply(parse_numeric_cell)
        if normalized["depreciation"].isna().any():
            raise ValueError("depreciation column must contain numeric values")
        if (normalized["depreciation"] < 0).any():
            raise ValueError("depreciation cannot be negative")

    return normalized[["account", "type", "subtype", "amount", "depreciation"]].reset_index(drop=True)


def normalize_trial_balance_dataframe(df):
    entries = []
    pending_account = None

    for row in df.itertuples(index=False, name=None):
        cells = [value for value in row if not is_blank_cell(value)]
        if not cells:
            continue

        numeric_values = [parse_numeric_cell(value) for value in cells]
        numeric_values = [value for value in numeric_values if value is not None]
        text_values = [str(value).strip() for value in cells if parse_numeric_cell(value) is None and str(value).strip()]

        if text_values and numeric_values:
            if should_skip_derived_label(text_values[0]):
                pending_account = None
                continue
            details = infer_trial_balance_account(text_values[0])
            entries.append(
                {
                    **details,
                    "amount": abs(float(numeric_values[-1])),
                    "depreciation": 0.0,
                }
            )
            pending_account = None
            continue

        if text_values:
            normalized_label = normalize_account_key(text_values[0])
            if should_skip_derived_label(text_values[0]) or normalized_label == "factory overheads":
                pending_account = None
                continue
            pending_account = text_values[0]
            continue

        if numeric_values and pending_account:
            details = infer_trial_balance_account(pending_account)
            entries.append(
                {
                    **details,
                    "amount": abs(float(numeric_values[0])),
                    "depreciation": 0.0,
                }
            )
            pending_account = None

    if not entries:
        raise ValueError("could not detect ledger rows in the uploaded trial balance")

    return aggregate_ledger_dataframe(pd.DataFrame(entries))


def read_external_dataframe(uploaded_file):
    filename = (uploaded_file.filename or "").lower()
    suffix = Path(filename).suffix

    if suffix in {".csv", ".txt"}:
        return read_tabular_dataframe(uploaded_file, pd.read_csv)
    if suffix in {".xls", ".xlsx"}:
        try:
            return read_tabular_dataframe(uploaded_file, pd.read_excel)
        except ImportError as exc:
            raise ValueError("xlsx support requires openpyxl to be installed") from exc
    if suffix == ".json":
        return pd.read_json(uploaded_file)
    if suffix == ".pdf":
        if pdfplumber is None:
            raise ValueError("pdf support is not installed")

        rows = []
        with pdfplumber.open(uploaded_file) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables() or []
                for table in tables:
                    rows.extend(table)

                if not tables:
                    text = page.extract_text() or ""
                    for line in text.splitlines():
                        parts = [part.strip() for part in line.split() if part.strip()]
                        if parts:
                            rows.append(parts)

        if not rows:
            raise ValueError("no readable rows found in pdf")

        header, *data_rows = rows
        return pd.DataFrame(data_rows, columns=header) if data_rows else pd.DataFrame(rows)
    if suffix in {".doc", ".docx"}:
        if docx is None:
            raise ValueError("word support is not installed")

        document = docx.Document(uploaded_file)
        rows = []
        for table in document.tables:
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])

        if rows:
            header, *data_rows = rows
            return pd.DataFrame(data_rows, columns=header) if data_rows else pd.DataFrame(rows)

        text_rows = [line for line in (paragraph.text.strip() for paragraph in document.paragraphs) if line]
        if not text_rows:
            raise ValueError("no readable rows found in word document")
        return pd.read_csv(StringIO("\n".join(text_rows)))

    raise ValueError("unsupported file type; use CSV, XLS, XLSX, TXT, JSON, PDF, or Word")


def normalize_ledger_dataframe(df):
    column_roles = {detect_column_role(column) for column in df.columns}
    column_roles.discard(None)

    if column_roles:
        return normalize_structured_ledger_dataframe(df)
    return normalize_trial_balance_dataframe(df)


def calc(df):
    required = {"type", "amount"}
    missing = required.difference(df.columns)
    if missing:
        raise ValueError(f"missing required columns: {', '.join(sorted(missing))}")

    df = df.copy()
    df["type"] = df["type"].astype(str).str.lower().str.strip()
    df["amount"] = pd.to_numeric(df["amount"], errors="coerce")
    if "subtype" in df.columns:
        df["subtype"] = df["subtype"].astype(str).str.lower().str.strip()
    else:
        df["subtype"] = ""

    if "depreciation" in df.columns:
        df["depreciation"] = pd.to_numeric(df["depreciation"], errors="coerce")
        if df["depreciation"].isna().any():
            raise ValueError("depreciation column must contain numeric values")
        if (df["depreciation"] < 0).any():
            raise ValueError("depreciation cannot be negative")
    else:
        df["depreciation"] = 0.0

    if df["amount"].isna().any():
        raise ValueError("amount column must contain numeric values")

    current_assets = float(
        df.loc[(df["type"] == "asset") & (df["subtype"] != "non-current"), "amount"].sum()
    )
    non_current_assets_gross = float(
        df.loc[(df["type"] == "asset") & (df["subtype"] == "non-current"), "amount"].sum()
    )
    accumulated_depreciation = float(
        df.loc[(df["type"] == "asset") & (df["subtype"] == "non-current"), "depreciation"].sum()
    )
    net_non_current_assets = max(0.0, non_current_assets_gross - accumulated_depreciation)
    total_assets = current_assets + net_non_current_assets

    current_liabilities = float(
        df.loc[(df["type"] == "liability") & (df["subtype"] != "non-current"), "amount"].sum()
    )
    non_current_liabilities = float(
        df.loc[(df["type"] == "liability") & (df["subtype"] == "non-current"), "amount"].sum()
    )
    total_liabilities = current_liabilities + non_current_liabilities

    revenue = float(df.loc[df["type"] == "revenue", "amount"].sum())
    expenses = float(df.loc[df["type"] == "expense", "amount"].sum())

    if "account" in df.columns:
        account_totals = (
            df.assign(account_key=df["account"].astype(str).str.strip().str.lower())
            .groupby("account_key")["amount"]
            .sum()
            .to_dict()
        )
    else:
        account_totals = {}

    def amount_by_account(*names):
        return float(sum(account_totals.get(str(name).strip().lower(), 0) for name in names))

    raw_materials_opening = amount_by_account("Raw Materials Opening Stock", "Opening Raw Materials")
    raw_materials_purchases = amount_by_account("Purchases", "Purchases of Raw Materials")
    raw_materials_carriage = amount_by_account("Carriage Inwards")
    raw_materials_returns = amount_by_account("Returns Outwards")
    raw_materials_closing = amount_by_account("Closing Raw Materials")
    raw_materials_consumed = raw_materials_opening + raw_materials_purchases + raw_materials_carriage - raw_materials_returns - raw_materials_closing
    direct_manufacturing_labor = amount_by_account("Direct Labour", "Direct Manufacturing Labor")
    factory_indirect_labor = amount_by_account("Factory Indirect Labor")
    factory_utilities = amount_by_account("Factory Utilities")
    depreciation_factory_equipment = amount_by_account("Depreciation of Factory Equipment")
    other_factory_overheads = amount_by_account("Factory Expenses", "Factory Overheads")
    total_factory_overheads = (
        factory_indirect_labor
        + factory_utilities
        + depreciation_factory_equipment
        + other_factory_overheads
    )
    prime_cost = raw_materials_consumed + direct_manufacturing_labor
    total_factory_cost = prime_cost + total_factory_overheads
    opening_wip = amount_by_account("Opening Work in Progress")
    closing_wip = amount_by_account("Closing Work in Progress")
    cost_of_goods_manufactured = total_factory_cost + opening_wip - closing_wip

    # Return both snake_case and camelCase keys for easy frontend compatibility.
    return {
        "revenue": revenue,
        "expenses": expenses,
        "assets_current": current_assets,
        "assets_non_current_gross": non_current_assets_gross,
        "accumulated_depreciation": accumulated_depreciation,
        "assets_non_current_net": net_non_current_assets,
        "total_assets": total_assets,
        "liabilities_current": current_liabilities,
        "liabilities_non_current": non_current_liabilities,
        "total_liabilities": total_liabilities,
        "assetsCurrent": current_assets,
        "assetsNonCurrentGross": non_current_assets_gross,
        "nonCurrentAccumulatedDepreciation": accumulated_depreciation,
        "assetsNonCurrent": net_non_current_assets,
        "totalAssets": total_assets,
        "liabilitiesCurrent": current_liabilities,
        "liabilitiesNonCurrent": non_current_liabilities,
        "totalLiabilities": total_liabilities,
        "expense": expenses,
        "raw_materials_opening": raw_materials_opening,
        "raw_materials_purchases": raw_materials_purchases,
        "raw_materials_carriage": raw_materials_carriage,
        "raw_materials_returns": raw_materials_returns,
        "raw_materials_closing": raw_materials_closing,
        "raw_materials_consumed": raw_materials_consumed,
        "direct_manufacturing_labor": direct_manufacturing_labor,
        "factory_indirect_labor": factory_indirect_labor,
        "factory_utilities": factory_utilities,
        "depreciation_factory_equipment": depreciation_factory_equipment,
        "other_factory_overheads": other_factory_overheads,
        "total_factory_overheads": total_factory_overheads,
        "prime_cost": prime_cost,
        "total_factory_cost": total_factory_cost,
        "opening_wip": opening_wip,
        "closing_wip": closing_wip,
        "cost_of_goods_manufactured": cost_of_goods_manufactured,
        "rawMaterialsOpening": raw_materials_opening,
        "rawMaterialsPurchases": raw_materials_purchases,
        "rawMaterialsCarriage": raw_materials_carriage,
        "rawMaterialsReturns": raw_materials_returns,
        "rawMaterialsClosing": raw_materials_closing,
        "rawMaterialsConsumed": raw_materials_consumed,
        "directManufacturingLabor": direct_manufacturing_labor,
        "factoryIndirectLabor": factory_indirect_labor,
        "factoryUtilities": factory_utilities,
        "depreciationFactoryEquipment": depreciation_factory_equipment,
        "otherFactoryOverheads": other_factory_overheads,
        "totalFactoryOverheads": total_factory_overheads,
        "primeCost": prime_cost,
        "totalFactoryCost": total_factory_cost,
        "openingWip": opening_wip,
        "closingWip": closing_wip,
        "costOfGoodsManufactured": cost_of_goods_manufactured,
    }


# ---------------- AUTH ----------------

@app.route("/register", methods=["POST"])
@limiter.limit("10 per minute")
def register():
    data = request.get_json(silent=True) or {}
    org_name = (data.get("org") or "").strip()
    email = (data.get("email") or "").strip().lower()
    password = data.get("password") or ""
    raw_business_type = data.get("business_type")
    business_type = normalize_business_type(raw_business_type)

    try:
        partner_names = normalize_partner_names(data.get("partner_names"))
    except ValueError as exc:
        return error_response(str(exc))

    if not org_name or not email or not password:
        return error_response("org, email, and password are required")
    if len(password) < 8:
        return error_response("password must be at least 8 characters")
    validation_error = validate_company_setup(business_type, partner_names)
    if validation_error:
        return error_response(validation_error)

    if User.query.filter_by(email=email).first():
        return error_response("email already exists", 409)

    hashed_password = bcrypt.generate_password_hash(password).decode()
    try:
        org = Organization(name=org_name, billing_email=email)
        db.session.add(org)
        db.session.flush()
        apply_subscription_plan(org, "free", status="free")
        company = Company(org_id=org.id, name=org_name, business_type=business_type)
        db.session.add(company)
        db.session.flush()
        replace_company_partners(company, partner_names)
        if raw_business_type not in {None, ""}:
            set_company_onboarding_state(company.id, True)
        seed_chart_of_accounts(company)
        seed_integration_connections(company)
        user = User(email=email, password=hashed_password, role="owner", org_id=org.id, default_company_id=company.id)
        db.session.add(user)
        db.session.flush()
        assign_company_membership(user, company, role="owner", is_default=True)
        if not safe_commit():
            return error_response("database error during registration", 503)
    except IntegrityError:
        db.session.rollback()
        return error_response("email already exists", 409)
    except SQLAlchemyError:
        db.session.rollback()
        return error_response("database error during registration", 503)

    return {"msg": "registered", "company": serialize_company(company)}


@app.route("/password-reset/request", methods=["POST"])
@limiter.limit("5 per hour")
def request_password_reset():
    data = request.get_json(silent=True) or {}
    email = (data.get("email") or "").strip().lower()
    if not email:
        return error_response("email is required")

    response = {"msg": "If an account exists for that email, reset instructions have been prepared."}
    user = User.query.filter_by(email=email).first()
    if not user:
        return response

    raw_token = issue_password_reset_token(user)
    if not safe_commit():
        return error_response("database error while preparing password reset", 503)

    reset_link = build_password_reset_link(raw_token)
    email_sent = send_password_reset_email(user, reset_link)
    log(user.id, "requested password reset")

    if email_sent:
        response["delivery"] = "email"
    elif should_return_password_reset_preview():
        response["delivery"] = "preview"
        response["reset_token"] = raw_token
        response["reset_link"] = reset_link

    return response


@app.route("/password-reset/confirm", methods=["POST"])
@limiter.limit("10 per hour")
def confirm_password_reset():
    data = request.get_json(silent=True) or {}
    raw_token = (data.get("token") or "").strip()
    password = data.get("password") or ""

    if not raw_token or not password:
        return error_response("token and password are required")
    if len(password) < 8:
        return error_response("password must be at least 8 characters")

    record = get_valid_password_reset_record(raw_token)
    if not record:
        return error_response("invalid or expired reset token", 400)

    user = db.session.get(User, record.user_id)
    if not user:
        return error_response("invalid or expired reset token", 400)

    now = datetime.datetime.now(datetime.UTC)
    user.password = bcrypt.generate_password_hash(password).decode()
    record.used_at = now
    other_tokens = PasswordResetToken.query.filter(
        PasswordResetToken.user_id == user.id,
        PasswordResetToken.used_at.is_(None),
        PasswordResetToken.id != record.id,
    ).all()
    for token in other_tokens:
        token.used_at = now
    ActiveSession.query.filter_by(user_id=user.id).delete()
    if not safe_commit():
        return error_response("database error while resetting password", 503)

    log(user.id, "reset password")
    return {"msg": "Password reset complete. Sign in with your new password."}


@app.route("/login", methods=["POST"])
@limiter.limit("20 per minute")
def login():
    data = request.get_json(silent=True) or {}
    email = (data.get("email") or "").strip().lower()
    password = data.get("password") or ""

    if not email or not password:
        return error_response("email and password are required")

    user = User.query.filter_by(email=email).first()

    if not user or not bcrypt.check_password_hash(user.password, password):
        return error_response("invalid email or password", 401)

    touch_session(user.id)
    log(user.id, "logged in")
    return {"token": build_access_token(user)}


@app.route("/refresh", methods=["POST"])
@jwt_required()
def refresh():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    touch_session(user.id)
    return {"token": build_access_token(user)}


# ---------------- INVITE USER ----------------

@app.route("/invite", methods=["POST"])
@jwt_required()
@limiter.limit("30 per minute")
def invite():
    me = get_user_from_token()
    if not me:
        return error_response("invalid token", 401)

    if me.role not in ["owner", "admin"]:
        return error_response("not allowed", 403)

    data = request.get_json(silent=True) or {}
    invite_email = (data.get("email") or "").strip().lower()
    if not invite_email:
        return error_response("email is required")

    if User.query.filter_by(email=invite_email).first():
        return error_response("email already exists", 409)

    hashed_password = bcrypt.generate_password_hash("temp123").decode()

    role = (data.get("role") or "member").strip().lower()
    if role not in VALID_ROLES:
        return error_response("invalid role")

    try:
        membership_specs = normalize_membership_specs(
            data.get("memberships"),
            data.get("company_ids"),
            parse_company_id(data.get("company_id")) or me.default_company_id or resolve_company_for_user(me).id,
            role,
        )
    except ValueError as exc:
        return error_response(str(exc))

    company_map = membership_company_map([spec["company_id"] for spec in membership_specs])
    if len(company_map) != len(membership_specs) or any(company.org_id != me.org_id for company in company_map.values()):
        return error_response("company not found", 404)

    default_company_id = next((spec["company_id"] for spec in membership_specs if spec["is_default"]), membership_specs[0]["company_id"])
    user = User(email=invite_email, password=hashed_password, role=role, org_id=me.org_id, default_company_id=default_company_id)
    db.session.add(user)
    db.session.flush()
    apply_user_company_memberships(user, membership_specs)
    if not safe_commit():
        return error_response("database error while inviting user", 503)

    log(me.id, "invited user")
    return {"msg": "user added"}


# ---------------- REPORT ENGINE ----------------

@app.route("/analyze", methods=["POST"])
@jwt_required()
@limiter.limit("30 per minute")
def analyze():
    status = maintenance_state()
    if status["maintenance"]:
        return error_response(status["message"], 503)

    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    org = db.session.get(Organization, user.org_id)
    if not org:
        return error_response("organization not found", 404)

    if org.usage >= FREE_USAGE_LIMIT:
        return error_response("limit reached", 403)

    uploaded_file = request.files.get("file")
    if not uploaded_file:
        return error_response("file is required")

    company = resolve_company_for_user(user, request.form.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    try:
        df = read_external_dataframe(uploaded_file)
        lowered_columns = {str(column).strip().lower() for column in df.columns}
        if "type" in lowered_columns and "amount" in lowered_columns:
            df = normalize_ledger_dataframe(df)
        result = calc(df)
    except Exception as exc:
        suffix = Path((uploaded_file.filename or "").lower()).suffix
        label = "csv" if suffix in {".csv", ".txt"} else "file"
        return error_response(f"invalid {label}: {exc}")

    org.usage += 1
    db.session.add(Report(org_id=org.id, company_id=company.id, data=json.dumps(result)))
    if not safe_commit():
        return error_response("database error while saving report", 503)

    log(user.id, f"generated report for {company.name}")
    return result


@app.route("/extract-ledger", methods=["POST"])
@jwt_required()
@limiter.limit("30 per minute")
def extract_ledger():
    status = maintenance_state()
    if status["maintenance"]:
        return error_response(status["message"], 503)

    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    uploaded_file = request.files.get("file")
    if not uploaded_file:
        return error_response("file is required")

    try:
        raw_df = read_external_dataframe(uploaded_file)
        normalized_df = normalize_ledger_dataframe(raw_df)
        summary = calc(normalized_df)
    except Exception as exc:
        suffix = Path((uploaded_file.filename or "").lower()).suffix
        label = "csv" if suffix in {".csv", ".txt"} else "file"
        return error_response(f"invalid {label}: {exc}")

    ledger_rows = []
    for idx, row in enumerate(normalized_df.to_dict(orient="records"), start=1):
        ledger_rows.append(
            {
                "id": idx,
                "account": str(row.get("account", "")).strip(),
                "type": str(row.get("type", "")).strip(),
                "subtype": str(row.get("subtype", "")).strip(),
                "amount": float(row.get("amount", 0) or 0),
                "depreciation": float(row.get("depreciation", 0) or 0),
            }
        )

    log(user.id, "extracted external ledger file")
    return {"ledger_rows": ledger_rows, "summary": summary}


@app.route("/companies")
@jwt_required()
def list_companies():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    company_pairs = accessible_companies_for_user(user)
    if not company_pairs:
        company = get_or_create_default_company(user.org_id)
        assign_company_membership(user, company, role=user.role, is_default=True)
        safe_commit()
        company_pairs = accessible_companies_for_user(user)

    return jsonify([serialize_company(company, membership=membership) for company, membership in company_pairs])


@app.route("/companies", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin")
def create_company():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    org = db.session.get(Organization, user.org_id)
    if not org:
        return error_response("organization not found", 404)

    data = request.get_json(silent=True) or {}
    name = (data.get("name") or "").strip()
    business_type = normalize_business_type(data.get("business_type"))

    try:
        partner_names = normalize_partner_names(data.get("partner_names"))
    except ValueError as exc:
        return error_response(str(exc))

    if not name:
        return error_response("company name is required")
    validation_error = validate_company_setup(business_type, partner_names)
    if validation_error:
        return error_response(validation_error)
    company_limit = int(org.max_companies or get_plan_definition(org.plan_code)["max_companies"])
    existing_company_count = Company.query.filter_by(org_id=user.org_id).count()
    if existing_company_count >= company_limit:
        return error_response("current plan limit reached; upgrade billing to add more companies", 403)

    company = Company(org_id=user.org_id, name=name, business_type=business_type)
    db.session.add(company)
    db.session.flush()
    replace_company_partners(company, partner_names)
    set_company_onboarding_state(company.id, True)
    seed_chart_of_accounts(company)
    seed_integration_connections(company)
    assign_company_membership(user, company, role=user.role, is_default=False)
    admin_users = User.query.filter(User.org_id == user.org_id, User.role.in_(["owner", "admin"])).all()
    for org_user in admin_users:
        assign_company_membership(org_user, company, role=org_user.role, is_default=False)
    if not safe_commit():
        return error_response("database error while creating company", 503)

    log(user.id, f"created company {name}", company_id=company.id)
    membership = company_membership_for_user(user, company.id)
    return serialize_company(company, membership=membership), 201


@app.route("/companies/<int:company_id>/setup", methods=["PUT"])
@jwt_required()
@roles_required("owner", "admin")
def update_company_setup(company_id):
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    company = resolve_company_for_user(user, company_id)
    if not company:
        return error_response("company not found", 404)

    data = request.get_json(silent=True) or {}
    business_type = normalize_business_type(data.get("business_type"))

    try:
        partner_names = normalize_partner_names(data.get("partner_names"))
    except ValueError as exc:
        return error_response(str(exc))

    validation_error = validate_company_setup(business_type, partner_names)
    if validation_error:
        return error_response(validation_error)

    company.business_type = business_type
    replace_company_partners(company, partner_names)
    set_company_onboarding_state(company.id, True)
    seed_chart_of_accounts(company)
    seed_integration_connections(company)
    if not safe_commit():
        return error_response("database error while saving company setup", 503)

    log(user.id, f"configured company {company.name} as {business_type}", company_id=company.id)
    membership = company_membership_for_user(user, company.id)
    return serialize_company(company, membership=membership)


def normalized_trial_balance_amount(item):
    amount = float(item.get("net_balance", 0) or 0)
    return round(-amount if item.get("normal_balance") == "credit" else amount, 2)


def build_company_ai_snapshot(company):
    finance = calculate_finance_summary(company)
    tax = calculate_tax_summary(company)
    accounting = build_accounting_overview(company)
    workforce = build_workforce_overview(company)
    inventory = build_inventory_summary(company)
    projects = build_project_summary(company)
    reports = aggregate_org_reports(company.org_id, company.id)

    trial_items = accounting["trial_balance"]["items"]
    current_assets = round(
        sum(
            max(0.0, normalized_trial_balance_amount(item))
            for item in trial_items
            if item["category"] == "asset" and item["subtype"] == "current"
        ),
        2,
    )
    current_liabilities = round(
        sum(
            max(0.0, normalized_trial_balance_amount(item))
            for item in trial_items
            if item["category"] == "liability" and item["subtype"] == "current"
        ),
        2,
    )
    cash_balance = round(
        sum(max(0.0, normalized_trial_balance_amount(item)) for item in trial_items if item["code"] == "1000"),
        2,
    )
    annual_revenue = round(float(reports.get("revenue", 0) or 0), 2)
    annual_expenses = round(float(reports.get("expenses", 0) or 0), 2)
    monthly_inflow = round(max(float(finance["collected_this_month"]), annual_revenue / 12 if annual_revenue else 0), 2)
    baseline_outflow = float(finance["paid_this_month"]) + float(workforce["payroll_this_month"])
    fallback_outflow = annual_expenses / 12 if annual_expenses else (float(finance["open_payables"]) / 3 if finance["open_payables"] else 0)
    monthly_outflow = round(max(baseline_outflow, fallback_outflow, 1.0), 2)
    current_ratio = round(current_assets / current_liabilities, 2) if current_liabilities > 0 else None
    cash_runway_months = round(cash_balance / monthly_outflow, 1) if monthly_outflow > 0 else None

    tax_drag = round(max(float(tax["net_tax_due"]), 0.0) / 3, 2) if float(tax["net_tax_due"]) > 0 else 0.0
    projected_cash = cash_balance
    forecast = []
    for month_number in range(1, 4):
        projected_cash = round(projected_cash + monthly_inflow - monthly_outflow - tax_drag, 2)
        forecast.append(
            {
                "month": month_number,
                "label": f"{month_number * 30} days",
                "projected_cash": projected_cash,
            }
        )

    return {
        "company_id": company.id,
        "company_name": company.name,
        "metrics": {
            "cash_balance": cash_balance,
            "current_assets": current_assets,
            "current_liabilities": current_liabilities,
            "current_ratio": current_ratio,
            "annual_revenue": annual_revenue,
            "annual_expenses": annual_expenses,
            "monthly_inflow": monthly_inflow,
            "monthly_outflow": monthly_outflow,
            "cash_runway_months": cash_runway_months,
        },
        "finance": finance,
        "tax": tax,
        "workforce": workforce,
        "inventory": {
            "inventory_value": inventory["inventory_value"],
            "low_stock_count": inventory["low_stock_count"],
            "open_purchase_orders": inventory["open_purchase_orders"],
        },
        "projects": {
            "total_revenue": projects["total_revenue"],
            "total_cost": projects["total_cost"],
            "total_margin": projects["total_margin"],
        },
        "forecast": forecast,
    }


def build_ai_cfo_alerts(snapshot):
    finance = snapshot["finance"]
    tax = snapshot["tax"]
    inventory = snapshot["inventory"]
    workforce = snapshot["workforce"]
    metrics = snapshot["metrics"]
    alerts = []

    if metrics["current_ratio"] is not None and metrics["current_ratio"] < 1:
        alerts.append(
            {
                "severity": "high",
                "title": "Liquidity risk",
                "message": f"Current ratio is {metrics['current_ratio']}, so short-term obligations are outpacing liquid assets.",
                "recommendation": "Accelerate collections, delay discretionary spend, and review payable timing this week.",
            }
        )
    if float(finance["overdue_receivables"]) > 0:
        alerts.append(
            {
                "severity": "medium",
                "title": "Collections pressure",
                "message": f"There are {finance['overdue_invoice_count']} overdue invoices totaling {round(float(finance['overdue_receivables']), 2)}.",
                "recommendation": "Trigger a collections sequence and escalate the oldest invoices to the owner or finance lead.",
            }
        )
    if float(finance["overdue_payables"]) > 0:
        alerts.append(
            {
                "severity": "medium",
                "title": "Supplier obligations overdue",
                "message": f"Overdue payables total {round(float(finance['overdue_payables']), 2)} across {finance['overdue_bill_count']} bills.",
                "recommendation": "Prioritize critical vendors and schedule payment rails before terms deteriorate.",
            }
        )
    if int(finance["bank_unmatched_count"]) > 0:
        alerts.append(
            {
                "severity": "medium",
                "title": "Bank close backlog",
                "message": f"{finance['bank_unmatched_count']} bank transactions still need reconciliation.",
                "recommendation": "Run the reconciliation workspace and clear unmatched items before month-end reporting.",
            }
        )
    if int(inventory["low_stock_count"]) > 0:
        alerts.append(
            {
                "severity": "medium",
                "title": "Reorder risk",
                "message": f"{inventory['low_stock_count']} inventory items are at or below reorder point.",
                "recommendation": "Release purchase orders for critical SKUs and review safety stock on the fastest movers.",
            }
        )
    if float(tax["net_tax_due"]) > 0:
        alerts.append(
            {
                "severity": "medium",
                "title": "Tax liability due",
                "message": f"Estimated net tax due is {round(float(tax['net_tax_due']), 2)} in {tax['jurisdiction_code']}.",
                "recommendation": "Reserve cash now and prepare the filing pack before the next statutory deadline.",
            }
        )
    if snapshot["forecast"] and snapshot["forecast"][-1]["projected_cash"] < 0:
        alerts.append(
            {
                "severity": "high",
                "title": "Projected cash shortfall",
                "message": f"Cash is forecast to fall to {snapshot['forecast'][-1]['projected_cash']} within 90 days.",
                "recommendation": "Cut non-essential outflow, pull collections forward, or line up short-term financing immediately.",
            }
        )
    if float(workforce["contractor_1099_exposure"]) >= 600:
        alerts.append(
            {
                "severity": "low",
                "title": "1099 review",
                "message": f"Contractor exposure this period is {round(float(workforce['contractor_1099_exposure']), 2)}.",
                "recommendation": "Confirm contractor tax details now so year-end compliance does not become a scramble.",
            }
        )

    return alerts


def build_ai_cfo_overview(company):
    snapshot = build_company_ai_snapshot(company)
    alerts = build_ai_cfo_alerts(snapshot)
    top_actions = [alert["recommendation"] for alert in alerts[:4]]
    metrics = snapshot["metrics"]
    narrative = (
        f"{company.name} is carrying {metrics['cash_balance']} in cash with "
        f"{snapshot['finance']['open_receivables']} outstanding receivables and "
        f"{snapshot['finance']['open_payables']} open payables. "
        f"Projected 90-day cash is {snapshot['forecast'][-1]['projected_cash'] if snapshot['forecast'] else metrics['cash_balance']}."
    )
    return {
        **snapshot,
        "alerts": alerts,
        "top_actions": top_actions,
        "narrative": narrative,
    }


def answer_ai_cfo_question(question, overview):
    lowered = str(question or "").strip().lower()
    metrics = overview["metrics"]
    finance = overview["finance"]
    tax = overview["tax"]
    inventory = overview["inventory"]
    workforce = overview["workforce"]

    if "profit" in lowered or "margin" in lowered:
        return (
            f"Recorded annual revenue is {metrics['annual_revenue']} against expenses of {metrics['annual_expenses']}. "
            f"Project margin currently stands at {overview['projects']['total_margin']}."
        )
    if "cash" in lowered or "runway" in lowered or "liquidity" in lowered:
        return (
            f"Cash balance is {metrics['cash_balance']} with a monthly outflow run rate of {metrics['monthly_outflow']}. "
            f"Estimated cash runway is {metrics['cash_runway_months']} months."
        )
    if "tax" in lowered or "vat" in lowered or "filing" in lowered:
        return (
            f"Estimated net tax due is {tax['net_tax_due']} under {tax['jurisdiction_code']}. "
            f"Prepare the next {tax['filing_frequency']} filing before cash is committed elsewhere."
        )
    if "inventory" in lowered or "stock" in lowered or "purchase order" in lowered:
        return (
            f"{inventory['low_stock_count']} items are at reorder risk and there are "
            f"{inventory['open_purchase_orders']} open purchase orders."
        )
    if "payroll" in lowered or "staff" in lowered or "contractor" in lowered:
        return (
            f"Payroll cash this month is {workforce['payroll_this_month']} and contractor 1099 exposure is "
            f"{workforce['contractor_1099_exposure']}."
        )
    if overview["top_actions"]:
        return f"{overview['narrative']} Top action: {overview['top_actions'][0]}"
    return overview["narrative"]


def serialize_background_job(job):
    payload = {}
    result = {}
    try:
        payload = json.loads(job.payload_json or "{}")
    except (TypeError, ValueError):
        payload = {}
    try:
        result = json.loads(job.result_json or "{}")
    except (TypeError, ValueError):
        result = {}
    return {
        "id": job.id,
        "org_id": job.org_id,
        "company_id": job.company_id,
        "requested_by": job.requested_by,
        "job_type": job.job_type,
        "status": job.status,
        "provider": job.provider,
        "payload": payload,
        "result": result,
        "error_message": job.error_message or "",
        "started_at": job.started_at.isoformat() if job.started_at else None,
        "completed_at": job.completed_at.isoformat() if job.completed_at else None,
        "created_at": job.created_at.isoformat() if job.created_at else None,
    }


def run_background_job(job_id):
    job = db.session.get(BackgroundJob, job_id)
    if not job:
        return None
    if job.status in JOB_TERMINAL_STATUSES:
        return serialize_background_job(job)

    job.status = "running"
    job.started_at = datetime.datetime.now(datetime.UTC)
    safe_commit()

    try:
        company = Company.query.filter_by(id=job.company_id, org_id=job.org_id).first()
        if not company:
            raise ValueError("company not found")
        payload = {}
        try:
            payload = json.loads(job.payload_json or "{}")
        except (TypeError, ValueError):
            payload = {}

        if job.job_type == "finance_digest":
            result = build_ai_cfo_overview(company)
        elif job.job_type == "tax_filing_package":
            profile = get_or_create_tax_profile(company)
            result = build_tax_filing_package(
                company,
                profile,
                period_start=parse_iso_date(payload.get("period_start"), "period_start", default=None),
                period_end=parse_iso_date(payload.get("period_end"), "period_end", default=None),
            )
        elif job.job_type == "accountant_brief":
            result = build_accountant_toolkit(company)
        else:
            raise ValueError("unsupported job type")

        job.status = "completed"
        job.result_json = json.dumps(result)
        job.error_message = None
    except Exception as exc:
        job.status = "failed"
        job.error_message = str(exc)
    finally:
        job.completed_at = datetime.datetime.now(datetime.UTC)
        safe_commit()

    return serialize_background_job(job)


if celery_app:
    @celery_app.task(name="financial_dashboard.process_background_job")
    def process_background_job(job_id):
        return run_background_job(job_id)
else:
    def process_background_job(job_id):  # pragma: no cover - fallback shim
        return run_background_job(job_id)


def plan_code_from_stripe_object(obj):
    metadata = obj.get("metadata") or {}
    metadata_plan = (metadata.get("plan_code") or "").strip().lower()
    if metadata_plan in PLAN_DEFINITIONS:
        return metadata_plan

    items = obj.get("items") or {}
    for item in items.get("data", []) if isinstance(items, dict) else []:
        price = (item or {}).get("price") or {}
        price_id = price.get("id") if isinstance(price, dict) else str(price or "")
        for plan_code, configured_price_id in STRIPE_PRICE_IDS.items():
            if configured_price_id and configured_price_id == price_id:
                return plan_code
    return None


def resolve_org_from_billing_object(obj):
    metadata = obj.get("metadata") or {}
    raw_org_id = metadata.get("org_id")
    try:
        org_id = int(raw_org_id)
    except (TypeError, ValueError):
        org_id = None

    if org_id:
        org = db.session.get(Organization, org_id)
        if org:
            return org

    customer_id = (obj.get("customer") or "").strip()
    subscription_id = (obj.get("subscription") or obj.get("id") or "").strip()
    if subscription_id:
        org = Organization.query.filter_by(stripe_subscription_id=subscription_id).first()
        if org:
            return org
    if customer_id:
        return Organization.query.filter_by(stripe_customer_id=customer_id).first()
    return None


def apply_billing_event(event):
    event_type = (event.get("type") or "").strip()
    obj = (event.get("data") or {}).get("object") or {}
    org = resolve_org_from_billing_object(obj)
    if not org:
        return False

    plan_code = plan_code_from_stripe_object(obj) or org.plan_code or "free"
    customer_id = (obj.get("customer") or "").strip() or None

    if event_type == "checkout.session.completed":
        apply_subscription_plan(
            org,
            plan_code,
            status="active",
            stripe_customer_id=customer_id,
            stripe_subscription_id=(obj.get("subscription") or "").strip() or None,
        )
    elif event_type in {"customer.subscription.created", "customer.subscription.updated"}:
        apply_subscription_plan(
            org,
            plan_code,
            status=(obj.get("status") or "active").strip().lower(),
            stripe_customer_id=customer_id,
            stripe_subscription_id=(obj.get("id") or "").strip() or None,
        )
    elif event_type == "customer.subscription.deleted":
        apply_subscription_plan(org, "free", status="cancelled", stripe_customer_id=customer_id, stripe_subscription_id=None)
    elif event_type == "invoice.paid":
        org.subscription_status = "active"
        org.subscription_updated_at = datetime.datetime.now(datetime.UTC)
        if customer_id:
            org.stripe_customer_id = customer_id
    elif event_type == "invoice.payment_failed":
        org.subscription_status = "past_due"
        org.subscription_updated_at = datetime.datetime.now(datetime.UTC)
    else:
        return False

    safe_commit()
    return True


@app.route("/billing/plans")
@jwt_required()
def billing_plans():
    return {"items": list(PLAN_DEFINITIONS.values())}


@app.route("/billing/summary")
@jwt_required()
def billing_summary():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    org = db.session.get(Organization, user.org_id)
    if not org:
        return error_response("organization not found", 404)

    return {
        **serialize_subscription(org),
        "company_count": Company.query.filter_by(org_id=user.org_id).count(),
        "accessible_company_count": len(ensure_user_company_memberships(user)),
        "usage": int(org.usage or 0),
        "mpesa_configured": mpesa_is_configured(),
        "local_currency": "KES",
    }


@app.route("/billing/checkout-session", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin")
@limiter.limit("10 per minute")
def create_billing_checkout_session():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)
    if not stripe or not STRIPE_SECRET_KEY:
        return error_response("stripe is not configured", 503)

    data = request.get_json(silent=True) or {}
    plan_code = (data.get("plan_code") or "").strip().lower()
    if plan_code not in {"pro", "ai"}:
        return error_response("plan_code must be pro or ai")

    price_id = STRIPE_PRICE_IDS.get(plan_code)
    if not price_id:
        return error_response("selected Stripe price is not configured", 503)

    org = db.session.get(Organization, user.org_id)
    if not org:
        return error_response("organization not found", 404)

    success_url = (data.get("success_url") or f"{frontend_base_url()}/?billing=success").strip()
    cancel_url = (data.get("cancel_url") or f"{frontend_base_url()}/?billing=cancelled").strip()
    customer_id = org.stripe_customer_id
    if not customer_id:
        customer = stripe.Customer.create(
            email=org.billing_email or user.email,
            name=org.name,
            metadata={"org_id": str(org.id)},
        )
        customer_id = customer.get("id")
        org.stripe_customer_id = customer_id
        safe_commit()

    session = stripe.checkout.Session.create(
        mode="subscription",
        customer=customer_id,
        line_items=[{"price": price_id, "quantity": 1}],
        success_url=success_url,
        cancel_url=cancel_url,
        metadata={"org_id": str(org.id), "plan_code": plan_code},
    )
    return {"checkout_url": session.get("url"), "session_id": session.get("id")}


@app.route("/billing/webhook", methods=["POST"])
def stripe_billing_webhook():
    if stripe and STRIPE_WEBHOOK_SECRET:
        payload = request.get_data()
        signature = request.headers.get("Stripe-Signature", "")
        try:
            event = stripe.Webhook.construct_event(payload, signature, STRIPE_WEBHOOK_SECRET)
        except Exception:
            return error_response("invalid webhook signature", 400)
    else:
        event = request.get_json(silent=True) or {}

    handled = apply_billing_event(event)
    return {"received": True, "handled": handled}


@app.route("/billing/mpesa/checkout", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin")
@limiter.limit("10 per minute")
def create_mpesa_checkout():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    data = request.get_json(silent=True) or {}
    plan_code = (data.get("plan_code") or "").strip().lower()
    if plan_code not in {"pro", "ai"}:
        return error_response("plan_code must be pro or ai")

    try:
        phone_number = normalize_phone_number(data.get("phone_number"))
    except ValueError as exc:
        return error_response(str(exc))

    org = db.session.get(Organization, user.org_id)
    company = resolve_company_for_user(user, data.get("company_id"))
    if not org:
        return error_response("organization not found", 404)
    if not company:
        return error_response("company not found", 404)

    try:
        response = initiate_mpesa_stk_push(plan_code, phone_number)
    except urllib.error.HTTPError as exc:
        return error_response(f"M-Pesa request failed ({exc.code})", 502)
    except urllib.error.URLError:
        return error_response("unable to reach M-Pesa right now", 502)
    except Exception as exc:
        return error_response(str(exc), 400)

    payment_request = BillingPaymentRequest(
        org_id=org.id,
        company_id=company.id,
        requested_by=user.id,
        provider="mpesa",
        plan_code=plan_code,
        currency_code="KES",
        amount=float(response["amount"]),
        phone_number=phone_number,
        external_reference=response.get("reference"),
        merchant_request_id=response.get("merchant_request_id"),
        checkout_request_id=response.get("checkout_request_id"),
        status=response.get("status", "pending"),
        provider_response_json=json.dumps(response.get("raw") or {}),
    )
    db.session.add(payment_request)
    if not safe_commit():
        return error_response("database error while creating M-Pesa checkout", 503)

    log(user.id, f"initiated mpesa checkout for {plan_code}", company_id=company.id)
    return {
        **serialize_billing_payment_request(payment_request),
        "customer_message": response.get("customer_message", ""),
        "preview_mode": response.get("status") == "preview",
    }, 201


@app.route("/billing/mpesa/requests/<int:payment_request_id>")
@jwt_required()
def get_mpesa_payment_request(payment_request_id):
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    payment_request = BillingPaymentRequest.query.filter_by(id=payment_request_id, org_id=user.org_id).first()
    if not payment_request:
        return error_response("payment request not found", 404)
    if not company_membership_for_user(user, payment_request.company_id):
        return error_response("not allowed", 403)
    return serialize_billing_payment_request(payment_request)


@app.route("/billing/mpesa/callback", methods=["POST"])
def mpesa_callback():
    payload = request.get_json(silent=True) or {}
    stk_payload = (((payload.get("Body") or {}).get("stkCallback")) or {})
    checkout_request_id = (stk_payload.get("CheckoutRequestID") or "").strip()
    payment_request = BillingPaymentRequest.query.filter_by(checkout_request_id=checkout_request_id).first()
    if not payment_request:
        return {"received": True, "matched": False}

    payment_request.callback_payload_json = json.dumps(payload)
    result_code = str(stk_payload.get("ResultCode", ""))
    if result_code in {"0", "00"}:
        payment_request.status = "paid"
        org = db.session.get(Organization, payment_request.org_id)
        if org:
            apply_subscription_plan(org, payment_request.plan_code, status="active")
    else:
        payment_request.status = "failed"
    if not safe_commit():
        return error_response("database error while processing callback", 503)
    return {"received": True, "matched": True}


@app.route("/ops/jobs")
@jwt_required()
def list_background_jobs():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    query = BackgroundJob.query.filter_by(org_id=user.org_id)
    company = resolve_company_for_user(user, request.args.get("company_id"))
    if request.args.get("company_id"):
        if not company:
            return error_response("company not found", 404)
        query = query.filter_by(company_id=company.id)

    jobs = query.order_by(BackgroundJob.created_at.desc(), BackgroundJob.id.desc()).limit(25).all()
    visible_company_ids = {membership.company_id for membership in ensure_user_company_memberships(user)}
    return {"items": [serialize_background_job(job) for job in jobs if job.company_id in visible_company_ids]}


@app.route("/ops/jobs", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
@plan_required("pro")
@limiter.limit("30 per minute")
def create_background_job_route():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    data = request.get_json(silent=True) or {}
    job_type = (data.get("job_type") or "").strip().lower()
    if job_type not in JOB_TYPES:
        return error_response("unsupported job type")

    company = resolve_company_for_user(user, data.get("company_id"))
    if not company:
        return error_response("company not found", 404)
    if not has_company_permission(user, company, "jobs:run"):
        return error_response("not allowed", 403)

    job = BackgroundJob(
        org_id=company.org_id,
        company_id=company.id,
        requested_by=user.id,
        job_type=job_type,
        payload_json=json.dumps(
            {
                "period_start": data.get("period_start"),
                "period_end": data.get("period_end"),
            }
        ),
        provider="celery" if ASYNC_JOB_MODE == "celery" and celery_app else "inline",
    )
    db.session.add(job)
    if not safe_commit():
        return error_response("database error while creating background job", 503)

    if job.provider == "celery" and celery_app:
        process_background_job.delay(job.id)
        return serialize_background_job(job), 202

    result = run_background_job(job.id)
    return result, 201


@app.route("/ops/jobs/<int:job_id>")
@jwt_required()
def get_background_job_route(job_id):
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    job = BackgroundJob.query.filter_by(id=job_id, org_id=user.org_id).first()
    if not job:
        return error_response("job not found", 404)
    if not company_membership_for_user(user, job.company_id):
        return error_response("not allowed", 403)
    return serialize_background_job(job)


@app.route("/ai-cfo/overview")
@jwt_required()
def ai_cfo_overview():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)
    if not has_company_permission(user, company, "company:view"):
        return error_response("not allowed", 403)

    return build_ai_cfo_overview(company)


@app.route("/ai-cfo/ask", methods=["POST"])
@jwt_required()
@limiter.limit("30 per minute")
def ai_cfo_ask():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    data = request.get_json(silent=True) or {}
    question = (data.get("question") or "").strip()
    if not question:
        return error_response("question is required")

    company = resolve_company_for_user(user, data.get("company_id"))
    if not company:
        return error_response("company not found", 404)
    if not has_company_permission(user, company, "ai:ask"):
        return error_response("not allowed", 403)

    org = db.session.get(Organization, user.org_id)
    if not org or not bool(org.ai_assistant_enabled):
        return error_response("AI CFO chat is available on the AI plan", 403)

    overview = build_ai_cfo_overview(company)
    return {
        "answer": answer_ai_cfo_question(question, overview),
        "narrative": overview["narrative"],
        "top_actions": overview["top_actions"],
    }


# ---------------- FINANCE OPS ----------------

@app.route("/finance/summary")
@jwt_required()
def finance_summary():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    summary = calculate_finance_summary(company)
    summary["company_id"] = company.id
    return summary


@app.route("/finance/invoices")
@jwt_required()
def list_invoices():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    invoices, _ = refresh_finance_documents(company.id)
    scoped = [invoice for invoice in invoices if invoice.org_id == user.org_id]
    scoped.sort(key=lambda invoice: (invoice.issue_date or today_utc_date(), invoice.id), reverse=True)
    return {"items": [serialize_invoice(invoice) for invoice in scoped]}


@app.route("/finance/invoices", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant", "cashier")
def create_invoice():
    user = get_user_from_token()
    data = request.get_json(silent=True) or {}
    company = resolve_company_for_user(user, data.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    customer_name = (data.get("customer_name") or "").strip()
    customer_email = (data.get("customer_email") or "").strip().lower() or None
    if not customer_name:
        return error_response("customer_name is required")

    try:
        issue_date = parse_iso_date(data.get("issue_date"), "issue_date", today_utc_date())
        due_date = parse_iso_date(data.get("due_date"), "due_date", issue_date + datetime.timedelta(days=14))
        tax_rate = parse_money(data.get("tax_rate", 0), "tax_rate")
        items, subtotal = normalize_document_items(data.get("items"), "invoice")
    except ValueError as exc:
        return error_response(str(exc))

    requested_status = (data.get("status") or "draft").strip().lower()
    if requested_status not in {"draft", "sent"}:
        return error_response("invoice status must be draft or sent")

    tax_amount = round(subtotal * (tax_rate / 100), 2)
    total_amount = round(subtotal + tax_amount, 2)
    invoice = Invoice(
        org_id=user.org_id,
        company_id=company.id,
        invoice_number=generate_document_number(Invoice, company.id, "INV"),
        customer_name=customer_name,
        customer_email=customer_email,
        status=requested_status,
        issue_date=issue_date,
        due_date=due_date,
        subtotal=subtotal,
        tax_rate=tax_rate,
        tax_amount=tax_amount,
        total_amount=total_amount,
        balance_due=total_amount,
        notes=(data.get("notes") or "").strip() or None,
        created_by=user.id,
        last_sent_at=datetime.datetime.now(datetime.UTC) if requested_status == "sent" else None,
    )
    db.session.add(invoice)
    db.session.flush()

    for item in items:
        db.session.add(
            InvoiceItem(
                invoice_id=invoice.id,
                description=item["description"],
                quantity=item["quantity"],
                unit_price=item["unit_price"],
                amount=item["amount"],
            )
        )

    refresh_invoice_status(invoice)
    if requested_status == "sent":
        journal_lines = [
            {"account_code": "1100", "debit": total_amount, "credit": 0, "description": customer_name},
            {"account_code": "4000", "debit": 0, "credit": subtotal, "description": customer_name},
        ]
        if tax_amount > 0:
            journal_lines.append({"account_code": "2100", "debit": 0, "credit": tax_amount, "description": customer_name})
        post_operational_entry(
            company,
            user,
            source_type="invoice_issue",
            source_id=invoice.id,
            memo=f"Invoice {invoice.invoice_number} issued",
            lines=journal_lines,
            entry_date=issue_date,
            reference=invoice.invoice_number,
        )
    if not safe_commit():
        return error_response("database error while creating invoice", 503)

    log(user.id, f"created invoice {invoice.invoice_number}")
    return serialize_invoice(invoice), 201


@app.route("/finance/invoices/<int:invoice_id>/status", methods=["PATCH"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant", "cashier")
def update_invoice_status(invoice_id):
    user = get_user_from_token()
    invoice = get_invoice_for_user(user, invoice_id)
    if not invoice:
        return error_response("invoice not found", 404)

    data = request.get_json(silent=True) or {}
    next_status = (data.get("status") or "").strip().lower()
    if next_status not in {"draft", "sent", "cancelled"}:
        return error_response("invalid invoice status")
    if next_status == "cancelled" and total_customer_payments(invoice.id) > 0:
        return error_response("cannot cancel an invoice that has payments")

    previous_status = invoice.status
    invoice.status = next_status
    if next_status == "sent":
        invoice.last_sent_at = datetime.datetime.now(datetime.UTC)
        if previous_status == "draft":
            journal_lines = [
                {"account_code": "1100", "debit": float(invoice.total_amount or 0), "credit": 0, "description": invoice.customer_name},
                {"account_code": "4000", "debit": 0, "credit": float(invoice.subtotal or 0), "description": invoice.customer_name},
            ]
            if float(invoice.tax_amount or 0) > 0:
                journal_lines.append({"account_code": "2100", "debit": 0, "credit": float(invoice.tax_amount or 0), "description": invoice.customer_name})
            post_operational_entry(
                resolve_company_for_user(user, invoice.company_id),
                user,
                source_type="invoice_issue",
                source_id=invoice.id,
                memo=f"Invoice {invoice.invoice_number} issued",
                lines=journal_lines,
                entry_date=invoice.issue_date or today_utc_date(),
                reference=invoice.invoice_number,
            )
    elif next_status == "cancelled" and previous_status in {"sent", "partial", "overdue"}:
        company = resolve_company_for_user(user, invoice.company_id)
        reverse_source_entries(company, user, "invoice_issue", invoice.id, f"Invoice {invoice.invoice_number} cancelled")
    if not safe_commit():
        return error_response("database error while updating invoice status", 503)

    refresh_invoice_status(invoice)
    safe_commit()
    log(user.id, f"updated invoice {invoice.invoice_number} to {next_status}")
    return serialize_invoice(invoice)


@app.route("/finance/invoices/<int:invoice_id>/payments", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant", "cashier")
def create_invoice_payment(invoice_id):
    user = get_user_from_token()
    invoice = get_invoice_for_user(user, invoice_id)
    if not invoice:
        return error_response("invoice not found", 404)

    data = request.get_json(silent=True) or {}
    try:
        refresh_invoice_status(invoice)
        amount = parse_money(data.get("amount", invoice.balance_due), "amount")
        payment_date = parse_iso_date(data.get("payment_date"), "payment_date", today_utc_date())
        apply_customer_payment(
            invoice,
            amount=amount,
            payment_date=payment_date,
            reference=data.get("reference") or "",
            notes=data.get("notes") or "",
        )
        payment_record = CustomerPayment.query.filter_by(invoice_id=invoice.id).order_by(CustomerPayment.id.desc()).first()
        company = resolve_company_for_user(user, invoice.company_id)
        post_operational_entry(
            company,
            user,
            source_type="invoice_payment",
            source_id=payment_record.id if payment_record else None,
            memo=f"Payment received for {invoice.invoice_number}",
            lines=[
                {"account_code": "1000", "debit": amount, "credit": 0, "description": invoice.customer_name},
                {"account_code": "1100", "debit": 0, "credit": amount, "description": invoice.customer_name},
            ],
            entry_date=payment_date,
            reference=invoice.invoice_number,
        )
    except ValueError as exc:
        db.session.rollback()
        return error_response(str(exc))

    if not safe_commit():
        return error_response("database error while recording payment", 503)

    log(user.id, f"recorded payment on {invoice.invoice_number}")
    return serialize_invoice(invoice)


@app.route("/finance/receivables")
@jwt_required()
def receivables_summary():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    invoices, _ = refresh_finance_documents(company.id)
    scoped = [invoice for invoice in invoices if invoice.org_id == user.org_id]
    return calculate_aging(scoped, "invoice_number")


@app.route("/finance/bills")
@jwt_required()
def list_bills():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    _, bills = refresh_finance_documents(company.id)
    scoped = [bill for bill in bills if bill.org_id == user.org_id]
    scoped.sort(key=lambda bill: (bill.issue_date or today_utc_date(), bill.id), reverse=True)
    return {"items": [serialize_bill(bill) for bill in scoped]}


@app.route("/finance/bills", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
def create_bill():
    user = get_user_from_token()
    data = request.get_json(silent=True) or {}
    company = resolve_company_for_user(user, data.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    vendor_name = (data.get("vendor_name") or "").strip()
    if not vendor_name:
        return error_response("vendor_name is required")

    try:
        issue_date = parse_iso_date(data.get("issue_date"), "issue_date", today_utc_date())
        due_date = parse_iso_date(data.get("due_date"), "due_date", issue_date + datetime.timedelta(days=30))
        tax_rate = parse_money(data.get("tax_rate", 0), "tax_rate")
        items, subtotal = normalize_document_items(data.get("items"), "bill")
    except ValueError as exc:
        return error_response(str(exc))

    requested_status = (data.get("status") or "draft").strip().lower()
    if requested_status not in {"draft", "approved"}:
        return error_response("bill status must be draft or approved")

    tax_amount = round(subtotal * (tax_rate / 100), 2)
    total_amount = round(subtotal + tax_amount, 2)
    get_or_create_vendor_profile(
        company,
        vendor_name,
        {
            "email": data.get("vendor_email"),
            "tax_id": data.get("vendor_tax_id"),
            "default_payment_rail": data.get("default_payment_rail"),
            "is_1099_eligible": data.get("is_1099_eligible"),
            "tax_form_type": data.get("tax_form_type"),
            "tin_status": data.get("tin_status"),
        },
    )
    bill = VendorBill(
        org_id=user.org_id,
        company_id=company.id,
        bill_number=generate_document_number(VendorBill, company.id, "BILL"),
        vendor_name=vendor_name,
        status=requested_status,
        issue_date=issue_date,
        due_date=due_date,
        subtotal=subtotal,
        tax_rate=tax_rate,
        tax_amount=tax_amount,
        total_amount=total_amount,
        balance_due=total_amount,
        notes=(data.get("notes") or "").strip() or None,
        created_by=user.id,
        approved_at=datetime.datetime.now(datetime.UTC) if requested_status == "approved" else None,
    )
    db.session.add(bill)
    db.session.flush()

    for item in items:
        db.session.add(
            VendorBillItem(
                bill_id=bill.id,
                description=item["description"],
                quantity=item["quantity"],
                unit_price=item["unit_price"],
                amount=item["amount"],
            )
        )

    refresh_bill_status(bill)
    if requested_status == "approved":
        journal_lines = [
            {"account_code": "5200", "debit": subtotal, "credit": 0, "description": vendor_name},
            {"account_code": "2000", "debit": 0, "credit": total_amount, "description": vendor_name},
        ]
        if tax_amount > 0:
            journal_lines.insert(1, {"account_code": "1250", "debit": tax_amount, "credit": 0, "description": vendor_name})
        post_operational_entry(
            company,
            user,
            source_type="bill_issue",
            source_id=bill.id,
            memo=f"Vendor bill {bill.bill_number} approved",
            lines=journal_lines,
            entry_date=issue_date,
            reference=bill.bill_number,
        )
    if not safe_commit():
        return error_response("database error while creating bill", 503)

    log(user.id, f"created bill {bill.bill_number}")
    return serialize_bill(bill), 201


@app.route("/finance/bills/<int:bill_id>/status", methods=["PATCH"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
def update_bill_status(bill_id):
    user = get_user_from_token()
    bill = get_bill_for_user(user, bill_id)
    if not bill:
        return error_response("bill not found", 404)

    data = request.get_json(silent=True) or {}
    next_status = (data.get("status") or "").strip().lower()
    if next_status not in {"draft", "approved", "cancelled"}:
        return error_response("invalid bill status")
    if next_status == "cancelled" and total_vendor_payments(bill.id) > 0:
        return error_response("cannot cancel a bill that has payments")

    previous_status = bill.status
    bill.status = next_status
    if next_status == "approved":
        bill.approved_at = datetime.datetime.now(datetime.UTC)
        if previous_status == "draft":
            company = resolve_company_for_user(user, bill.company_id)
            journal_lines = [
                {"account_code": "5200", "debit": float(bill.subtotal or 0), "credit": 0, "description": bill.vendor_name},
                {"account_code": "2000", "debit": 0, "credit": float(bill.total_amount or 0), "description": bill.vendor_name},
            ]
            if float(bill.tax_amount or 0) > 0:
                journal_lines.insert(1, {"account_code": "1250", "debit": float(bill.tax_amount or 0), "credit": 0, "description": bill.vendor_name})
            post_operational_entry(
                company,
                user,
                source_type="bill_issue",
                source_id=bill.id,
                memo=f"Vendor bill {bill.bill_number} approved",
                lines=journal_lines,
                entry_date=bill.issue_date or today_utc_date(),
                reference=bill.bill_number,
            )
    elif next_status == "cancelled" and previous_status in {"approved", "partial", "overdue"}:
        company = resolve_company_for_user(user, bill.company_id)
        reverse_source_entries(company, user, "bill_issue", bill.id, f"Vendor bill {bill.bill_number} cancelled")
    if not safe_commit():
        return error_response("database error while updating bill status", 503)

    refresh_bill_status(bill)
    safe_commit()
    log(user.id, f"updated bill {bill.bill_number} to {next_status}")
    return serialize_bill(bill)


@app.route("/finance/bills/<int:bill_id>/payments", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
def create_bill_payment(bill_id):
    user = get_user_from_token()
    bill = get_bill_for_user(user, bill_id)
    if not bill:
        return error_response("bill not found", 404)

    data = request.get_json(silent=True) or {}
    try:
        refresh_bill_status(bill)
        amount = parse_money(data.get("amount", bill.balance_due), "amount")
        payment_date = parse_iso_date(data.get("payment_date"), "payment_date", today_utc_date())
        apply_vendor_payment(
            bill,
            amount=amount,
            payment_date=payment_date,
            reference=data.get("reference") or "",
            notes=data.get("notes") or "",
        )
        payment_record = VendorPayment.query.filter_by(bill_id=bill.id).order_by(VendorPayment.id.desc()).first()
        company = resolve_company_for_user(user, bill.company_id)
        post_operational_entry(
            company,
            user,
            source_type="bill_payment",
            source_id=payment_record.id if payment_record else None,
            memo=f"Payment sent for {bill.bill_number}",
            lines=[
                {"account_code": "2000", "debit": amount, "credit": 0, "description": bill.vendor_name},
                {"account_code": "1000", "debit": 0, "credit": amount, "description": bill.vendor_name},
            ],
            entry_date=payment_date,
            reference=bill.bill_number,
        )
    except ValueError as exc:
        db.session.rollback()
        return error_response(str(exc))

    if not safe_commit():
        return error_response("database error while recording bill payment", 503)

    log(user.id, f"recorded payment on {bill.bill_number}")
    return serialize_bill(bill)


@app.route("/finance/payables")
@jwt_required()
def payables_summary():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    _, bills = refresh_finance_documents(company.id)
    scoped = [bill for bill in bills if bill.org_id == user.org_id]
    return calculate_aging(scoped, "bill_number")


@app.route("/finance/bank-transactions")
@jwt_required()
def list_bank_transactions():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    transactions = (
        BankFeedTransaction.query.filter_by(org_id=user.org_id, company_id=company.id)
        .order_by(BankFeedTransaction.posted_at.desc(), BankFeedTransaction.id.desc())
        .limit(100)
        .all()
    )
    return {"items": [serialize_bank_transaction(transaction) for transaction in transactions]}


@app.route("/finance/banking/providers")
@jwt_required()
def banking_providers():
    return {
        "plaid": {
            "enabled": plaid_enabled(),
            "environment": (os.getenv("PLAID_ENV", "sandbox") or "sandbox").strip().lower(),
            "products": plaid_products(),
            "country_codes": plaid_country_codes(),
        }
    }


@app.route("/finance/banking/connections")
@jwt_required()
def banking_connections():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    rows = (
        BankConnection.query.filter_by(org_id=user.org_id, company_id=company.id)
        .order_by(BankConnection.created_at.desc(), BankConnection.id.desc())
        .all()
    )
    return {"items": [serialize_bank_connection(row) for row in rows]}


@app.route("/finance/banking/plaid/link-token", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
@plan_required("pro")
def create_plaid_link_token():
    user = get_user_from_token()
    data = request.get_json(silent=True) or {}
    company = resolve_company_for_user(user, data.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    try:
        response = call_plaid(
            "/link/token/create",
            {
                "client_name": "FinancePro",
                "language": "en",
                "country_codes": plaid_country_codes(),
                "products": plaid_products(),
                "user": {"client_user_id": f"{user.org_id}:{user.id}:{company.id}"},
                "transactions": {"days_requested": 730},
                **({"webhook": os.getenv("PLAID_WEBHOOK_URL")} if os.getenv("PLAID_WEBHOOK_URL") else {}),
            },
        )
    except ValueError as exc:
        return error_response(str(exc), 503)

    return {"link_token": response.get("link_token")}


@app.route("/finance/banking/plaid/exchange-token", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
@plan_required("pro")
def exchange_plaid_public_token():
    user = get_user_from_token()
    data = request.get_json(silent=True) or {}
    company = resolve_company_for_user(user, data.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    public_token = (data.get("public_token") or "").strip()
    if not public_token:
        return error_response("public_token is required")

    try:
        exchange = call_plaid(
            "/item/public_token/exchange",
            {
                "public_token": public_token,
            },
        )
    except ValueError as exc:
        return error_response(str(exc), 503)

    item_id = exchange.get("item_id")
    access_token = exchange.get("access_token")
    if not item_id or not access_token:
        return error_response("Plaid did not return an item_id or access_token", 503)

    existing = BankConnection.query.filter_by(item_id=item_id).first()
    if existing:
        existing.company_id = company.id
        existing.org_id = user.org_id
        existing.access_token = access_token
        existing.institution_name = (data.get("institution_name") or existing.institution_name or "").strip() or None
        existing.status = "connected"
        existing.updated_at = datetime.datetime.now(datetime.UTC)
        connection = existing
    else:
        connection = BankConnection(
            org_id=user.org_id,
            company_id=company.id,
            provider="plaid",
            item_id=item_id,
            access_token=access_token,
            institution_name=(data.get("institution_name") or "").strip() or None,
            status="connected",
        )
        db.session.add(connection)

    if not safe_commit():
        return error_response("database error while saving bank connection", 503)

    log(user.id, f"connected bank feed for company {company.name} via Plaid")
    return serialize_bank_connection(connection), 201


@app.route("/finance/banking/plaid/sync", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
@plan_required("pro")
def sync_plaid_transactions():
    user = get_user_from_token()
    data = request.get_json(silent=True) or {}
    company = resolve_company_for_user(user, data.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    connection_id = data.get("connection_id")
    query = BankConnection.query.filter_by(org_id=user.org_id, company_id=company.id, provider="plaid")
    connection = query.filter_by(id=connection_id).first() if connection_id else query.order_by(BankConnection.id.desc()).first()
    if not connection:
        return error_response("no Plaid bank connection found for this company", 404)

    try:
        result = sync_plaid_connection(connection)
    except ValueError as exc:
        return error_response(str(exc), 503)

    log(user.id, f"synced Plaid transactions for {company.name}")
    return {
        "connection": serialize_bank_connection(connection),
        **result,
    }


@app.route("/finance/bank-feed/import", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
def import_bank_feed():
    user = get_user_from_token()
    company = resolve_company_for_user(user, request.form.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    uploaded_file = request.files.get("file")
    if not uploaded_file:
        return error_response("file is required")

    try:
        raw_df = read_external_dataframe(uploaded_file)
        normalized_df = normalize_bank_feed_dataframe(raw_df)
    except Exception as exc:
        return error_response(f"invalid bank feed: {exc}")

    imported = 0
    skipped = 0
    imported_rows = []
    for row in normalized_df.to_dict(orient="records"):
        existing = BankFeedTransaction.query.filter_by(
            org_id=user.org_id,
            company_id=company.id,
            posted_at=row["posted_at"],
            description=str(row["description"]).strip(),
            amount=round(float(row["amount"] or 0), 2),
            reference=(str(row.get("reference", "")).strip() or None),
        ).first()
        if existing:
            skipped += 1
            continue

        transaction = BankFeedTransaction(
            org_id=user.org_id,
            company_id=company.id,
            posted_at=row["posted_at"],
            description=str(row["description"]).strip(),
            amount=round(float(row["amount"] or 0), 2),
            reference=(str(row.get("reference", "")).strip() or None),
            raw_payload=json.dumps(
                {
                    "posted_at": iso_date(row["posted_at"]),
                    "description": str(row["description"]).strip(),
                    "amount": round(float(row["amount"] or 0), 2),
                    "reference": str(row.get("reference", "")).strip(),
                }
            ),
        )
        db.session.add(transaction)
        db.session.flush()
        imported += 1
        imported_rows.append(serialize_bank_transaction(transaction))

    if not safe_commit():
        return error_response("database error while importing bank feed", 503)

    log(user.id, f"imported bank feed with {imported} transactions")
    return {"imported": imported, "skipped": skipped, "items": imported_rows}


@app.route("/finance/reconciliation/suggestions")
@jwt_required()
def reconciliation_suggestions():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    return {"items": build_reconciliation_suggestions(company)}


@app.route("/finance/reconciliation/match", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
def reconcile_transaction():
    user = get_user_from_token()
    data = request.get_json(silent=True) or {}
    transaction_id = data.get("transaction_id")
    entity_type = (data.get("entity_type") or "").strip().lower()
    entity_id = data.get("entity_id")

    if not transaction_id or not entity_id or entity_type not in {"invoice", "bill"}:
        return error_response("transaction_id, entity_type, and entity_id are required")

    transaction = get_bank_transaction_for_user(user, int(transaction_id))
    if not transaction:
        return error_response("bank transaction not found", 404)
    if transaction.status not in {"unmatched", "rule_matched"}:
        return error_response("bank transaction is already matched", 400)

    payment_date = parse_iso_date(data.get("payment_date"), "payment_date", transaction.posted_at or today_utc_date())
    amount = round(abs(float(transaction.amount or 0)), 2)
    if amount <= 0:
        return error_response("bank transaction amount must be non-zero")

    try:
        if entity_type == "invoice":
            if float(transaction.amount or 0) < 0:
                return error_response("outflow transaction cannot be matched to an invoice")
            invoice = get_invoice_for_user(user, int(entity_id))
            if not invoice or invoice.company_id != transaction.company_id:
                return error_response("invoice not found", 404)
            apply_customer_payment(
                invoice,
                amount=amount,
                payment_date=payment_date,
                reference=transaction.reference or transaction.description,
                source="bank_feed",
                notes="Matched from bank feed reconciliation",
                bank_transaction_id=transaction.id,
            )
            transaction.status = "matched"
            transaction.matched_invoice_id = invoice.id
            transaction.matched_bill_id = None
            matched_payload = serialize_invoice(invoice)
            log(user.id, f"matched bank transaction {transaction.id} to invoice {invoice.invoice_number}")
        else:
            if float(transaction.amount or 0) > 0:
                return error_response("inflow transaction cannot be matched to a bill")
            bill = get_bill_for_user(user, int(entity_id))
            if not bill or bill.company_id != transaction.company_id:
                return error_response("bill not found", 404)
            apply_vendor_payment(
                bill,
                amount=amount,
                payment_date=payment_date,
                reference=transaction.reference or transaction.description,
                source="bank_feed",
                notes="Matched from bank feed reconciliation",
                bank_transaction_id=transaction.id,
            )
            transaction.status = "matched"
            transaction.matched_invoice_id = None
            transaction.matched_bill_id = bill.id
            matched_payload = serialize_bill(bill)
            log(user.id, f"matched bank transaction {transaction.id} to bill {bill.bill_number}")
    except ValueError as exc:
        db.session.rollback()
        return error_response(str(exc))

    if not safe_commit():
        return error_response("database error while reconciling transaction", 503)

    return {"transaction": serialize_bank_transaction(transaction), "matched": matched_payload}


@app.route("/finance/tax/summary")
@jwt_required()
def tax_summary():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    profile = get_or_create_tax_profile(company)
    summary = calculate_tax_summary(company, profile)
    summary["company_id"] = company.id
    return summary


@app.route("/finance/tax/profile")
@jwt_required()
def tax_profile():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    return serialize_tax_profile(get_or_create_tax_profile(company))


@app.route("/finance/tax/profile", methods=["PUT"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
def update_tax_profile():
    user = get_user_from_token()
    data = request.get_json(silent=True) or {}
    company = resolve_company_for_user(user, data.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    profile = get_or_create_tax_profile(company)
    jurisdiction_code = (data.get("jurisdiction_code") or profile.jurisdiction_code or "generic").strip().lower()
    filing_frequency = (data.get("filing_frequency") or profile.filing_frequency or "monthly").strip().lower()
    if filing_frequency not in VALID_TAX_FILING_FREQUENCIES:
        return error_response("invalid filing_frequency")

    try:
        indirect_tax_rate = parse_money(data.get("indirect_tax_rate", profile.indirect_tax_rate), "indirect_tax_rate")
        income_tax_rate = parse_money(data.get("income_tax_rate", profile.income_tax_rate), "income_tax_rate")
        period_start_month = int(data.get("period_start_month", profile.period_start_month or 1))
    except ValueError as exc:
        return error_response(str(exc))

    if period_start_month < 1 or period_start_month > 12:
        return error_response("period_start_month must be between 1 and 12")

    profile.jurisdiction_code = jurisdiction_code
    profile.filing_frequency = filing_frequency
    profile.registration_number = (data.get("registration_number") or "").strip() or None
    profile.currency_code = (data.get("currency_code") or profile.currency_code or "USD").strip().upper()
    profile.sales_tax_name = (data.get("sales_tax_name") or profile.sales_tax_name or "Sales Tax").strip()
    profile.purchase_tax_name = (data.get("purchase_tax_name") or profile.purchase_tax_name or "Purchase Tax Credit").strip()
    profile.indirect_tax_rate = indirect_tax_rate
    profile.income_tax_rate = income_tax_rate
    profile.period_start_month = period_start_month

    if not safe_commit():
        return error_response("database error while updating tax profile", 503)

    log(user.id, f"updated tax profile for {company.name}")
    return serialize_tax_profile(profile)


@app.route("/finance/tax/filing-preview")
@jwt_required()
def tax_filing_preview():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    profile = get_or_create_tax_profile(company)
    preview = build_tax_filing_preview(
        company,
        profile,
        period_start=request.args.get("period_start"),
        period_end=request.args.get("period_end"),
    )
    preview["profile"] = serialize_tax_profile(profile)
    return preview


@app.route("/finance/chart-of-accounts")
@jwt_required()
def list_chart_of_accounts():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    seed_chart_of_accounts(company)
    safe_commit()
    accounts = LedgerAccount.query.filter_by(company_id=company.id).order_by(LedgerAccount.code.asc()).all()
    return {"items": [serialize_ledger_account(account) for account in accounts]}


@app.route("/finance/chart-of-accounts", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
def create_chart_of_account():
    user = get_user_from_token()
    data = request.get_json(silent=True) or {}
    company = resolve_company_for_user(user, data.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    code = (data.get("code") or "").strip()
    name = (data.get("name") or "").strip()
    category = (data.get("category") or "").strip().lower()
    normal_balance = (data.get("normal_balance") or ("credit" if category in {"liability", "equity", "revenue"} else "debit")).strip().lower()

    if not code or not name:
        return error_response("code and name are required")
    if category not in VALID_ACCOUNT_CATEGORIES:
        return error_response("invalid account category")
    if normal_balance not in {"debit", "credit"}:
        return error_response("normal_balance must be debit or credit")
    if LedgerAccount.query.filter_by(company_id=company.id, code=code).first():
        return error_response("account code already exists", 409)

    account = LedgerAccount(
        org_id=company.org_id,
        company_id=company.id,
        code=code,
        name=name,
        category=category,
        subtype=(data.get("subtype") or "").strip() or None,
        normal_balance=normal_balance,
        description=(data.get("description") or "").strip() or None,
        is_system=False,
        is_active=parse_bool(data.get("is_active"), True),
    )
    db.session.add(account)
    if not safe_commit():
        return error_response("database error while creating account", 503)

    log(user.id, f"created chart of account {code}")
    return serialize_ledger_account(account), 201


@app.route("/finance/chart-of-accounts/seed", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
def seed_chart_of_accounts_route():
    user = get_user_from_token()
    data = request.get_json(silent=True) or {}
    company = resolve_company_for_user(user, data.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    created = seed_chart_of_accounts(company)
    if not safe_commit():
        return error_response("database error while seeding chart of accounts", 503)

    log(user.id, f"seeded chart of accounts for {company.name}")
    return {"created": created}


@app.route("/finance/accounting/overview")
@jwt_required()
def accounting_overview():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    return build_accounting_overview(company)


@app.route("/finance/journal-entries")
@jwt_required()
def list_journal_entries():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    entries = (
        JournalEntry.query.filter_by(company_id=company.id)
        .order_by(JournalEntry.entry_date.desc(), JournalEntry.id.desc())
        .limit(50)
        .all()
    )
    return {"items": [serialize_journal_entry(entry) for entry in entries]}


@app.route("/finance/journal-entries", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
def create_journal_entry_route():
    user = get_user_from_token()
    data = request.get_json(silent=True) or {}
    company = resolve_company_for_user(user, data.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    memo = (data.get("memo") or "").strip()
    if not memo:
        return error_response("memo is required")

    try:
        entry_date = parse_iso_date(data.get("entry_date"), "entry_date", today_utc_date())
        entry = post_journal_entry(
            company,
            user,
            entry_date=entry_date,
            memo=memo,
            lines=data.get("lines"),
            source_type="manual",
            source_id=None,
            reference=data.get("reference"),
        )
    except ValueError as exc:
        db.session.rollback()
        return error_response(str(exc))

    if not safe_commit():
        return error_response("database error while posting journal entry", 503)

    log(user.id, f"posted journal entry {entry.entry_number}")
    return serialize_journal_entry(entry), 201


@app.route("/finance/register")
@jwt_required()
def account_register():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    seed_chart_of_accounts(company)
    account = get_company_account(company.id, request.args.get("account_id"), request.args.get("account_code"))
    if not account:
        account = LedgerAccount.query.filter_by(company_id=company.id).order_by(LedgerAccount.code.asc()).first()
    if not account:
        return error_response("account not found", 404)

    return build_account_register(company, account)


@app.route("/finance/vendors")
@jwt_required()
def list_vendors():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    vendors = VendorProfile.query.filter_by(company_id=company.id).order_by(VendorProfile.vendor_name.asc()).all()
    return {"items": [serialize_vendor_profile(vendor) for vendor in vendors]}


@app.route("/finance/vendors", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
def create_vendor():
    user = get_user_from_token()
    data = request.get_json(silent=True) or {}
    company = resolve_company_for_user(user, data.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    try:
        vendor = get_or_create_vendor_profile(company, data.get("vendor_name"), data)
    except ValueError as exc:
        return error_response(str(exc))

    vendor.email = (data.get("email") or vendor.email or "").strip().lower() or None
    vendor.tax_id = (data.get("tax_id") or vendor.tax_id or "").strip() or None
    vendor.default_payment_rail = (data.get("default_payment_rail") or vendor.default_payment_rail or "ach").strip().lower()
    vendor.remittance_reference = (data.get("remittance_reference") or vendor.remittance_reference or "").strip() or None
    vendor.bank_last4 = ((data.get("bank_last4") or vendor.bank_last4 or "").strip()[-4:] or None)
    vendor.is_1099_eligible = parse_bool(data.get("is_1099_eligible"), vendor.is_1099_eligible)
    vendor.tax_form_type = (data.get("tax_form_type") or vendor.tax_form_type or "1099-NEC").strip().upper()
    vendor.tin_status = (data.get("tin_status") or vendor.tin_status or "pending").strip().lower()

    if vendor.default_payment_rail not in VALID_PAYMENT_RAILS:
        return error_response("invalid payment rail")

    if not safe_commit():
        return error_response("database error while saving vendor", 503)

    log(user.id, f"saved vendor profile {vendor.vendor_name}")
    return serialize_vendor_profile(vendor), 201


@app.route("/finance/vendors/1099-summary")
@jwt_required()
def vendor_1099_summary():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    year = request.args.get("year")
    try:
        parsed_year = int(year) if year else None
    except ValueError:
        return error_response("year must be numeric")

    return build_1099_summary(company, parsed_year)


@app.route("/finance/bill-pay/summary")
@jwt_required()
def bill_pay_summary():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    return build_bill_pay_summary(company)


@app.route("/finance/bill-pay/disbursements", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
@plan_required("pro")
def create_bill_disbursement():
    user = get_user_from_token()
    data = request.get_json(silent=True) or {}
    company = resolve_company_for_user(user, data.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    bill = get_bill_for_user(user, data.get("bill_id"))
    if not bill or bill.company_id != company.id:
        return error_response("bill not found", 404)

    try:
        refresh_bill_status(bill)
        amount = parse_money(data.get("amount", bill.balance_due), "amount")
        scheduled_date = parse_iso_date(data.get("scheduled_date"), "scheduled_date", today_utc_date())
    except ValueError as exc:
        return error_response(str(exc))

    if amount <= 0 or amount - float(bill.balance_due or 0) > 0.01:
        return error_response("scheduled amount must be positive and within the open bill balance")

    vendor = get_or_create_vendor_profile(company, bill.vendor_name, data)
    rail = (data.get("payment_rail") or vendor.default_payment_rail or "ach").strip().lower()
    if rail not in VALID_PAYMENT_RAILS:
        return error_response("invalid payment rail")

    compliance_status = "ready" if (not vendor.is_1099_eligible or vendor.tin_status == "verified" or vendor.tin_status == "received") else "needs_tin_review"
    disbursement = BillDisbursement(
        org_id=company.org_id,
        company_id=company.id,
        bill_id=bill.id,
        vendor_profile_id=vendor.id,
        payment_rail=rail,
        status="scheduled",
        scheduled_date=scheduled_date,
        amount=amount,
        reference=(data.get("reference") or bill.bill_number).strip(),
        compliance_status=compliance_status,
        created_by=user.id,
    )
    db.session.add(disbursement)
    if not safe_commit():
        return error_response("database error while scheduling disbursement", 503)

    log(user.id, f"scheduled bill pay for {bill.bill_number}")
    return serialize_disbursement(disbursement), 201


@app.route("/finance/bill-pay/disbursements/<int:disbursement_id>/execute", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
@plan_required("pro")
def execute_bill_disbursement(disbursement_id):
    user = get_user_from_token()
    disbursement = BillDisbursement.query.filter_by(id=disbursement_id, org_id=user.org_id).first()
    if not disbursement:
        return error_response("disbursement not found", 404)
    if disbursement.status == "completed":
        return error_response("disbursement already completed")

    bill = get_bill_for_user(user, disbursement.bill_id)
    if not bill:
        return error_response("bill not found", 404)

    try:
        payment_date = parse_iso_date((request.get_json(silent=True) or {}).get("payment_date"), "payment_date", today_utc_date())
        apply_vendor_payment(
            bill,
            amount=float(disbursement.amount or 0),
            payment_date=payment_date,
            reference=disbursement.reference or bill.bill_number,
            source=f"bill_pay_{disbursement.payment_rail}",
            notes=f"Executed through {disbursement.payment_rail}",
        )
        payment_record = VendorPayment.query.filter_by(bill_id=bill.id).order_by(VendorPayment.id.desc()).first()
        company = resolve_company_for_user(user, bill.company_id)
        post_operational_entry(
            company,
            user,
            source_type="bill_payment",
            source_id=payment_record.id if payment_record else None,
            memo=f"Payment sent for {bill.bill_number}",
            lines=[
                {"account_code": "2000", "debit": float(disbursement.amount or 0), "credit": 0, "description": bill.vendor_name},
                {"account_code": "1000", "debit": 0, "credit": float(disbursement.amount or 0), "description": bill.vendor_name},
            ],
            entry_date=payment_date,
            reference=bill.bill_number,
        )
    except ValueError as exc:
        db.session.rollback()
        return error_response(str(exc))

    disbursement.status = "completed"
    disbursement.processed_at = datetime.datetime.now(datetime.UTC)
    disbursement.confirmation_code = f"{disbursement.payment_rail.upper()}-{os.urandom(4).hex().upper()}"

    if not safe_commit():
        return error_response("database error while executing disbursement", 503)

    log(user.id, f"executed bill pay for {bill.bill_number}")
    return serialize_disbursement(disbursement)


@app.route("/finance/reconciliation/rules")
@jwt_required()
def list_reconciliation_rules():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    rules = ReconciliationRule.query.filter_by(company_id=company.id).order_by(ReconciliationRule.priority.asc(), ReconciliationRule.id.asc()).all()
    return {"items": [serialize_reconciliation_rule(rule) for rule in rules]}


@app.route("/finance/reconciliation/rules", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
@plan_required("pro")
def create_reconciliation_rule():
    user = get_user_from_token()
    data = request.get_json(silent=True) or {}
    company = resolve_company_for_user(user, data.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    name = (data.get("name") or "").strip()
    if not name:
        return error_response("name is required")

    direction = (data.get("direction") or "any").strip().lower()
    auto_action = (data.get("auto_action") or "suggest_account").strip().lower()
    if direction not in VALID_RECONCILIATION_DIRECTIONS:
        return error_response("invalid direction")
    if auto_action not in VALID_RECONCILIATION_ACTIONS:
        return error_response("invalid auto_action")

    rule = ReconciliationRule(
        org_id=company.org_id,
        company_id=company.id,
        name=name,
        keyword=(data.get("keyword") or "").strip() or None,
        direction=direction,
        min_amount=parse_money(data.get("min_amount", 0), "min_amount") if data.get("min_amount") not in {None, ""} else None,
        max_amount=parse_money(data.get("max_amount", 0), "max_amount") if data.get("max_amount") not in {None, ""} else None,
        auto_action=auto_action,
        target_reference=(data.get("target_reference") or "").strip() or None,
        exception_type=(data.get("exception_type") or "").strip() or None,
        priority=int(data.get("priority", 100) or 100),
        is_active=parse_bool(data.get("is_active"), True),
    )
    db.session.add(rule)
    if not safe_commit():
        return error_response("database error while saving reconciliation rule", 503)

    log(user.id, f"created reconciliation rule {rule.name}")
    return serialize_reconciliation_rule(rule), 201


@app.route("/finance/reconciliation/rules/auto-apply", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
@plan_required("pro")
def auto_apply_rules_route():
    user = get_user_from_token()
    data = request.get_json(silent=True) or {}
    company = resolve_company_for_user(user, data.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    result = auto_apply_reconciliation_rules(company, user)
    log(user.id, f"auto-applied reconciliation rules for {company.name}")
    return result


@app.route("/finance/reconciliation/workspace")
@jwt_required()
def reconciliation_workspace():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    return build_reconciliation_workspace(company)


@app.route("/finance/reconciliation/exceptions", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
def create_reconciliation_exception():
    user = get_user_from_token()
    data = request.get_json(silent=True) or {}
    company = resolve_company_for_user(user, data.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    transaction = get_bank_transaction_for_user(user, data.get("transaction_id"))
    if not transaction or transaction.company_id != company.id:
        return error_response("bank transaction not found", 404)

    exception_type = (data.get("exception_type") or "").strip().lower()
    if not exception_type:
        return error_response("exception_type is required")

    exception = ReconciliationException(
        org_id=company.org_id,
        company_id=company.id,
        bank_transaction_id=transaction.id,
        exception_type=exception_type,
        notes=(data.get("notes") or "").strip() or None,
        status="open",
        created_by=user.id,
    )
    db.session.add(exception)
    transaction.status = "exception"
    if not safe_commit():
        return error_response("database error while creating reconciliation exception", 503)

    log(user.id, f"flagged reconciliation exception on bank transaction {transaction.id}")
    return serialize_reconciliation_exception(exception), 201


@app.route("/finance/reconciliation/exceptions/<int:exception_id>/resolve", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
def resolve_reconciliation_exception(exception_id):
    user = get_user_from_token()
    exception = ReconciliationException.query.filter_by(id=exception_id, org_id=user.org_id).first()
    if not exception:
        return error_response("reconciliation exception not found", 404)

    transaction = get_bank_transaction_for_user(user, exception.bank_transaction_id)
    exception.status = "resolved"
    exception.resolved_at = datetime.datetime.now(datetime.UTC)
    if transaction and transaction.status == "exception":
        transaction.status = "unmatched"
    if not safe_commit():
        return error_response("database error while resolving reconciliation exception", 503)

    log(user.id, f"resolved reconciliation exception {exception_id}")
    return serialize_reconciliation_exception(exception)


@app.route("/finance/tax/jurisdictions")
@jwt_required()
def tax_jurisdictions():
    return {
        "items": [
            {"code": code, **payload}
            for code, payload in TAX_JURISDICTION_LIBRARY.items()
        ]
    }


@app.route("/finance/tax/filings")
@jwt_required()
def list_tax_filings():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)
    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    filings = TaxFiling.query.filter_by(company_id=company.id).order_by(TaxFiling.period_end.desc(), TaxFiling.id.desc()).all()
    return {"items": [serialize_tax_filing(filing) for filing in filings]}


@app.route("/finance/tax/filings", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
@plan_required("pro")
def prepare_tax_filing():
    user = get_user_from_token()
    data = request.get_json(silent=True) or {}
    company = resolve_company_for_user(user, data.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    profile = get_or_create_tax_profile(company)
    filing_type = (data.get("filing_type") or "indirect_tax").strip().lower()
    if filing_type not in VALID_TAX_FILING_TYPES:
        return error_response("invalid filing_type")

    try:
        package = build_tax_filing_package(
            company,
            profile,
            period_start=data.get("period_start"),
            period_end=data.get("period_end"),
        )
        period_start = parse_iso_date(package["preview"]["period_start"], "period_start")
        period_end = parse_iso_date(package["preview"]["period_end"], "period_end")
    except ValueError as exc:
        return error_response(str(exc))

    filing = TaxFiling(
        org_id=company.org_id,
        company_id=company.id,
        jurisdiction_code=profile.jurisdiction_code,
        filing_frequency=profile.filing_frequency,
        filing_type=filing_type,
        period_start=period_start,
        period_end=period_end,
        status="prepared",
        payload_json=json.dumps(package),
        prepared_by=user.id,
    )
    db.session.add(filing)
    if not safe_commit():
        return error_response("database error while preparing tax filing", 503)

    log(user.id, f"prepared tax filing for {company.name}")
    return serialize_tax_filing(filing), 201


@app.route("/finance/tax/filings/<int:filing_id>/submit", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
@plan_required("pro")
def submit_tax_filing(filing_id):
    user = get_user_from_token()
    filing = TaxFiling.query.filter_by(id=filing_id, org_id=user.org_id).first()
    if not filing:
        return error_response("tax filing not found", 404)

    filing.status = "submitted"
    filing.submitted_at = datetime.datetime.now(datetime.UTC)
    filing.reference = filing.reference or f"TAX-{filing.company_id}-{filing.id}"
    if not safe_commit():
        return error_response("database error while submitting tax filing", 503)

    log(user.id, f"submitted tax filing {filing.id}")
    return serialize_tax_filing(filing)


@app.route("/finance/workforce/overview")
@jwt_required()
def workforce_overview():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)
    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)
    return build_workforce_overview(company)


@app.route("/finance/workforce/employees")
@jwt_required()
def list_employees():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)
    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)
    rows = EmployeeProfile.query.filter_by(company_id=company.id).order_by(EmployeeProfile.full_name.asc()).all()
    return {"items": [serialize_employee(row) for row in rows]}


@app.route("/finance/workforce/employees", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
def create_employee():
    user = get_user_from_token()
    data = request.get_json(silent=True) or {}
    company = resolve_company_for_user(user, data.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    full_name = (data.get("full_name") or "").strip()
    pay_type = (data.get("pay_type") or "hourly").strip().lower()
    if not full_name:
        return error_response("full_name is required")
    if pay_type not in VALID_PAY_TYPES:
        return error_response("invalid pay_type")

    try:
        employee = EmployeeProfile(
            org_id=company.org_id,
            company_id=company.id,
            full_name=full_name,
            email=(data.get("email") or "").strip().lower() or None,
            pay_type=pay_type,
            hourly_rate=parse_money(data.get("hourly_rate", 0), "hourly_rate"),
            salary_amount=parse_money(data.get("salary_amount", 0), "salary_amount"),
            withholding_rate=parse_money(data.get("withholding_rate", 0), "withholding_rate"),
            benefit_rate=parse_money(data.get("benefit_rate", 0), "benefit_rate"),
            is_active=parse_bool(data.get("is_active"), True),
        )
    except ValueError as exc:
        return error_response(str(exc))

    db.session.add(employee)
    if not safe_commit():
        return error_response("database error while creating employee", 503)

    log(user.id, f"created employee {employee.full_name}")
    return serialize_employee(employee), 201


@app.route("/finance/workforce/contractors")
@jwt_required()
def list_contractors():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)
    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)
    rows = ContractorProfile.query.filter_by(company_id=company.id).order_by(ContractorProfile.full_name.asc()).all()
    return {"items": [serialize_contractor(row) for row in rows]}


@app.route("/finance/workforce/contractors", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
def create_contractor():
    user = get_user_from_token()
    data = request.get_json(silent=True) or {}
    company = resolve_company_for_user(user, data.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    full_name = (data.get("full_name") or "").strip()
    if not full_name:
        return error_response("full_name is required")

    try:
        contractor = ContractorProfile(
            org_id=company.org_id,
            company_id=company.id,
            full_name=full_name,
            email=(data.get("email") or "").strip().lower() or None,
            tax_id=(data.get("tax_id") or "").strip() or None,
            default_rate=parse_money(data.get("default_rate", 0), "default_rate"),
            is_1099_eligible=parse_bool(data.get("is_1099_eligible"), True),
            tax_form_type=(data.get("tax_form_type") or "1099-NEC").strip().upper(),
            is_active=parse_bool(data.get("is_active"), True),
        )
    except ValueError as exc:
        return error_response(str(exc))

    db.session.add(contractor)
    if not safe_commit():
        return error_response("database error while creating contractor", 503)

    log(user.id, f"created contractor {contractor.full_name}")
    return serialize_contractor(contractor), 201


@app.route("/finance/workforce/time")
@jwt_required()
def list_time_entries():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)
    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)
    rows = TimeEntry.query.filter_by(company_id=company.id).order_by(TimeEntry.work_date.desc(), TimeEntry.id.desc()).limit(50).all()
    return {"items": [serialize_time_entry(row) for row in rows]}


@app.route("/finance/workforce/time", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant", "cashier")
def create_time_entry():
    user = get_user_from_token()
    data = request.get_json(silent=True) or {}
    company = resolve_company_for_user(user, data.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    employee_id = data.get("employee_id")
    contractor_id = data.get("contractor_id")
    if not employee_id and not contractor_id:
        return error_response("employee_id or contractor_id is required")

    try:
        entry = TimeEntry(
            org_id=company.org_id,
            company_id=company.id,
            employee_id=int(employee_id) if employee_id not in {None, ""} else None,
            contractor_id=int(contractor_id) if contractor_id not in {None, ""} else None,
            project_id=int(data.get("project_id")) if data.get("project_id") not in {None, ""} else None,
            work_date=parse_iso_date(data.get("work_date"), "work_date", today_utc_date()),
            hours=parse_money(data.get("hours", 0), "hours"),
            hourly_cost=parse_money(data.get("hourly_cost", 0), "hourly_cost"),
            billable_rate=parse_money(data.get("billable_rate", 0), "billable_rate"),
            description=(data.get("description") or "").strip() or None,
            status=(data.get("status") or "submitted").strip().lower(),
        )
    except ValueError as exc:
        return error_response(str(exc))

    db.session.add(entry)
    if not safe_commit():
        return error_response("database error while creating time entry", 503)

    log(user.id, f"logged time entry {entry.id}")
    return serialize_time_entry(entry), 201


@app.route("/finance/workforce/mileage")
@jwt_required()
def list_mileage_entries():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)
    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)
    rows = MileageEntry.query.filter_by(company_id=company.id).order_by(MileageEntry.trip_date.desc(), MileageEntry.id.desc()).limit(50).all()
    return {"items": [serialize_mileage_entry(row) for row in rows]}


@app.route("/finance/workforce/mileage", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant", "cashier")
def create_mileage_entry():
    user = get_user_from_token()
    data = request.get_json(silent=True) or {}
    company = resolve_company_for_user(user, data.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    employee_id = data.get("employee_id")
    contractor_id = data.get("contractor_id")
    if not employee_id and not contractor_id:
        return error_response("employee_id or contractor_id is required")

    try:
        entry = MileageEntry(
            org_id=company.org_id,
            company_id=company.id,
            employee_id=int(employee_id) if employee_id not in {None, ""} else None,
            contractor_id=int(contractor_id) if contractor_id not in {None, ""} else None,
            project_id=int(data.get("project_id")) if data.get("project_id") not in {None, ""} else None,
            trip_date=parse_iso_date(data.get("trip_date"), "trip_date", today_utc_date()),
            miles=parse_money(data.get("miles", 0), "miles"),
            rate_per_mile=parse_money(data.get("rate_per_mile", 0.725), "rate_per_mile"),
            purpose=(data.get("purpose") or "").strip() or None,
            status=(data.get("status") or "submitted").strip().lower(),
        )
    except ValueError as exc:
        return error_response(str(exc))

    db.session.add(entry)
    if not safe_commit():
        return error_response("database error while creating mileage entry", 503)

    log(user.id, f"logged mileage entry {entry.id}")
    return serialize_mileage_entry(entry), 201


@app.route("/finance/workforce/payroll-runs")
@jwt_required()
def list_payroll_runs():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)
    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)
    runs = PayrollRun.query.filter_by(company_id=company.id).order_by(PayrollRun.pay_date.desc(), PayrollRun.id.desc()).all()
    return {"items": [serialize_payroll_run(run) for run in runs]}


@app.route("/finance/workforce/payroll-runs", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
def create_payroll_run():
    user = get_user_from_token()
    data = request.get_json(silent=True) or {}
    company = resolve_company_for_user(user, data.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    try:
        period_start = parse_iso_date(data.get("period_start"), "period_start", today_utc_date().replace(day=1))
        period_end = parse_iso_date(data.get("period_end"), "period_end", today_utc_date())
        pay_date = parse_iso_date(data.get("pay_date"), "pay_date", today_utc_date())
    except ValueError as exc:
        return error_response(str(exc))

    employees = EmployeeProfile.query.filter_by(company_id=company.id, is_active=True).all()
    if not employees:
        return error_response("no active employees to process")

    payroll = PayrollRun(
        org_id=company.org_id,
        company_id=company.id,
        payroll_number=generate_payroll_number(company.id),
        period_start=period_start,
        period_end=period_end,
        pay_date=pay_date,
        status="processed",
        created_by=user.id,
    )
    db.session.add(payroll)
    db.session.flush()

    gross_total = 0.0
    withholding_total = 0.0
    benefit_total = 0.0
    mileage_total = 0.0
    for employee in employees:
        time_entries = TimeEntry.query.filter(
            TimeEntry.company_id == company.id,
            TimeEntry.employee_id == employee.id,
            TimeEntry.work_date >= period_start,
            TimeEntry.work_date <= period_end,
        ).all()
        mileage_entries = MileageEntry.query.filter(
            MileageEntry.company_id == company.id,
            MileageEntry.employee_id == employee.id,
            MileageEntry.trip_date >= period_start,
            MileageEntry.trip_date <= period_end,
        ).all()
        hours = round(sum(float(entry.hours or 0) for entry in time_entries), 2)
        gross_pay = round(hours * float(employee.hourly_rate or 0), 2) if employee.pay_type == "hourly" else round(float(employee.salary_amount or 0), 2)
        withholding_amount = round(gross_pay * (float(employee.withholding_rate or 0) / 100), 2)
        benefit_amount = round(gross_pay * (float(employee.benefit_rate or 0) / 100), 2)
        mileage_reimbursement = round(sum(float(entry.miles or 0) * float(entry.rate_per_mile or 0) for entry in mileage_entries), 2)
        net_pay = round(gross_pay - withholding_amount + mileage_reimbursement, 2)
        gross_total += gross_pay
        withholding_total += withholding_amount
        benefit_total += benefit_amount
        mileage_total += mileage_reimbursement
        db.session.add(
            PayrollLine(
                payroll_run_id=payroll.id,
                employee_id=employee.id,
                regular_hours=hours,
                gross_pay=gross_pay,
                withholding_amount=withholding_amount,
                benefit_amount=benefit_amount,
                mileage_reimbursement=mileage_reimbursement,
                net_pay=net_pay,
            )
        )

    payroll.gross_pay = round(gross_total, 2)
    payroll.withholding_total = round(withholding_total, 2)
    payroll.benefit_total = round(benefit_total, 2)
    payroll.mileage_reimbursement_total = round(mileage_total, 2)
    payroll.net_cash = round(gross_total - withholding_total + mileage_total + benefit_total, 2)

    payroll_lines = [
        {"account_code": "5100", "debit": round(gross_total + benefit_total, 2), "credit": 0, "description": payroll.payroll_number},
        {"account_code": "1000", "debit": 0, "credit": round(payroll.net_cash, 2), "description": payroll.payroll_number},
    ]
    if mileage_total > 0:
        payroll_lines.insert(1, {"account_code": "5400", "debit": round(mileage_total, 2), "credit": 0, "description": payroll.payroll_number})
    if withholding_total > 0:
        payroll_lines.insert(-1, {"account_code": "2200", "debit": 0, "credit": round(withholding_total, 2), "description": payroll.payroll_number})

    post_operational_entry(
        company,
        user,
        source_type="payroll_run",
        source_id=payroll.id,
        memo=f"Payroll run {payroll.payroll_number}",
        lines=payroll_lines,
        entry_date=pay_date,
        reference=payroll.payroll_number,
    )

    if not safe_commit():
        return error_response("database error while processing payroll", 503)

    log(user.id, f"processed payroll run {payroll.payroll_number}")
    return serialize_payroll_run(payroll), 201


@app.route("/finance/inventory/summary")
@jwt_required()
def inventory_workspace_summary():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)
    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)
    return build_inventory_summary(company)


@app.route("/finance/inventory/items")
@jwt_required()
def list_inventory_items():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)
    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)
    items = InventoryItem.query.filter_by(company_id=company.id).order_by(InventoryItem.name.asc()).all()
    return {"items": [serialize_inventory_item(item) for item in items]}


@app.route("/finance/inventory/items", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
def create_inventory_item():
    user = get_user_from_token()
    data = request.get_json(silent=True) or {}
    company = resolve_company_for_user(user, data.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    sku = (data.get("sku") or "").strip()
    name = (data.get("name") or "").strip()
    if not sku or not name:
        return error_response("sku and name are required")
    if InventoryItem.query.filter_by(company_id=company.id, sku=sku).first():
        return error_response("sku already exists", 409)

    try:
        item = InventoryItem(
            org_id=company.org_id,
            company_id=company.id,
            sku=sku,
            name=name,
            category=(data.get("category") or "").strip() or None,
            quantity_on_hand=parse_money(data.get("quantity_on_hand", 0), "quantity_on_hand"),
            reorder_point=parse_money(data.get("reorder_point", 0), "reorder_point"),
            reorder_quantity=parse_money(data.get("reorder_quantity", 0), "reorder_quantity"),
            unit_cost=parse_money(data.get("unit_cost", 0), "unit_cost"),
            unit_price=parse_money(data.get("unit_price", 0), "unit_price"),
            preferred_vendor_name=(data.get("preferred_vendor_name") or "").strip() or None,
        )
    except ValueError as exc:
        return error_response(str(exc))

    db.session.add(item)
    db.session.flush()
    if float(item.quantity_on_hand or 0) > 0:
        db.session.add(
            InventoryMovement(
                org_id=company.org_id,
                company_id=company.id,
                inventory_item_id=item.id,
                movement_type="opening_balance",
                quantity_delta=float(item.quantity_on_hand or 0),
                unit_cost=float(item.unit_cost or 0),
                reference=item.sku,
            )
        )
    if not safe_commit():
        return error_response("database error while creating inventory item", 503)

    log(user.id, f"created inventory item {item.sku}")
    return serialize_inventory_item(item), 201


@app.route("/finance/purchase-orders")
@jwt_required()
def list_purchase_orders():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)
    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)
    orders = PurchaseOrder.query.filter_by(company_id=company.id).order_by(PurchaseOrder.issue_date.desc(), PurchaseOrder.id.desc()).all()
    return {"items": [serialize_purchase_order(order) for order in orders]}


@app.route("/finance/purchase-orders", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
def create_purchase_order():
    user = get_user_from_token()
    data = request.get_json(silent=True) or {}
    company = resolve_company_for_user(user, data.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    vendor_name = (data.get("vendor_name") or "").strip()
    if not vendor_name:
        return error_response("vendor_name is required")

    items = data.get("items") or []
    if not isinstance(items, list) or not items:
        return error_response("purchase order requires at least one item")

    try:
        order = PurchaseOrder(
            org_id=company.org_id,
            company_id=company.id,
            po_number=generate_document_number(PurchaseOrder, company.id, "PO"),
            vendor_name=vendor_name,
            status="draft",
            issue_date=parse_iso_date(data.get("issue_date"), "issue_date", today_utc_date()),
            expected_date=parse_iso_date(data.get("expected_date"), "expected_date", today_utc_date() + datetime.timedelta(days=7)),
            notes=(data.get("notes") or "").strip() or None,
            created_by=user.id,
        )
        db.session.add(order)
        db.session.flush()
        for item in items:
            quantity = parse_money(item.get("quantity", 0), "purchase order quantity")
            unit_cost = parse_money(item.get("unit_cost", 0), "purchase order unit_cost")
            sku = (item.get("sku") or "").strip()
            inventory_item = InventoryItem.query.filter_by(company_id=company.id, sku=sku).first() if sku else None
            description = (item.get("description") or "").strip() or (inventory_item.name if inventory_item else "Inventory item")
            db.session.add(
                PurchaseOrderLine(
                    purchase_order_id=order.id,
                    inventory_item_id=inventory_item.id if inventory_item else None,
                    sku=sku or (inventory_item.sku if inventory_item else None),
                    description=description,
                    quantity=quantity,
                    unit_cost=unit_cost,
                    received_quantity=0,
                )
            )
    except ValueError as exc:
        db.session.rollback()
        return error_response(str(exc))

    if not safe_commit():
        return error_response("database error while creating purchase order", 503)

    log(user.id, f"created purchase order {order.po_number}")
    return serialize_purchase_order(order), 201


@app.route("/finance/purchase-orders/<int:purchase_order_id>/submit", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
def submit_purchase_order(purchase_order_id):
    user = get_user_from_token()
    order = PurchaseOrder.query.filter_by(id=purchase_order_id, org_id=user.org_id).first()
    if not order:
        return error_response("purchase order not found", 404)

    order.status = "ordered"
    if not safe_commit():
        return error_response("database error while submitting purchase order", 503)

    log(user.id, f"submitted purchase order {order.po_number}")
    return serialize_purchase_order(order)


@app.route("/finance/purchase-orders/<int:purchase_order_id>/receive", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
def receive_purchase_order(purchase_order_id):
    user = get_user_from_token()
    order = PurchaseOrder.query.filter_by(id=purchase_order_id, org_id=user.org_id).first()
    if not order:
        return error_response("purchase order not found", 404)
    company = resolve_company_for_user(user, order.company_id)
    if not company:
        return error_response("company not found", 404)

    try:
        receive_purchase_order_items(company, user, order, (request.get_json(silent=True) or {}).get("items"))
    except ValueError as exc:
        db.session.rollback()
        return error_response(str(exc))

    if not safe_commit():
        return error_response("database error while receiving purchase order", 503)

    log(user.id, f"received purchase order {order.po_number}")
    return serialize_purchase_order(order)


@app.route("/finance/projects")
@jwt_required()
def list_projects():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)
    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)
    rows = Project.query.filter_by(company_id=company.id).order_by(Project.updated_at.desc(), Project.id.desc()).all()
    return {"items": [serialize_project(row) for row in rows]}


@app.route("/finance/projects", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
def create_project():
    user = get_user_from_token()
    data = request.get_json(silent=True) or {}
    company = resolve_company_for_user(user, data.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    project_code = (data.get("project_code") or "").strip()
    name = (data.get("name") or "").strip()
    if not project_code or not name:
        return error_response("project_code and name are required")
    if Project.query.filter_by(company_id=company.id, project_code=project_code).first():
        return error_response("project_code already exists", 409)

    try:
        project = Project(
            org_id=company.org_id,
            company_id=company.id,
            project_code=project_code,
            name=name,
            customer_name=(data.get("customer_name") or "").strip() or None,
            status=(data.get("status") or "active").strip().lower(),
            budget_revenue=parse_money(data.get("budget_revenue", 0), "budget_revenue"),
            budget_cost=parse_money(data.get("budget_cost", 0), "budget_cost"),
            notes=(data.get("notes") or "").strip() or None,
        )
    except ValueError as exc:
        return error_response(str(exc))

    db.session.add(project)
    if not safe_commit():
        return error_response("database error while creating project", 503)

    log(user.id, f"created project {project.project_code}")
    return serialize_project(project), 201


@app.route("/finance/projects/costs", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant", "cashier")
def create_project_cost():
    user = get_user_from_token()
    data = request.get_json(silent=True) or {}
    company = resolve_company_for_user(user, data.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    project = Project.query.filter_by(id=data.get("project_id"), company_id=company.id).first()
    if not project:
        return error_response("project not found", 404)

    entry_type = (data.get("entry_type") or "cost").strip().lower()
    if entry_type not in {"cost", "revenue"}:
        return error_response("entry_type must be cost or revenue")

    try:
        entry = ProjectCostEntry(
            org_id=company.org_id,
            company_id=company.id,
            project_id=project.id,
            entry_type=entry_type,
            description=(data.get("description") or "").strip() or f"{entry_type.title()} entry",
            amount=parse_money(data.get("amount", 0), "amount"),
            reference=(data.get("reference") or "").strip() or None,
            work_date=parse_iso_date(data.get("work_date"), "work_date", today_utc_date()),
        )
    except ValueError as exc:
        return error_response(str(exc))

    db.session.add(entry)
    if not safe_commit():
        return error_response("database error while creating project cost entry", 503)

    log(user.id, f"posted {entry_type} entry to project {project.project_code}")
    return serialize_project_cost_entry(entry), 201


@app.route("/finance/projects/summary")
@jwt_required()
def project_summary():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)
    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)
    return build_project_summary(company)


@app.route("/finance/accountant/toolkit")
@jwt_required()
def accountant_toolkit():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)
    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)
    return build_accountant_toolkit(company)


@app.route("/finance/integrations")
@jwt_required()
def list_integrations():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)
    company = resolve_company_for_user(user, request.args.get("company_id"))
    if not company:
        return error_response("company not found", 404)
    seed_integration_connections(company)
    safe_commit()
    rows = IntegrationConnection.query.filter_by(company_id=company.id).order_by(IntegrationConnection.provider.asc()).all()
    return {"items": [serialize_integration_connection(row) for row in rows]}


@app.route("/finance/integrations", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
@plan_required("pro")
def connect_integration():
    user = get_user_from_token()
    data = request.get_json(silent=True) or {}
    company = resolve_company_for_user(user, data.get("company_id"))
    if not company:
        return error_response("company not found", 404)

    provider = (data.get("provider") or "").strip().lower()
    if not provider:
        return error_response("provider is required")

    connection = IntegrationConnection.query.filter_by(company_id=company.id, provider=provider).first()
    if not connection:
        category = next((item["category"] for item in INTEGRATION_CATALOG if item["provider"] == provider), "other")
        connection = IntegrationConnection(
            org_id=company.org_id,
            company_id=company.id,
            provider=provider,
            category=category,
        )
        db.session.add(connection)

    connection.status = "connected"
    connection.config_json = json.dumps(data.get("config") or {"connected_via": "workspace"})
    connection.last_synced_at = datetime.datetime.now(datetime.UTC)
    if not safe_commit():
        return error_response("database error while connecting integration", 503)

    log(user.id, f"connected integration {provider}")
    return serialize_integration_connection(connection), 201


@app.route("/finance/integrations/<int:integration_id>/sync", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
@plan_required("pro")
def sync_integration(integration_id):
    user = get_user_from_token()
    connection = IntegrationConnection.query.filter_by(id=integration_id, org_id=user.org_id).first()
    if not connection:
        return error_response("integration not found", 404)

    connection.status = "connected"
    connection.last_synced_at = datetime.datetime.now(datetime.UTC)
    if not safe_commit():
        return error_response("database error while syncing integration", 503)

    log(user.id, f"synced integration {connection.provider}")
    return serialize_integration_connection(connection)


# ---------------- API KEY ----------------

@app.route("/apikey", methods=["POST"])
@jwt_required()
def create_key():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    raw = os.urandom(24).hex()

    db.session.add(APIKey(org_id=user.org_id, key_hash=hash_key(raw)))
    if not safe_commit():
        return error_response("database error while creating API key", 503)

    log(user.id, "created api key")
    return {"api_key": raw}


# ---------------- ANALYTICS ----------------

@app.route("/me")
@jwt_required()
def me():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    memberships = ensure_user_company_memberships(user)
    company = resolve_company_for_user(user, user.default_company_id)
    org = db.session.get(Organization, user.org_id)
    payload = serialize_user(user)
    payload.update(
        {
            "default_company_id": company.id if company else user.default_company_id,
            "business_type": company.business_type if company else "sole_proprietor",
            "accessible_company_count": len(memberships),
            "subscription": serialize_subscription(org) if org else serialize_subscription(Organization(name="", usage=0)),
        }
    )
    return payload


@app.route("/me", methods=["DELETE"])
@jwt_required()
@limiter.limit("5 per hour")
def delete_my_account():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    data = request.get_json(silent=True) or {}
    password = data.get("password") or ""
    if not password:
        return error_response("password is required")
    if not bcrypt.check_password_hash(user.password, password):
        return error_response("invalid password", 401)

    if user.role == "owner":
        owner_count = User.query.filter_by(org_id=user.org_id, role="owner").count()
        if owner_count <= 1:
            return error_response("create another owner before deleting this account", 400)

    db.session.add(AuditLog(user_id=user.id, company_id=user.default_company_id, action="deleted own account"))
    PasswordResetToken.query.filter_by(user_id=user.id).delete()
    ActiveSession.query.filter_by(user_id=user.id).delete()
    UserCompanyMembership.query.filter_by(user_id=user.id).delete()
    db.session.delete(user)
    if not safe_commit():
        return error_response("database error while deleting account", 503)

    return {"msg": "account deleted"}


@app.route("/analytics")
@jwt_required()
def analytics():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    org = db.session.get(Organization, user.org_id)
    if not org:
        return error_response("organization not found", 404)

    return {
        "usage": org.usage,
        "reports": Report.query.filter_by(org_id=org.id).count(),
        "users": User.query.filter_by(org_id=org.id).count(),
        "active_users": active_user_count_for_org(org.id),
    }


@app.route("/dashboard")
@jwt_required()
def dashboard_summary():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    org = db.session.get(Organization, user.org_id)
    if not org:
        return error_response("organization not found", 404)

    company = resolve_company_for_user(user, request.args.get("company_id"))
    report_totals = aggregate_org_reports(org.id, company.id if company else None)
    return {
        "sales": report_totals["revenue"],
        "expenses": report_totals["expenses"],
        "profit": report_totals["profit"],
        "inventory_value": report_totals["total_assets"],
        "active_users": active_user_count_for_org(org.id),
        "company_id": company.id if company else None,
    }


@app.route("/reports/income")
@jwt_required()
@roles_required("owner", "admin", "manager", "accountant")
def income_statement():
    user = get_user_from_token()
    org = db.session.get(Organization, user.org_id)
    if not org:
        return error_response("organization not found", 404)

    totals = aggregate_org_reports(org.id)
    return {
        "revenue": totals["revenue"],
        "expenses": totals["expenses"],
        "profit": totals["profit"],
    }


@app.route("/inventory")
@jwt_required()
def inventory_summary():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    org = db.session.get(Organization, user.org_id)
    if not org:
        return error_response("organization not found", 404)

    totals = aggregate_org_reports(org.id)
    return {
        "inventory_value": totals["total_assets"],
        "note": "Derived from uploaded report assets. Add a dedicated inventory table for item-level stock.",
    }


# ---------------- LIVE USER COUNT ----------------

@app.route("/user-count")
@jwt_required()
def user_count():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    org = db.session.get(Organization, user.org_id)
    if not org:
        return error_response("organization not found", 404)

    count = active_user_count_for_org(org.id)
    registered = User.query.filter_by(org_id=org.id).count()
    return {"user_count": count, "active_users": count, "registered_users": registered}


@app.route("/session/ping", methods=["POST"])
@jwt_required()
def session_ping():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    touch_session(user.id)
    return {"ok": True}


@app.route("/logout", methods=["POST"])
@jwt_required()
def logout():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    clear_session(user.id)
    log(user.id, "logged out")
    return {"ok": True}


@app.route("/activity/recent")
@jwt_required()
def recent_activity():
    user = get_user_from_token()
    if not user:
        return error_response("invalid token", 401)

    try:
        limit = int(request.args.get("limit", "8"))
    except ValueError:
        limit = 8
    limit = max(1, min(limit, 50))

    rows = (
        db.session.query(AuditLog, User.email)
        .join(User, AuditLog.user_id == User.id)
        .filter(User.org_id == user.org_id)
        .order_by(AuditLog.time.desc())
        .limit(limit)
        .all()
    )

    return {
        "items": [
            {
                "email": email,
                "action": log_row.action,
                "time": log_row.time.isoformat() if log_row.time else None,
            }
            for log_row, email in rows
        ]
    }


# ---------------- ADMIN USERS ----------------

@app.route("/admin/users")
@jwt_required()
@roles_required("owner", "admin")
def users():
    me = get_user_from_token()
    return jsonify([serialize_user(user) for user in User.query.filter_by(org_id=me.org_id).order_by(User.id.asc())])


@app.route("/admin/users", methods=["POST"])
@jwt_required()
@roles_required("owner", "admin")
@limiter.limit("20 per minute")
def create_user():
    me = get_user_from_token()
    data = request.get_json(silent=True) or {}
    email = (data.get("email") or "").strip().lower()
    password = data.get("password") or ""
    role = (data.get("role") or "cashier").strip().lower()

    if not email or not password:
        return error_response("email and password are required")
    if len(password) < 8:
        return error_response("password must be at least 8 characters")
    if role not in VALID_ROLES:
        return error_response("invalid role")
    if User.query.filter_by(email=email).first():
        return error_response("email already exists", 409)

    fallback_company = resolve_company_for_user(me, data.get("company_id")) or resolve_company_for_user(me)
    try:
        membership_specs = normalize_membership_specs(
            data.get("memberships"),
            data.get("company_ids"),
            fallback_company.id if fallback_company else None,
            role,
        )
    except ValueError as exc:
        return error_response(str(exc))

    company_map = membership_company_map([spec["company_id"] for spec in membership_specs])
    if len(company_map) != len(membership_specs) or any(company.org_id != me.org_id for company in company_map.values()):
        return error_response("company not found", 404)

    hashed_password = bcrypt.generate_password_hash(password).decode()
    default_company_id = next((spec["company_id"] for spec in membership_specs if spec["is_default"]), membership_specs[0]["company_id"])
    user = User(email=email, password=hashed_password, role=role, org_id=me.org_id, default_company_id=default_company_id)
    db.session.add(user)
    db.session.flush()
    apply_user_company_memberships(user, membership_specs)
    if not safe_commit():
        return error_response("database error while creating user", 503)

    log(me.id, f"created user {email} with role {role}")
    return {"msg": "user created", "user": serialize_user(user)}, 201


@app.route("/admin/users/<int:user_id>/role", methods=["PATCH"])
@jwt_required()
@roles_required("owner", "admin")
@limiter.limit("30 per minute")
def update_user_role(user_id):
    me = get_user_from_token()
    data = request.get_json(silent=True) or {}
    role = (data.get("role") or "").strip().lower()
    if role not in VALID_ROLES:
        return error_response("invalid role")

    target = User.query.filter_by(id=user_id, org_id=me.org_id).first()
    if not target:
        return error_response("user not found", 404)
    if target.id == me.id and role not in {"owner", "admin"}:
        return error_response("cannot downgrade your own admin access", 400)
    if target.role == "owner" and role != "owner":
        owner_count = User.query.filter_by(org_id=me.org_id, role="owner").count()
        if owner_count <= 1:
            return error_response("create another owner before removing the final owner", 400)

    target.role = role
    memberships = UserCompanyMembership.query.filter_by(user_id=target.id).all()
    for membership in memberships:
        membership.role = role
    if not safe_commit():
        return error_response("database error while updating role", 503)

    log(me.id, f"updated role for {target.email} to {role}")
    return {"msg": "role updated"}


@app.route("/admin/users/<int:user_id>/companies", methods=["PUT"])
@jwt_required()
@roles_required("owner", "admin")
@limiter.limit("20 per minute")
def update_user_companies(user_id):
    me = get_user_from_token()
    target = User.query.filter_by(id=user_id, org_id=me.org_id).first()
    if not target:
        return error_response("user not found", 404)

    data = request.get_json(silent=True) or {}
    try:
        membership_specs = normalize_membership_specs(
            data.get("memberships"),
            data.get("company_ids"),
            parse_company_id(data.get("default_company_id")) or target.default_company_id or me.default_company_id,
            target.role,
        )
    except ValueError as exc:
        return error_response(str(exc))

    company_map = membership_company_map([spec["company_id"] for spec in membership_specs])
    if len(company_map) != len(membership_specs) or any(company.org_id != me.org_id for company in company_map.values()):
        return error_response("company not found", 404)

    try:
        apply_user_company_memberships(target, membership_specs)
    except ValueError as exc:
        return error_response(str(exc))
    if not safe_commit():
        return error_response("database error while updating company access", 503)

    log(me.id, f"updated company access for {target.email}")
    return {"msg": "company access updated", "user": serialize_user(target)}


@app.route("/admin/users/<int:user_id>", methods=["DELETE"])
@jwt_required()
@roles_required("owner", "admin")
@limiter.limit("20 per minute")
def delete_user(user_id):
    me = get_user_from_token()
    target = User.query.filter_by(id=user_id, org_id=me.org_id).first()
    if not target:
        return error_response("user not found", 404)
    if target.id == me.id:
        return error_response("cannot delete your own account", 400)
    if target.role == "owner":
        owner_count = User.query.filter_by(org_id=me.org_id, role="owner").count()
        if owner_count <= 1:
            return error_response("create another owner before deleting the final owner", 400)

    clear_session(target.id)
    PasswordResetToken.query.filter_by(user_id=target.id).delete()
    UserCompanyMembership.query.filter_by(user_id=target.id).delete()
    db.session.delete(target)
    if not safe_commit():
        return error_response("database error while deleting user", 503)

    log(me.id, f"deleted user {target.email}")
    return {"msg": "user deleted"}


# ---------------- HEALTH ----------------

@app.route("/")
def home():
    return {"status": "FULL SAAS RUNNING"}


@app.route("/health")
def health():
    try:
        db.session.execute(text("SELECT 1"))
        return {"status": "ok", "env": ENV, "database": "ok"}
    except SQLAlchemyError:
        db.session.rollback()
        return {"status": "degraded", "env": ENV, "database": "error"}, 503


@app.route("/system-status")
def system_status():
    return maintenance_state()


# ---------------- RUN ----------------

if __name__ == "__main__":
    port = int(os.getenv("PORT", "5000"))
    debug = ENV == "development"
    app.run(host="0.0.0.0", port=port, debug=debug)
